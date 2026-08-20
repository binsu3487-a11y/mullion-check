from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Dict, List, Optional, Sequence, Tuple
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.ticker as ticker
import numpy as np
import pandas as pd
import streamlit as st

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:
    Document = None  # type: ignore[assignment]
    DOCX_IMPORT_ERROR: Optional[Exception] = exc
else:
    DOCX_IMPORT_ERROR = None

try:
    import openseespy.opensees as ops
    import opsvis as opsv
except ImportError as exc:
    ops = None  # type: ignore[assignment]
    opsv = None  # type: ignore[assignment]
    OPENSEES_IMPORT_ERROR: Optional[Exception] = exc
else:
    OPENSEES_IMPORT_ERROR = None


# ============================================================
# 1. 固定設定（kgf、cm）
# ============================================================

ELEMENT_COUNT = 40
POSTPROCESS_POINTS_PER_ELEMENT = 41
KPA_TO_KGF_PER_CM2 = 1.0 / 98.0665
DEFAULT_GLASS_DENSITY_KG_M3 = 2500.0
WORD_FIGURE_HEIGHT_CM = 7.3
STRONG_DEFLECTION_DIVISOR = 240.0
WEAK_DEFLECTION_LIMIT_CM = 0.300  # 3 mm

MATERIAL_DATABASE: Dict[str, Dict[str, object]] = {
    "鋁合金": {
        "E": 710100.0,
        "grades": {
            "6105-T5": {"Fty": 2450.0, "Ftu": 2660.0, "Fb_allow": 1485.0},
            "6061-T6": {"Fty": 2450.0, "Ftu": 2660.0, "Fb_allow": 1485.0},
            "6063-T5": {"Fty": 1050.0, "Ftu": 1470.0, "Fb_allow": 682.0},
            "6063-T6": {"Fty": 1750.0, "Ftu": 2100.0, "Fb_allow": 1060.0},
            "3003-H14": {
                "Fty": 1195.218,
                "Ftu": 1406.139,
                "Fb_allow": 717.131,
            },
        },
    },
    "碳鋼": {
        "E": 2039432.426,
        "grades": {
            "A36": {"Fty": 2531.050, "Ftu": 4077.804, "Fb_allow": 1670.493},
        },
    },
}

SETTING_BLOCK_OPTIONS = {
    "距支承 L/4": 4.0,
    "距支承 L/6": 6.0,
    "距支承 L/8": 8.0,
    "自訂距離": None,
}

WIND_GLASS_MODE_OPTIONS = {
    "單片玻璃": "single",
    "上下同尺寸玻璃": "same_double",
    "上下玻璃獨立輸入": "separate",
}

WIND_GLASS_MODE_DESCRIPTIONS = {
    "single": "風壓只計一片玻璃的 tributary load；玻璃自重亦計此片玻璃。",
    "same_double": "上、下玻璃尺寸相同；風壓採兩片玻璃 tributary load 疊加，自重仍只計一片玻璃。",
    "separate": "上、下玻璃尺寸分別輸入；風壓分別建立 tributary load 後疊加，自重只計上方玻璃。",
}


# ============================================================
# 2. 共用函式
# ============================================================


def status_text(passed: bool) -> str:
    return "PASS（通過）" if passed else "FAIL（不通過）"


def trapezoid_integral(y, x) -> float:
    """相容新版與舊版 NumPy 的梯形積分。

    NumPy 2.4 已移除 np.trapz；NumPy 2.0+ 使用 np.trapezoid。
    若執行環境仍是較舊 NumPy，則自動退回 np.trapz。
    """
    if hasattr(np, "trapezoid"):
        return float(np.trapezoid(y, x))
    return float(np.trapz(y, x))


def clear_analysis_result() -> None:
    st.session_state.pop("analysis_result", None)


def figure_to_png_bytes(figure: plt.Figure) -> bytes:
    buffer = BytesIO()
    figure.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    buffer.seek(0)
    return buffer.getvalue()


def dataframe_to_csv_bytes(dataframe: pd.DataFrame) -> bytes:
    return dataframe.round(3).to_csv(index=False, float_format="%.3f").encode("utf-8-sig")


def round_dataframe_for_display(dataframe: pd.DataFrame) -> pd.DataFrame:
    """Streamlit 表格中的數值至多顯示小數點後 3 位。"""
    formatted = dataframe.copy()

    def format_value(value: object) -> object:
        if isinstance(value, (float, np.floating)):
            number = float(value)
            if not np.isfinite(number):
                return "-"
            if abs(number) < 0.0005:
                number = 0.0
            return f"{number:.3f}".rstrip("0").rstrip(".")
        return value

    for column in formatted.columns:
        formatted[column] = formatted[column].map(format_value)
    return formatted


def sanitize_report_filename(raw_name: str) -> str:
    name = (raw_name or "").strip()
    if name.lower().endswith(".docx"):
        name = name[:-5]
    name = re.sub(r'[<>:"/\\|?*]+', "_", name)
    name = name.strip(" .")
    if not name:
        name = "transom_analysis_report"
    return f"{name}.docx"


def wind_allowable_deflection(length_cm: float) -> Tuple[float, str]:
    """橫料受風壓方向之容許變形：L/240。"""
    return length_cm / STRONG_DEFLECTION_DIVISOR, "L/240"


def glass_self_weight_kgf(
    glass_width_cm: float,
    glass_height_cm: float,
    total_glass_thickness_mm: float,
    glass_density_kg_m3: float,
) -> float:
    thickness_cm = total_glass_thickness_mm / 10.0
    volume_cm3 = glass_width_cm * glass_height_cm * thickness_cm
    return volume_cm3 * glass_density_kg_m3 / 1_000_000.0


def tributary_width_cm(
    x_cm: np.ndarray | float,
    transom_length_cm: float,
    glass_width_cm: float,
    glass_height_cm: float,
) -> np.ndarray:
    """龜殼式 45 度 tributary width。

    風壓分析以橫料「淨支承跨距」0~L 為梁座標，再按相同比例映射到
    玻璃水平座標 0~Wg：

        x_glass = (x / L) * Wg
        b(x) = min(x_glass, Wg - x_glass, Hg/2)

    這樣即使玻璃寬度 Wg 與淨跨 L 不同，支承兩端仍有 b(0)=b(L)=0，
    而中央最大 tributary width 仍由實際玻璃尺寸 Wg、Hg 決定。
    """
    x = np.asarray(x_cm, dtype=float)
    L = float(transom_length_cm)
    W = float(glass_width_cm)
    H = float(glass_height_cm)

    if L <= 0.0 or W <= 0.0 or H <= 0.0:
        return np.zeros_like(x, dtype=float)

    # 將橫料淨跨座標映射到玻璃寬度座標。
    xi = np.clip(x / L, 0.0, 1.0)
    x_glass = xi * W

    width = np.minimum.reduce(
        [
            x_glass,
            W - x_glass,
            np.full_like(x_glass, H / 2.0),
        ]
    )
    width = np.maximum(width, 0.0)

    # 超出實際梁跨範圍時不施加載重。
    outside = (x < 0.0) | (x > L)
    if np.any(outside):
        width = np.where(outside, 0.0, width)
    return width


def exact_wind_line_load_kgf_per_cm(
    x_cm: np.ndarray | float,
    wind_pressure_kpa: float,
    transom_length_cm: float,
    glass_width_cm: float,
    glass_height_cm: float,
    tributary_factor: float = 1.0,
) -> np.ndarray:
    tributary = tributary_width_cm(
        x_cm,
        transom_length_cm,
        glass_width_cm,
        glass_height_cm,
    )
    return (
        float(tributary_factor)
        * wind_pressure_kpa
        * KPA_TO_KGF_PER_CM2
        * tributary
    )


def wind_load_shape_text(glass_width_cm: float, glass_height_cm: float) -> str:
    W = float(glass_width_cm)
    H = float(glass_height_cm)
    if np.isclose(W, H, rtol=0.0, atol=1.0e-9):
        return "三角形"
    return "梯形" if W > H else "三角形"


def wind_mode_label(mode: str) -> str:
    for label, value in WIND_GLASS_MODE_OPTIONS.items():
        if value == mode:
            return label
    return str(mode)


def apply_point_load_to_beam(
    point_x_cm: float,
    point_load_kgf: float,
    node_x: np.ndarray,
) -> None:
    """將集中載重精確放在節點，或以 -beamPoint 放入所在元素。"""
    if ops is None:
        raise RuntimeError("分析套件尚未載入。")

    distances = np.abs(node_x - point_x_cm)
    nearest_index = int(np.argmin(distances))
    if distances[nearest_index] <= 1.0e-9:
        node_tag = nearest_index + 1
        ops.load(node_tag, 0.0, -point_load_kgf, 0.0)
        return

    element_index = int(np.searchsorted(node_x, point_x_cm, side="right") - 1)
    element_index = max(0, min(element_index, len(node_x) - 2))
    x_i = float(node_x[element_index])
    x_j = float(node_x[element_index + 1])
    xL = (point_x_cm - x_i) / (x_j - x_i)
    element_tag = element_index + 1
    ops.eleLoad(
        "-ele",
        element_tag,
        "-type",
        "-beamPoint",
        -point_load_kgf,
        xL,
        0.0,
    )


def _force_plot_formats() -> Tuple[Dict[str, object], Dict[str, object]]:
    main = {
        "color": "blue",
        "linestyle": "solid",
        "linewidth": 1.8,
        "marker": "",
        "markersize": 0,
    }
    reference = {
        "color": "blue",
        "linestyle": "solid",
        "linewidth": 0.6,
        "marker": "",
        "markersize": 0,
    }
    return main, reference


def build_analysis_node_coordinates(
    transom_length_cm: float,
    event_x_coordinates: Sequence[float] = (),
) -> np.ndarray:
    """建立分析節點；setting block 等重要位置會直接成為節點。"""
    L = float(transom_length_cm)
    values = [float(value) for value in np.linspace(0.0, L, ELEMENT_COUNT + 1)]
    values.extend([0.0, L])
    values.extend(float(value) for value in event_x_coordinates)
    values = sorted(value for value in values if -1.0e-9 <= value <= L + 1.0e-9)

    unique: List[float] = []
    for value in values:
        value = min(max(value, 0.0), L)
        if not unique or abs(value - unique[-1]) > 1.0e-9:
            unique.append(value)
    return np.asarray(unique, dtype=float)


def build_piecewise_wind_loads_for_nodes(
    node_x: np.ndarray,
    wind_q_function,
) -> np.ndarray:
    """依每一梁元素範圍，將 tributary line load 轉為該元素的等值均布載重。"""
    q_elements = np.zeros(len(node_x) - 1, dtype=float)
    for index in range(len(node_x) - 1):
        x_i = float(node_x[index])
        x_j = float(node_x[index + 1])
        sample_x = np.linspace(x_i, x_j, 9)
        sample_q = np.asarray(wind_q_function(sample_x), dtype=float)
        q_elements[index] = trapezoid_integral(sample_q, sample_x) / (x_j - x_i)
    return q_elements


def extract_opensees_beam_response(
    node_x: np.ndarray,
    elastic_modulus_kgf_cm2: float,
    inertia_cm4: float,
    q_elements: Optional[np.ndarray],
    point_loads: Sequence[Tuple[float, float]],
) -> Dict[str, object]:
    """由 OpenSees 求得的節點位移／轉角重建梁內部反應。

    分布載重在每個元素內為常數，因此以 Euler-Bernoulli 梁形函數加上
    均布載重 particular solution 重建變形、彎矩與剪力。集中載重位置已
    被建立為節點，因此不會落在元素內部。
    """
    if ops is None:
        raise RuntimeError("分析套件尚未載入。")

    E = float(elastic_modulus_kgf_cm2)
    I = float(inertia_cm4)
    node_tags = list(range(1, len(node_x) + 1))
    uy = {tag: float(ops.nodeDisp(tag, 2)) for tag in node_tags}
    rz = {tag: float(ops.nodeDisp(tag, 3)) for tag in node_tags}

    all_x: List[float] = []
    all_v: List[float] = []
    all_m: List[float] = []
    all_s: List[float] = []
    all_q: List[float] = []

    for element_index in range(len(node_x) - 1):
        node_i = element_index + 1
        node_j = element_index + 2
        x_i = float(node_x[element_index])
        x_j = float(node_x[element_index + 1])
        length = x_j - x_i
        if length <= 0.0:
            raise RuntimeError("分析網格出現零長度元素。")

        v_i = uy[node_i]
        v_j = uy[node_j]
        theta_i = rz[node_i]
        theta_j = rz[node_j]
        xi = np.linspace(0.0, 1.0, POSTPROCESS_POINTS_PER_ELEMENT)

        # OpenSees 2D 梁沿 +X 建立時，local y 與 global Y 同向。
        q_down = 0.0 if q_elements is None else float(q_elements[element_index])
        q_local_y = -q_down

        n1 = 1.0 - 3.0 * xi**2 + 2.0 * xi**3
        n2 = xi - 2.0 * xi**2 + xi**3
        n3 = 3.0 * xi**2 - 2.0 * xi**3
        n4 = -xi**2 + xi**3

        load_bubble = (
            q_local_y
            * length**4
            / (24.0 * E * I)
            * xi**2
            * (1.0 - xi) ** 2
        )
        v_values = (
            n1 * v_i
            + n2 * length * theta_i
            + n3 * v_j
            + n4 * length * theta_j
            + load_bubble
        )
        x_values = x_i + xi * length

        curvature = (
            (-6.0 + 12.0 * xi) / length**2 * v_i
            + (-4.0 + 6.0 * xi) / length * theta_i
            + (6.0 - 12.0 * xi) / length**2 * v_j
            + (-2.0 + 6.0 * xi) / length * theta_j
            + q_local_y
            * length**2
            / (24.0 * E * I)
            * (2.0 - 12.0 * xi + 12.0 * xi**2)
        )
        moment_values = E * I * curvature

        third_derivative = (
            12.0 / length**3 * v_i
            + 6.0 / length**2 * theta_i
            - 12.0 / length**3 * v_j
            + 6.0 / length**2 * theta_j
            + q_local_y
            * length
            / (24.0 * E * I)
            * (-12.0 + 24.0 * xi)
        )
        shear_values = E * I * third_derivative

        # 保留元素交界兩側的值，才能呈現集中載重位置的剪力跳躍。
        all_x.extend(x_values.tolist())
        all_v.extend(v_values.tolist())
        all_m.extend(moment_values.tolist())
        all_s.extend(shear_values.tolist())
        all_q.extend(np.full_like(x_values, q_down).tolist())

    x_array = np.asarray(all_x, dtype=float)
    v_array = np.asarray(all_v, dtype=float)
    m_array = np.asarray(all_m, dtype=float)
    s_array = np.asarray(all_s, dtype=float)
    q_array = np.asarray(all_q, dtype=float)

    max_disp_index = int(np.argmax(np.abs(v_array)))
    max_moment_index = int(np.argmax(np.abs(m_array)))
    max_shear_index = int(np.argmax(np.abs(s_array)))

    ops.reactions()
    left_reaction = float(ops.nodeReaction(1, 2))
    right_reaction = float(ops.nodeReaction(len(node_x), 2))

    distributed_load = 0.0
    if q_elements is not None:
        distributed_load = float(
            np.sum(np.asarray(q_elements, dtype=float) * np.diff(node_x))
        )
    point_load_total = float(sum(float(P) for _, P in point_loads))

    return {
        "x": x_array,
        "q": q_array,
        "deflection": v_array,
        "moment": m_array,
        "shear": s_array,
        "total_distributed_load": distributed_load,
        "total_point_load": point_load_total,
        "total_load": distributed_load + point_load_total,
        "left_reaction": left_reaction,
        "right_reaction": right_reaction,
        "max_displacement": float(v_array[max_disp_index]),
        "max_abs_displacement": abs(float(v_array[max_disp_index])),
        "max_displacement_x": float(x_array[max_disp_index]),
        "max_moment": float(m_array[max_moment_index]),
        "max_abs_moment": abs(float(m_array[max_moment_index])),
        "max_moment_x": float(x_array[max_moment_index]),
        "max_shear": float(s_array[max_shear_index]),
        "max_abs_shear": abs(float(s_array[max_shear_index])),
        "max_shear_x": float(x_array[max_shear_index]),
    }


def run_opensees_axis_model(
    axis_name: str,
    transom_length_cm: float,
    elastic_modulus_kgf_cm2: float,
    area_cm2: float,
    inertia_cm4: float,
    q_function,
    point_loads: Sequence[Tuple[float, float]],
    exact_q_for_plot=None,
    mesh_mode: str = "uniform",
) -> Dict[str, object]:
    if ops is None or opsv is None:
        raise RuntimeError("必要分析套件尚未安裝。")

    L = float(transom_length_cm)

    # 風壓需要沿梁長離散分布載重；玻璃自重只有兩個 setting block
    # 集中力，因此弱軸採「支承 + 集中力位置」的必要節點即可。
    # 這樣每一支梁元素內皆無集中力，OpenSees 的節點載重與元素內力
    # 關係最直接，也避免不必要的離散節點。
    if mesh_mode == "point_load_events":
        event_values = [0.0, L] + [float(xp) for xp, _ in point_loads]
        node_x = np.asarray(sorted(set(round(value, 12) for value in event_values)), dtype=float)
        if len(node_x) < 2 or abs(float(node_x[0])) > 1.0e-9 or abs(float(node_x[-1]) - L) > 1.0e-9:
            raise RuntimeError(f"{axis_name} 的集中載重節點配置錯誤。")
    elif mesh_mode == "uniform":
        node_x = build_analysis_node_coordinates(
            L,
            [float(xp) for xp, _ in point_loads],
        )
    else:
        raise ValueError(f"未知的網格模式：{mesh_mode}")

    number_of_elements = len(node_x) - 1
    q_elements = (
        None
        if q_function is None
        else build_piecewise_wind_loads_for_nodes(node_x, q_function)
    )

    ops.wipe()
    ops.model("basic", "-ndm", 2, "-ndf", 3)
    ops.geomTransf("Linear", 1)

    for index, x_value in enumerate(node_x, start=1):
        ops.node(index, float(x_value), 0.0)

    # 左端 Pin：UX/UY 固定、RZ 自由；右端 Roller：UY 固定。
    ops.fix(1, 1, 1, 0)
    ops.fix(len(node_x), 0, 1, 0)

    for element_index in range(number_of_elements):
        ops.element(
            "elasticBeamColumn",
            element_index + 1,
            element_index + 1,
            element_index + 2,
            area_cm2,
            elastic_modulus_kgf_cm2,
            inertia_cm4,
            1,
        )

    ops.timeSeries("Linear", 1)
    ops.pattern("Plain", 1, 1)

    if q_elements is not None:
        for element_index, q_value in enumerate(q_elements, start=1):
            if abs(float(q_value)) > 1.0e-14:
                ops.eleLoad(
                    "-ele",
                    element_index,
                    "-type",
                    "-beamUniform",
                    -float(q_value),
                    0.0,
                )

    # setting block 位置已被加入 node_x，因此集中力會直接施加在節點上。
    for point_x, point_load in point_loads:
        apply_point_load_to_beam(
            float(point_x),
            float(point_load),
            node_x,
        )

    # 此模型只有單點支承拘束；弱軸集中載重模型使用 Plain handler
    # 即可，風壓模型則保留原本的 Transformation handler。
    if mesh_mode == "point_load_events":
        ops.constraints("Plain")
    else:
        ops.constraints("Transformation")
    ops.numberer("RCM")
    ops.system("BandGeneral")
    ops.algorithm("Linear")
    ops.integrator("LoadControl", 1.0)
    ops.analysis("Static")

    analysis_code = ops.analyze(1)
    if analysis_code != 0:
        raise RuntimeError(f"{axis_name}分析失敗，代碼 = {analysis_code}。")

    response = extract_opensees_beam_response(
        node_x=node_x,
        elastic_modulus_kgf_cm2=elastic_modulus_kgf_cm2,
        inertia_cm4=inertia_cm4,
        q_elements=q_elements,
        point_loads=point_loads,
    )

    # 基本靜力平衡檢查。若 OpenSees 反力與施加載重不平衡，
    # 直接停止，不讓錯誤結果進入畫面或 Word 報告。
    total_downward_load = float(response["total_load"])
    reaction_sum = float(response["left_reaction"]) + float(response["right_reaction"])
    balance_error = abs(reaction_sum - total_downward_load)
    balance_tolerance = max(1.0e-5, 1.0e-6 * max(total_downward_load, 1.0))
    if balance_error > balance_tolerance:
        raise RuntimeError(
            f"{axis_name} 力平衡檢查失敗："
            f"左反力 + 右反力 = {reaction_sum:.3f} kgf，"
            f"總外力 = {total_downward_load:.3f} kgf。"
        )
    response["reaction_sum"] = reaction_sum
    response["force_balance_error"] = balance_error

    node_displacements = np.asarray(
        [float(ops.nodeDisp(tag, 2)) for tag in range(1, len(node_x) + 1)],
        dtype=float,
    )
    node_rotations = np.asarray(
        [float(ops.nodeDisp(tag, 3)) for tag in range(1, len(node_x) + 1)],
        dtype=float,
    )

    left_reaction_fe = float(response["left_reaction"])
    right_reaction_fe = float(response["right_reaction"])

    max_abs_disp = float(response["max_abs_displacement"])
    max_abs_moment = float(response["max_abs_moment"])
    max_abs_shear = float(response["max_abs_shear"])

    deformation_target = max(0.08 * L, 1.0)
    force_target = max(0.10 * L, 1.0)
    sfac_defo = deformation_target / max_abs_disp if max_abs_disp > 1.0e-12 else 1.0
    sfac_m = force_target / max_abs_moment if max_abs_moment > 1.0e-12 else 1.0
    sfac_v = force_target / max_abs_shear if max_abs_shear > 1.0e-12 else 1.0

    fmt_main, fmt_ref = _force_plot_formats()

    # --------------------------------------------------------
    # 模型與載重圖
    # --------------------------------------------------------
    fig_model, ax_model = plt.subplots(figsize=(8.6, 3.4))
    opsv.plot_model(
        ax=ax_model,
        node_supports=True,
        node_labels=0,
        element_labels=0,
    )
    ax_model.set_title(f"{axis_name} Model and Loading")
    ax_model.set_xlabel("Transom X Coordinate (cm)")
    ax_model.set_ylabel("Load schematic")

    load_height = max(0.12 * L, 10.0)
    if q_elements is not None and exact_q_for_plot is not None:
        sample_x = np.linspace(0.0, L, 25)
        sample_q = np.asarray(exact_q_for_plot(sample_x), dtype=float)
        q_max = float(np.max(sample_q)) if len(sample_q) else 0.0
        if q_max > 1.0e-12:
            for x_value, q_value in zip(sample_x, sample_q):
                if q_value <= 1.0e-12:
                    continue
                arrow_length = load_height * (0.25 + 0.75 * q_value / q_max)
                ax_model.annotate(
                    "",
                    xy=(x_value, 0.0),
                    xytext=(x_value, arrow_length),
                    arrowprops={"arrowstyle": "->", "linewidth": 0.9},
                )
            ax_model.text(
                0.02,
                0.93,
                f"Wind load, qmax = {q_max:.3f} kgf/cm",
                transform=ax_model.transAxes,
                ha="left",
                va="top",
                fontsize=9,
                bbox={"boxstyle": "round,pad=0.2", "facecolor": "white", "alpha": 0.9},
            )
    else:
        for point_x, point_load in point_loads:
            ax_model.annotate(
                "",
                xy=(float(point_x), 0.0),
                xytext=(float(point_x), load_height),
                arrowprops={"arrowstyle": "-|>", "linewidth": 1.5},
            )
            ax_model.text(
                float(point_x),
                load_height * 1.03,
                f"P={float(point_load):.3f} kgf",
                ha="center",
                va="bottom",
                fontsize=9,
            )

    ax_model.set_xlim(-0.03 * L, 1.03 * L)
    ax_model.set_ylim(-0.18 * load_height, 1.35 * load_height)
    ax_model.grid(True, linestyle="--", alpha=0.25)
    ax_model.set_aspect("auto")
    fig_model.tight_layout()

    # --------------------------------------------------------
    # 變形圖
    # --------------------------------------------------------
    fig_defo, ax_defo = plt.subplots(figsize=(8.6, 4.2))
    opsv.plot_defo(
        sfac=sfac_defo,
        unDefoFlag=1,
        node_supports=True,
        ax=ax_defo,
    )
    max_x = float(response["max_displacement_x"])
    max_disp = float(response["max_displacement"])
    ax_defo.plot(max_x, max_disp * sfac_defo, marker="o", markersize=6, zorder=7)
    ax_defo.annotate(
        f"Max |δ| = {abs(max_disp):.3f} cm\nX = {max_x:.3f} cm",
        xy=(max_x, max_disp * sfac_defo),
        xytext=(10, -28),
        textcoords="offset points",
        arrowprops={"arrowstyle": "->", "linewidth": 0.8},
        fontsize=9,
        bbox={"boxstyle": "round,pad=0.2", "facecolor": "white", "alpha": 0.85},
    )
    ax_defo.set_title(f"{axis_name} Deformed Shape")
    ax_defo.set_xlabel("Transom X Coordinate (cm)")
    ax_defo.yaxis.set_major_formatter(
        ticker.FuncFormatter(lambda value, pos: f"{value / sfac_defo:.3f}")
    )
    ax_defo.set_ylabel("True Deflection (cm)")
    ax_defo.set_xlim(-0.03 * L, 1.03 * L)
    ax_defo.set_aspect("auto")
    ax_defo.grid(True, linestyle="--", alpha=0.25)
    fig_defo.tight_layout()

    # --------------------------------------------------------
    # 彎矩圖
    # --------------------------------------------------------
    fig_moment, ax_moment = plt.subplots(figsize=(8.6, 4.2))
    opsv.section_force_diagram_2d(
        "M",
        sfac=sfac_m,
        nep=17,
        ax=ax_moment,
        fmt_secforce1=fmt_main,
        fmt_secforce2=fmt_ref,
        ref_vert_lines=False,
        end_max_values=False,
        node_supports=True,
        number_format=".3f",
    )
    ax_moment.set_title(
        f"{axis_name} Bending Moment Diagram\n"
        f"Maximum |M| = {max_abs_moment:.3f} kgf-cm"
    )
    ax_moment.set_xlabel("Transom X Coordinate (cm)")
    ax_moment.yaxis.set_major_formatter(
        ticker.FuncFormatter(lambda value, pos: f"{value / sfac_m:.3f}")
    )
    ax_moment.set_ylabel("True Bending Moment (kgf-cm)")
    ax_moment.set_xlim(-0.03 * L, 1.03 * L)
    ax_moment.set_aspect("auto")
    ax_moment.grid(True, linestyle="--", alpha=0.25)
    fig_moment.tight_layout()

    # --------------------------------------------------------
    # 剪力圖
    # --------------------------------------------------------
    fig_shear, ax_shear = plt.subplots(figsize=(8.6, 4.2))
    opsv.section_force_diagram_2d(
        "V",
        sfac=sfac_v,
        nep=17,
        ax=ax_shear,
        fmt_secforce1=fmt_main,
        fmt_secforce2=fmt_ref,
        ref_vert_lines=False,
        end_max_values=False,
        node_supports=True,
        number_format=".3f",
    )
    ax_shear.set_title(
        f"{axis_name} Shear Force Diagram\n"
        f"Maximum |V| = {max_abs_shear:.3f} kgf"
    )
    ax_shear.set_xlabel("Transom X Coordinate (cm)")
    ax_shear.yaxis.set_major_formatter(
        ticker.FuncFormatter(lambda value, pos: f"{value / sfac_v:.3f}")
    )
    ax_shear.set_ylabel("True Shear Force (kgf)")
    ax_shear.set_xlim(-0.03 * L, 1.03 * L)
    ax_shear.set_aspect("auto")
    ax_shear.grid(True, linestyle="--", alpha=0.25)
    fig_shear.tight_layout()

    return {
        "node_x": node_x,
        "q_elements": q_elements,
        "response": response,
        "node_displacements": node_displacements,
        "node_rotations": node_rotations,
        "left_reaction_fe": left_reaction_fe,
        "right_reaction_fe": right_reaction_fe,
        "figures": {
            "model": fig_model,
            "deformation": fig_defo,
            "moment": fig_moment,
            "shear": fig_shear,
        },
    }

def make_tributary_area_figure(
    transom_length_cm: float,
    upper_width_cm: float,
    upper_height_cm: float,
    wind_glass_mode: str,
    lower_width_cm: Optional[float] = None,
    lower_height_cm: Optional[float] = None,
) -> plt.Figure:
    """繪製上、下玻璃的 45° Tributary Area 示意圖。"""
    L = float(transom_length_cm)
    upper_w = float(upper_width_cm)
    upper_h = float(upper_height_cm)

    has_lower = wind_glass_mode in {"same_double", "separate"}
    if wind_glass_mode == "same_double":
        lower_w = upper_w
        lower_h = upper_h
    elif wind_glass_mode == "separate":
        if lower_width_cm is None or lower_height_cm is None:
            raise ValueError("上下玻璃獨立輸入模式需要下方玻璃尺寸。")
        lower_w = float(lower_width_cm)
        lower_h = float(lower_height_cm)
    else:
        lower_w = 0.0
        lower_h = 0.0

    max_h = max(upper_h, lower_h if has_lower else 0.0)
    max_w = max(upper_w, lower_w if has_lower else 0.0, L)
    fig, ax = plt.subplots(figsize=(8.6, 6.2 if has_lower else 4.8))

    def draw_glass(W: float, H: float, y0: float, direction: float, label: str) -> None:
        glass_left = 0.5 * (L - W)
        x_local = np.linspace(0.0, W, 1001)
        x_global = glass_left + x_local
        boundary = np.minimum.reduce(
            [x_local, W - x_local, np.full_like(x_local, H / 2.0)]
        )
        boundary = np.clip(boundary, 0.0, H)

        if direction > 0:
            y_boundary = y0 + boundary
            y_far = y0 + H
            rect_y = y0
            text_y = y0 + 0.72 * H
        else:
            y_boundary = y0 - boundary
            y_far = y0 - H
            rect_y = y0 - H
            text_y = y0 - 0.72 * H

        ax.add_patch(
            patches.Rectangle(
                (glass_left, rect_y), W, H,
                fill=False, linewidth=1.7, edgecolor="black", zorder=5,
            )
        )
        ax.fill_between(
            x_global, y0, y_boundary,
            color="lightskyblue", alpha=0.45, zorder=2,
        )
        ax.plot(x_global, y_boundary, color="steelblue", linewidth=2.0, zorder=4)

        opposite = H - boundary
        y_opposite = y_far - direction * opposite
        ax.plot(
            x_global, y_opposite,
            color="gray", linewidth=1.0, linestyle="--", alpha=0.75, zorder=3,
        )

        ax.text(
            glass_left + W / 2.0, text_y, label,
            ha="center", va="center", fontsize=10,
            bbox={"boxstyle": "round,pad=0.2", "facecolor": "white", "alpha": 0.82},
        )

        dim_y = y0 + direction * H * 1.08
        ax.annotate(
            "", xy=(glass_left, dim_y), xytext=(glass_left + W, dim_y),
            arrowprops={"arrowstyle": "<->", "linewidth": 1.0},
        )
        ax.text(
            glass_left + W / 2.0,
            dim_y + direction * 0.018 * H,
            f"W = {W:.3f} cm",
            ha="center",
            va="bottom" if direction > 0 else "top",
            fontsize=9,
        )

        dim_x = glass_left - max(0.035 * W, 3.0)
        ax.annotate(
            "", xy=(dim_x, y0), xytext=(dim_x, y_far),
            arrowprops={"arrowstyle": "<->", "linewidth": 1.0},
        )
        ax.text(
            dim_x - max(0.012 * W, 1.5),
            y0 + direction * H / 2.0,
            f"H = {H:.3f} cm",
            rotation=90,
            ha="right", va="center", fontsize=9,
        )

    draw_glass(upper_w, upper_h, 0.0, 1.0, "Glass" if not has_lower else "Upper Glass")
    if has_lower:
        draw_glass(lower_w, lower_h, 0.0, -1.0, "Lower Glass")

    ax.plot([0.0, L], [0.0, 0.0], color="blue", linewidth=3.0, zorder=7)
    support_h = max(0.05 * max_h, 3.0)
    support_w = max(0.035 * L, 3.0)
    for x_support in (0.0, L):
        tri = patches.Polygon(
            [
                (x_support, -0.02 * support_h),
                (x_support - support_w / 2.0, -support_h),
                (x_support + support_w / 2.0, -support_h),
            ],
            closed=True,
            facecolor="mediumorchid",
            edgecolor="purple",
            linewidth=1.0,
            zorder=9,
        )
        ax.add_patch(tri)

    ax.set_title("Tributary Area")
    ax.set_xlabel("Horizontal coordinate (cm)")
    ax.set_ylabel("Glass height (cm)")
    margin_x = max(0.07 * max_w, 5.0)
    ax.set_xlim(0.5 * (L - max_w) - margin_x, 0.5 * (L + max_w) + margin_x)
    if has_lower:
        ax.set_ylim(-lower_h * 1.18, upper_h * 1.18)
    else:
        ax.set_ylim(-upper_h * 0.13, upper_h * 1.18)
    ax.set_aspect("equal", adjustable="box")
    ax.grid(True, linestyle="--", alpha=0.20)
    fig.tight_layout()
    return fig


def make_tributary_load_figure(
    transom_length_cm: float,
    wind_pressure_kpa: float,
    upper_width_cm: float,
    upper_height_cm: float,
    node_x: np.ndarray,
    wind_glass_mode: str,
    lower_width_cm: Optional[float] = None,
    lower_height_cm: Optional[float] = None,
) -> plt.Figure:
    """繪製合成後的三角形／梯形風壓線載重。"""
    L = float(transom_length_cm)
    x = np.linspace(0.0, L, 1201)

    q_upper = exact_wind_line_load_kgf_per_cm(
        x, wind_pressure_kpa, L, upper_width_cm, upper_height_cm, 1.0
    )
    if wind_glass_mode == "same_double":
        q_lower = q_upper.copy()
    elif wind_glass_mode == "separate":
        if lower_width_cm is None or lower_height_cm is None:
            raise ValueError("上下玻璃獨立輸入模式需要下方玻璃尺寸。")
        q_lower = exact_wind_line_load_kgf_per_cm(
            x, wind_pressure_kpa, L, lower_width_cm, lower_height_cm, 1.0
        )
    else:
        q_lower = np.zeros_like(q_upper)

    q_total = q_upper + q_lower
    q_max = float(np.max(q_total)) if len(q_total) else 0.0

    fig, ax = plt.subplots(figsize=(8.6, 4.5))
    ax.fill_between(x, 0.0, q_total, color="lightskyblue", alpha=0.42, zorder=1)
    ax.plot(x, q_total, color="steelblue", linewidth=2.0, zorder=4, label="Combined Wind Load")

    if wind_glass_mode == "separate":
        ax.plot(x, q_upper, linewidth=1.1, linestyle="--", alpha=0.8, label="Upper Glass Load")
        ax.plot(x, q_lower, linewidth=1.1, linestyle=":", alpha=0.8, label="Lower Glass Load")
        ax.legend(loc="upper right", fontsize=8)

    if q_max > 1.0e-12:
        arrow_x = np.linspace(0.0, L, 17)
        arrow_q = np.interp(arrow_x, x, q_total)
        for x_value, q_value in zip(arrow_x, arrow_q):
            if q_value <= 0.02 * q_max:
                continue
            ax.annotate(
                "", xy=(float(x_value), 0.0), xytext=(float(x_value), float(q_value)),
                arrowprops={"arrowstyle": "->", "linewidth": 1.0, "color": "steelblue"},
                zorder=5,
            )

    ax.plot([0.0, L], [0.0, 0.0], color="blue", linewidth=3.0, zorder=6)
    ax.scatter(node_x, np.zeros_like(node_x), s=13, color="blue", zorder=7)

    support_h = max(0.13 * q_max, 0.15) if q_max > 0.0 else 0.15
    support_w = max(0.035 * L, 3.0)
    for x_support, label in [(0.0, "Pin"), (L, "Roller")]:
        tri = patches.Polygon(
            [
                (x_support, -0.02 * support_h),
                (x_support - support_w / 2.0, -support_h),
                (x_support + support_w / 2.0, -support_h),
            ],
            closed=True,
            facecolor="mediumorchid",
            edgecolor="purple",
            linewidth=1.0,
            zorder=8,
        )
        ax.add_patch(tri)
        ax.text(x_support, -1.12 * support_h, label, ha="center", va="top", fontsize=9)

    def figure_shape_text(width_cm: float, height_cm: float) -> str:
        W = float(width_cm)
        H = float(height_cm)
        if np.isclose(W, H, rtol=0.0, atol=1.0e-9):
            return "Triangular"
        return "Trapezoidal" if W > H else "Triangular"

    if wind_glass_mode == "single":
        load_text = f"{figure_shape_text(upper_width_cm, upper_height_cm)} Wind Load"
    elif wind_glass_mode == "same_double":
        load_text = f"{figure_shape_text(upper_width_cm, upper_height_cm)} Wind Load × 2"
    else:
        upper_shape = figure_shape_text(upper_width_cm, upper_height_cm)
        lower_shape = figure_shape_text(float(lower_width_cm), float(lower_height_cm))
        load_text = f"Upper: {upper_shape} + Lower: {lower_shape}"

    ax.text(
        0.02, 0.94,
        f"{load_text}   qmax = {q_max:.3f} kgf/cm",
        transform=ax.transAxes,
        ha="left", va="top", fontsize=10,
        bbox={"boxstyle": "round,pad=0.25", "facecolor": "white", "alpha": 0.92},
    )

    ax.set_title("Wind Load Distribution")
    ax.set_xlabel("Transom X Coordinate (cm)")
    ax.set_ylabel("Line Load q(x) (kgf/cm)")
    ax.set_xlim(-0.035 * L, 1.035 * L)
    y_top = max(q_max * 1.15, 1.0)
    ax.set_ylim(-1.45 * support_h, y_top)
    ax.grid(True, linestyle="--", alpha=0.25)
    fig.tight_layout()
    return fig


# ============================================================
# 3. 主分析函式
# ============================================================


def run_analysis(
    transom_length_cm: float,
    wind_glass_mode: str,
    upper_glass_width_cm: float,
    upper_glass_height_cm: float,
    lower_glass_width_cm: Optional[float],
    lower_glass_height_cm: Optional[float],
    total_glass_thickness_mm: float,
    glass_density_kg_m3: float,
    wind_pressure_kpa: float,
    setting_block_distance_cm: float,
    material_name: str,
    grade_name: str,
    profiles: Sequence[Dict[str, object]],
) -> Dict[str, object]:
    if ops is None or opsv is None:
        raise RuntimeError("必要的分析套件尚未安裝。")
    if not profiles:
        raise ValueError("至少需要一組斷面資料。")
    if wind_glass_mode not in {"single", "same_double", "separate"}:
        raise ValueError("風壓玻璃配置無效。")

    positive_inputs = {
        "橫料支承跨距": transom_length_cm,
        "上方／單片玻璃水平寬度": upper_glass_width_cm,
        "上方／單片玻璃垂直高度": upper_glass_height_cm,
        "玻璃總實體厚度": total_glass_thickness_mm,
        "玻璃密度": glass_density_kg_m3,
        "風壓": wind_pressure_kpa,
    }
    if wind_glass_mode == "separate":
        positive_inputs["下方玻璃水平寬度"] = float(lower_glass_width_cm or 0.0)
        positive_inputs["下方玻璃垂直高度"] = float(lower_glass_height_cm or 0.0)

    for name, value in positive_inputs.items():
        if not np.isfinite(float(value)) or float(value) <= 0.0:
            raise ValueError(f"{name} 必須大於 0。")

    profiles = [dict(item) for item in profiles]
    for profile in profiles:
        for field in ("A", "Istrong", "Sstrong"):
            value = float(profile[field])
            if not np.isfinite(value) or value <= 0.0:
                raise ValueError(f"{profile['name']} 的 {field} 必須大於 0。")

    # 玻璃自重只由第一組斷面承受：單一斷面時就是該橫料；
    # 上、下橫料分開時，第一組固定為上橫料。
    for field in ("Iweak", "Sweak"):
        value = float(profiles[0][field])
        if not np.isfinite(value) or value <= 0.0:
            raise ValueError(f"{profiles[0]['name']} 的 {field} 必須大於 0。")

    L = float(transom_length_cm)
    upper_w = float(upper_glass_width_cm)
    upper_h = float(upper_glass_height_cm)
    if wind_glass_mode == "same_double":
        lower_w = upper_w
        lower_h = upper_h
    elif wind_glass_mode == "separate":
        if lower_glass_width_cm is None or lower_glass_height_cm is None:
            raise ValueError("上下玻璃獨立輸入模式需要下方玻璃尺寸。")
        lower_w = float(lower_glass_width_cm)
        lower_h = float(lower_glass_height_cm)
    else:
        lower_w = None
        lower_h = None

    if not 0.0 < setting_block_distance_cm < L / 2.0:
        raise ValueError("Setting block 距左／右支承的距離必須介於 0 與 L/2 之間。")

    if material_name not in MATERIAL_DATABASE:
        raise ValueError("材料種類不存在。")
    material = MATERIAL_DATABASE[material_name]
    grades = material["grades"]
    assert isinstance(grades, dict)
    if grade_name not in grades:
        raise ValueError("材料牌號不存在。")
    grade = grades[grade_name]
    assert isinstance(grade, dict)

    E = float(material["E"])
    fb_allow = float(grade["Fb_allow"])

    # 上、下橫料分開時：
    # - 風壓方向：上、下橫料共同作用，強軸剛度取兩者加總。
    # - 玻璃自重方向：僅上橫料承受，弱軸剛度只採上橫料。
    split_profiles = len(profiles) > 1
    area_strong = sum(float(item["A"]) for item in profiles)
    i_strong_total = sum(float(item["Istrong"]) for item in profiles)

    gravity_profile = profiles[0]
    area_weak = float(gravity_profile["A"])
    i_weak_effective = float(gravity_profile["Iweak"])

    # --------------------------------------------------------
    # 風壓 → 各片玻璃 tributary load 分別計算後疊加
    # --------------------------------------------------------
    def q_upper(x_values):
        return exact_wind_line_load_kgf_per_cm(
            x_values,
            float(wind_pressure_kpa),
            L,
            upper_w,
            upper_h,
            1.0,
        )

    def q_lower(x_values):
        if wind_glass_mode == "single":
            return np.zeros_like(np.asarray(x_values, dtype=float))
        if wind_glass_mode == "same_double":
            return q_upper(x_values)
        assert lower_w is not None and lower_h is not None
        return exact_wind_line_load_kgf_per_cm(
            x_values,
            float(wind_pressure_kpa),
            L,
            lower_w,
            lower_h,
            1.0,
        )

    def q_exact(x_values):
        return q_upper(x_values) + q_lower(x_values)

    # --------------------------------------------------------
    # 玻璃自重 → 兩個 setting block 集中載重
    # 單片／上下同尺寸：以輸入的單片玻璃計算。
    # 上下分別輸入：只以上方玻璃計算自重。
    # --------------------------------------------------------
    glass_weight = glass_self_weight_kgf(
        upper_w,
        upper_h,
        float(total_glass_thickness_mm),
        float(glass_density_kg_m3),
    )
    point_load = glass_weight / 2.0
    left_block_x = float(setting_block_distance_cm)
    right_block_x = L - float(setting_block_distance_cm)
    weak_point_loads = [(left_block_x, point_load), (right_block_x, point_load)]

    # --------------------------------------------------------
    # OpenSees 分析
    # --------------------------------------------------------
    strong_model = run_opensees_axis_model(
        axis_name="Strong Axis / Wind",
        transom_length_cm=L,
        elastic_modulus_kgf_cm2=E,
        area_cm2=area_strong,
        inertia_cm4=i_strong_total,
        q_function=q_exact,
        point_loads=[],
        exact_q_for_plot=q_exact,
    )
    strong_response = strong_model["response"]
    assert isinstance(strong_response, dict)

    weak_model = run_opensees_axis_model(
        axis_name="Weak Axis / Glass Self-weight",
        transom_length_cm=L,
        elastic_modulus_kgf_cm2=E,
        area_cm2=area_weak,
        inertia_cm4=i_weak_effective,
        q_function=None,
        point_loads=weak_point_loads,
        exact_q_for_plot=None,
        mesh_mode="point_load_events",
    )
    weak_response = weak_model["response"]
    assert isinstance(weak_response, dict)

    node_x = np.asarray(strong_model["node_x"], dtype=float)
    q_elements = strong_model["q_elements"]
    assert isinstance(q_elements, np.ndarray)

    tributary_area_figure = make_tributary_area_figure(
        L,
        upper_w,
        upper_h,
        wind_glass_mode,
        lower_w,
        lower_h,
    )
    tributary_load_figure = make_tributary_load_figure(
        L,
        float(wind_pressure_kpa),
        upper_w,
        upper_h,
        node_x,
        wind_glass_mode,
        lower_w,
        lower_h,
    )

    # --------------------------------------------------------
    # 變形檢核
    # --------------------------------------------------------
    strong_allowable, strong_formula = wind_allowable_deflection(L)
    strong_deflection = float(strong_response["max_abs_displacement"])
    strong_deflection_utilization = strong_deflection / strong_allowable
    strong_deflection_pass = strong_deflection <= strong_allowable

    weak_allowable = WEAK_DEFLECTION_LIMIT_CM
    weak_formula = "3 mm（0.300 cm）"
    weak_deflection = float(weak_response["max_abs_displacement"])
    weak_deflection_utilization = weak_deflection / weak_allowable
    weak_deflection_pass = weak_deflection <= weak_allowable

    # --------------------------------------------------------
    # 上、下橫料應力分配與雙軸檢核
    # --------------------------------------------------------
    profile_stress_results: List[Dict[str, object]] = []
    max_m_strong = float(strong_response["max_abs_moment"])
    max_m_weak = float(weak_response["max_abs_moment"])

    for profile_index, profile in enumerate(profiles):
        m_strong_share = max_m_strong * float(profile["Istrong"]) / i_strong_total

        carries_gravity = (not split_profiles) or profile_index == 0
        m_weak_share = max_m_weak if carries_gravity else 0.0

        fb_strong = m_strong_share / float(profile["Sstrong"])
        fb_weak = m_weak_share / float(profile["Sweak"]) if carries_gravity else 0.0
        combined_stress = fb_strong + fb_weak
        combined_utilization = combined_stress / fb_allow
        passed = combined_utilization <= 1.0
        profile_stress_results.append(
            {
                **profile,
                "Mstrong_share": m_strong_share,
                "Mweak_share": m_weak_share,
                "carries_gravity": carries_gravity,
                "fb_strong": fb_strong,
                "fb_weak": fb_weak,
                "combined_stress": combined_stress,
                "strong_utilization": fb_strong / fb_allow,
                "weak_utilization": fb_weak / fb_allow,
                "combined_utilization": combined_utilization,
                "passed": passed,
                "result": status_text(passed),
            }
        )

    stress_control = max(
        profile_stress_results,
        key=lambda item: float(item["combined_utilization"]),
    )
    stress_pass = all(bool(item["passed"]) for item in profile_stress_results)
    overall_pass = stress_pass and strong_deflection_pass and weak_deflection_pass

    # --------------------------------------------------------
    # 載重摘要
    # --------------------------------------------------------
    exact_x = np.linspace(0.0, L, 2001)
    exact_q_upper = q_upper(exact_x)
    exact_q_lower = q_lower(exact_x)
    exact_q = exact_q_upper + exact_q_lower

    exact_total_wind = trapezoid_integral(exact_q, exact_x)
    max_exact_q = float(np.max(exact_q))
    upper_max_trib = float(np.max(tributary_width_cm(exact_x, L, upper_w, upper_h)))
    lower_max_trib = 0.0
    if wind_glass_mode != "single":
        assert lower_w is not None and lower_h is not None
        lower_max_trib = float(np.max(tributary_width_cm(exact_x, L, lower_w, lower_h)))

    upper_shape = wind_load_shape_text(upper_w, upper_h)
    if wind_glass_mode == "single":
        load_shape = upper_shape
        wind_consideration = "單片玻璃 tributary load"
        gravity_glass_text = "單片玻璃"
    elif wind_glass_mode == "same_double":
        load_shape = f"{upper_shape} × 2"
        wind_consideration = "上下同尺寸玻璃 tributary load 疊加"
        gravity_glass_text = "一片玻璃"
    else:
        assert lower_w is not None and lower_h is not None
        lower_shape = wind_load_shape_text(lower_w, lower_h)
        load_shape = f"上方：{upper_shape}；下方：{lower_shape}（疊加）"
        wind_consideration = "上下玻璃 tributary load 分別計算後疊加"
        gravity_glass_text = "上方玻璃"

    model_summary: Dict[str, object] = {
        "橫料支承跨距 L (cm)": L,
        "風壓玻璃配置": wind_mode_label(wind_glass_mode),
        "上方／單片玻璃水平寬度 (cm)": upper_w,
        "上方／單片玻璃垂直高度 (cm)": upper_h,
    }
    if wind_glass_mode == "same_double":
        model_summary["下方玻璃尺寸"] = "同上方玻璃"
    elif wind_glass_mode == "separate":
        model_summary["下方玻璃水平寬度 (cm)"] = float(lower_w)
        model_summary["下方玻璃垂直高度 (cm)"] = float(lower_h)

    model_summary.update(
        {
            "玻璃總實體厚度 (mm)": float(total_glass_thickness_mm),
            "玻璃密度 (kg/m³)": float(glass_density_kg_m3),
            "自重計算玻璃": gravity_glass_text,
            "風壓 p (kPa)": float(wind_pressure_kpa),
            "風壓載重考慮": wind_consideration,
            "龜殼式線載重形狀": load_shape,
            "上方／單片最大 tributary width (cm)": upper_max_trib,
            "下方最大 tributary width (cm)": lower_max_trib if wind_glass_mode != "single" else "-",
            "最大線載重 qmax (kgf/cm)": max_exact_q,
            "橫料總風力 (kgf)": exact_total_wind,
            "玻璃總自重 (kgf)": glass_weight,
            "單一 setting block 集中力 (kgf)": point_load,
            "左 setting block X (cm)": left_block_x,
            "右 setting block X (cm)": right_block_x,
            "材料": material_name,
            "牌號": grade_name,
            "E (kgf/cm²)": E,
            "Fb_allow (kgf/cm²)": fb_allow,
            "斷面配置": "上、下橫料分開" if split_profiles else "單一斷面",
            "風壓方向有效 A (cm²)": area_strong,
            "風壓方向總 Istrong (cm⁴)": i_strong_total,
            "自重承載構件": str(gravity_profile["name"]),
            "自重方向有效 A (cm²)": area_weak,
            "自重方向有效 Iweak (cm⁴)": i_weak_effective,
            "強軸風壓容許變形": "L/240",
            "弱軸自重容許變形": "3 mm (0.300 cm)",
        }
    )

    strong_results = {
        "max_abs_displacement": strong_deflection,
        "max_displacement_x": float(strong_response["max_displacement_x"]),
        "max_abs_moment": max_m_strong,
        "max_moment_x": float(strong_response["max_moment_x"]),
        "max_abs_shear": float(strong_response["max_abs_shear"]),
        "max_shear_x": float(strong_response["max_shear_x"]),
        "left_reaction": float(strong_response["left_reaction"]),
        "right_reaction": float(strong_response["right_reaction"]),
        "left_reaction_fe": float(strong_model["left_reaction_fe"]),
        "right_reaction_fe": float(strong_model["right_reaction_fe"]),
        "reaction_sum": float(strong_response["reaction_sum"]),
        "force_balance_error": float(strong_response["force_balance_error"]),
        "allowable_deflection": strong_allowable,
        "allowable_formula": strong_formula,
        "deflection_utilization": strong_deflection_utilization,
        "deflection_pass": strong_deflection_pass,
    }

    weak_results = {
        "max_abs_displacement": weak_deflection,
        "max_displacement_x": float(weak_response["max_displacement_x"]),
        "max_abs_moment": max_m_weak,
        "max_moment_x": float(weak_response["max_moment_x"]),
        "max_abs_shear": float(weak_response["max_abs_shear"]),
        "max_shear_x": float(weak_response["max_shear_x"]),
        "left_reaction": float(weak_response["left_reaction"]),
        "right_reaction": float(weak_response["right_reaction"]),
        "left_reaction_fe": float(weak_model["left_reaction_fe"]),
        "right_reaction_fe": float(weak_model["right_reaction_fe"]),
        "reaction_sum": float(weak_response["reaction_sum"]),
        "force_balance_error": float(weak_response["force_balance_error"]),
        "allowable_deflection": weak_allowable,
        "allowable_formula": weak_formula,
        "deflection_utilization": weak_deflection_utilization,
        "deflection_pass": weak_deflection_pass,
    }

    stress_results = {
        "Fb_allow": fb_allow,
        "control_profile": str(stress_control["name"]),
        "fb_strong": float(stress_control["fb_strong"]),
        "fb_weak": float(stress_control["fb_weak"]),
        "combined_stress": float(stress_control["combined_stress"]),
        "strong_utilization": float(stress_control["strong_utilization"]),
        "weak_utilization": float(stress_control["weak_utilization"]),
        "combined_utilization": float(stress_control["combined_utilization"]),
        "stress_pass": stress_pass,
        "profiles": profile_stress_results,
    }

    figures = {
        "tributary_area": tributary_area_figure,
        "tributary_load": tributary_load_figure,
        "strong_model": strong_model["figures"]["model"],
        "strong_deformation": strong_model["figures"]["deformation"],
        "strong_moment": strong_model["figures"]["moment"],
        "strong_shear": strong_model["figures"]["shear"],
        "weak_model": weak_model["figures"]["model"],
        "weak_deformation": weak_model["figures"]["deformation"],
        "weak_moment": weak_model["figures"]["moment"],
        "weak_shear": weak_model["figures"]["shear"],
    }

    return {
        "overall_pass": overall_pass,
        "stress_pass": stress_pass,
        "strong_deflection_pass": strong_deflection_pass,
        "weak_deflection_pass": weak_deflection_pass,
        "model_summary": model_summary,
        "profiles": profiles,
        "strong_results": strong_results,
        "weak_results": weak_results,
        "stress_results": stress_results,
        "strong_response": strong_response,
        "weak_response": weak_response,
        "q_elements": q_elements,
        "node_x": node_x,
        "wind_glass_mode": wind_glass_mode,
        "wind_mode_label": wind_mode_label(wind_glass_mode),
        "wind_mode_description": WIND_GLASS_MODE_DESCRIPTIONS[wind_glass_mode],
        "figures": figures,
    }


# ============================================================
# 4. Word 報告
# ============================================================


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _format_docx_paragraph(paragraph, font_size: float, bold: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(font_size)
        run.bold = bold


def _format_docx_number(value: object) -> str:
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        number = float(value)
        if not np.isfinite(number):
            return "-"
        if abs(number) < 0.0005:
            number = 0.0
        return f"{number:.3f}"
    return str(value)


def _add_docx_table(document, dataframe: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, column in enumerate(dataframe.columns):
        cell = table.rows[0].cells[index]
        cell.text = str(column)
        _set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            _format_docx_paragraph(paragraph, 10.5, bold=True)

    for row in dataframe.itertuples(index=False, name=None):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = _format_docx_number(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                _format_docx_paragraph(paragraph, 10.0)

    document.add_paragraph("")


def build_word_report(result: Dict[str, object]) -> bytes:
    if Document is None:
        raise RuntimeError("尚未安裝 python-docx，無法建立 Word 報告。")

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    title = document.add_heading("帷幕橫料分析報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated = document.add_paragraph(
        f"報告產生時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER

    verdict = document.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = verdict.add_run(
        f"整體判定：{status_text(bool(result['overall_pass']))}　｜　"
        f"雙軸應力：{status_text(bool(result['stress_pass']))}　｜　"
        f"強軸變形：{status_text(bool(result['strong_deflection_pass']))}　｜　"
        f"弱軸變形：{status_text(bool(result['weak_deflection_pass']))}"
    )
    run.bold = True

    summary = result["model_summary"]
    strong = result["strong_results"]
    weak = result["weak_results"]
    stress = result["stress_results"]
    assert isinstance(summary, dict)
    assert isinstance(strong, dict)
    assert isinstance(weak, dict)
    assert isinstance(stress, dict)

    document.add_heading("1. 輸入與模型摘要", level=1)
    summary_df = pd.DataFrame([{"項目": key, "數值": value} for key, value in summary.items()])
    _add_docx_table(document, summary_df)

    profile_rows = result.get("profiles", [])
    if isinstance(profile_rows, list) and profile_rows:
        document.add_paragraph("斷面資料")
        profile_df = pd.DataFrame(profile_rows)[["name", "A", "Istrong", "Sstrong", "Iweak", "Sweak"]].rename(
            columns={
                "name": "斷面",
                "A": "A (cm²)",
                "Istrong": "Istrong (cm⁴)",
                "Sstrong": "Sstrong (cm³)",
                "Iweak": "Iweak (cm⁴)",
                "Sweak": "Sweak (cm³)",
            }
        )
        _add_docx_table(document, profile_df)

    document.add_heading("2. 強軸－風壓分析", level=1)
    strong_df = pd.DataFrame(
        [
            {"項目": "最大變形 |δ|", "數值": strong["max_abs_displacement"], "單位": "cm", "位置 X (cm)": strong["max_displacement_x"]},
            {"項目": "最大彎矩 |M|", "數值": strong["max_abs_moment"], "單位": "kgf-cm", "位置 X (cm)": strong["max_moment_x"]},
            {"項目": "最大剪力 |V|", "數值": strong["max_abs_shear"], "單位": "kgf", "位置 X (cm)": strong["max_shear_x"]},
            {"項目": "左支承反力", "數值": strong["left_reaction"], "單位": "kgf", "位置 X (cm)": 0.0},
            {"項目": "右支承反力", "數值": strong["right_reaction"], "單位": "kgf", "位置 X (cm)": summary["橫料支承跨距 L (cm)"]},
        ]
    )
    _add_docx_table(document, strong_df)

    document.add_heading("3. 弱軸－玻璃自重分析", level=1)
    weak_df = pd.DataFrame(
        [
            {"項目": "最大變形 |δ|", "數值": weak["max_abs_displacement"], "單位": "cm", "位置 X (cm)": weak["max_displacement_x"]},
            {"項目": "最大彎矩 |M|", "數值": weak["max_abs_moment"], "單位": "kgf-cm", "位置 X (cm)": weak["max_moment_x"]},
            {"項目": "最大剪力 |V|", "數值": weak["max_abs_shear"], "單位": "kgf", "位置 X (cm)": weak["max_shear_x"]},
            {"項目": "左支承反力", "數值": weak["left_reaction"], "單位": "kgf", "位置 X (cm)": 0.0},
            {"項目": "右支承反力", "數值": weak["right_reaction"], "單位": "kgf", "位置 X (cm)": summary["橫料支承跨距 L (cm)"]},
        ]
    )
    _add_docx_table(document, weak_df)

    document.add_heading("4. 變形與雙軸應力檢核", level=1)
    check_df = pd.DataFrame(
        [
            {
                "檢核項目": "強軸風壓變形",
                "需求值": strong["max_abs_displacement"],
                "容許值": strong["allowable_deflection"],
                "公式": strong["allowable_formula"],
                "使用率": strong["deflection_utilization"],
                "結果": status_text(bool(strong["deflection_pass"])),
            },
            {
                "檢核項目": "弱軸自重變形",
                "需求值": weak["max_abs_displacement"],
                "容許值": weak["allowable_deflection"],
                "公式": weak["allowable_formula"],
                "使用率": weak["deflection_utilization"],
                "結果": status_text(bool(weak["deflection_pass"])),
            },
            {
                "檢核項目": f"雙軸彎曲應力（控制：{stress['control_profile']}）",
                "需求值": stress["combined_stress"],
                "容許值": stress["Fb_allow"],
                "公式": "|fb,strong| + |fb,weak|",
                "使用率": stress["combined_utilization"],
                "結果": status_text(bool(stress["stress_pass"])),
            },
        ]
    )
    _add_docx_table(document, check_df)

    profile_stress_rows = stress.get("profiles", [])
    if isinstance(profile_stress_rows, list) and profile_stress_rows:
        document.add_paragraph("各斷面應力檢核")
        profile_stress_df = pd.DataFrame(profile_stress_rows)[
            ["name", "carries_gravity", "fb_strong", "fb_weak", "combined_stress", "combined_utilization", "result"]
        ].rename(
            columns={
                "name": "斷面",
                "carries_gravity": "承受玻璃自重",
                "fb_strong": "強軸應力",
                "fb_weak": "弱軸應力",
                "combined_stress": "合成應力",
                "combined_utilization": "使用率",
                "result": "結果",
            }
        )
        _add_docx_table(document, profile_stress_df)

    document.add_page_break()
    document.add_heading("5. 圖表", level=1)

    figures = result["figures"]
    assert isinstance(figures, dict)
    figure_items = [
        ("龜殼式 Tributary Area 示意圖", "tributary_area"),
        ("強軸風壓載重模型", "tributary_load"),
        ("強軸變形圖", "strong_deformation"),
        ("強軸彎矩圖", "strong_moment"),
        ("強軸剪力圖", "strong_shear"),
        ("弱軸模型與玻璃自重集中載重", "weak_model"),
        ("弱軸變形圖", "weak_deformation"),
        ("弱軸彎矩圖", "weak_moment"),
        ("弱軸剪力圖", "weak_shear"),
    ]

    for index, (figure_title, figure_key) in enumerate(figure_items):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(figure_title)
        run.bold = True
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(11)

        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(
            BytesIO(figure_to_png_bytes(figures[figure_key])),
            height=Cm(WORD_FIGURE_HEIGHT_CM),
        )

        if index in {1, 3, 5, 7}:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


# ============================================================
# 5. Streamlit 結果顯示
# ============================================================


def render_dataframe_download(
    dataframe: pd.DataFrame,
    button_label: str,
    filename: str,
    key: str,
) -> None:
    st.download_button(
        button_label,
        data=dataframe_to_csv_bytes(dataframe),
        file_name=filename,
        mime="text/csv",
        key=key,
        use_container_width=True,
    )


def render_results(result: Dict[str, object], report_name_input: str) -> None:
    summary = result["model_summary"]
    strong = result["strong_results"]
    weak = result["weak_results"]
    stress = result["stress_results"]
    figures = result["figures"]
    assert isinstance(summary, dict)
    assert isinstance(strong, dict)
    assert isinstance(weak, dict)
    assert isinstance(stress, dict)
    assert isinstance(figures, dict)

    st.divider()
    st.header("分析結果")

    status_col1, status_col2, status_col3, status_col4 = st.columns(4)
    status_col1.metric("整體判定", "PASS" if bool(result["overall_pass"]) else "FAIL")
    status_col2.metric("雙軸應力", "PASS" if bool(result["stress_pass"]) else "FAIL")
    status_col3.metric("強軸風壓變形", "PASS" if bool(result["strong_deflection_pass"]) else "FAIL")
    status_col4.metric("弱軸自重變形", "PASS" if bool(result["weak_deflection_pass"]) else "FAIL")

    # Word 報告匯出位置與直料分析一致：放在主要結果區上方，不另開分頁。
    if DOCX_IMPORT_ERROR is not None:
        st.error("目前環境未安裝 python-docx，無法輸出 Word。")
        st.code("python -m pip install python-docx", language="bash")
    else:
        report_bytes = build_word_report(result)
        st.download_button(
            "匯出所有數據與圖表的 Word 報告",
            data=report_bytes,
            file_name=sanitize_report_filename(report_name_input),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_transom_word_report",
        )

    tab_summary, tab_wind, tab_gravity = st.tabs(
        ["結果摘要", "強軸－風壓", "弱軸－自重"]
    )

    with tab_summary:
        st.subheader("控制結果")
        result_df = pd.DataFrame(
            [
                {
                    "項目": "強軸最大變形",
                    "需求值": strong["max_abs_displacement"],
                    "容許值": strong["allowable_deflection"],
                    "使用率": strong["deflection_utilization"],
                    "結果": status_text(bool(strong["deflection_pass"])),
                },
                {
                    "項目": "弱軸最大變形",
                    "需求值": weak["max_abs_displacement"],
                    "容許值": weak["allowable_deflection"],
                    "使用率": weak["deflection_utilization"],
                    "結果": status_text(bool(weak["deflection_pass"])),
                },
                {
                    "項目": f"雙軸彎曲應力（{stress['control_profile']}）",
                    "需求值": stress["combined_stress"],
                    "容許值": stress["Fb_allow"],
                    "使用率": stress["combined_utilization"],
                    "結果": status_text(bool(stress["stress_pass"])),
                },
            ]
        )
        st.dataframe(
            round_dataframe_for_display(result_df),
            use_container_width=True,
            hide_index=True,
        )

        st.subheader("強軸與弱軸最大結果")
        axis_df = pd.DataFrame(
            [
                {
                    "方向": "強軸－風壓",
                    "最大位移 (cm)": strong["max_abs_displacement"],
                    "位移位置 X (cm)": strong["max_displacement_x"],
                    "最大彎矩 (kgf-cm)": strong["max_abs_moment"],
                    "彎矩位置 X (cm)": strong["max_moment_x"],
                    "最大剪力 (kgf)": strong["max_abs_shear"],
                    "剪力位置 X (cm)": strong["max_shear_x"],
                },
                {
                    "方向": "弱軸－自重",
                    "最大位移 (cm)": weak["max_abs_displacement"],
                    "位移位置 X (cm)": weak["max_displacement_x"],
                    "最大彎矩 (kgf-cm)": weak["max_abs_moment"],
                    "彎矩位置 X (cm)": weak["max_moment_x"],
                    "最大剪力 (kgf)": weak["max_abs_shear"],
                    "剪力位置 X (cm)": weak["max_shear_x"],
                },
            ]
        )
        st.dataframe(
            round_dataframe_for_display(axis_df),
            use_container_width=True,
            hide_index=True,
        )

        with st.expander("模型與輸入摘要"):
            summary_df = pd.DataFrame(
                [{"項目": key, "數值": value} for key, value in summary.items()]
            )
            st.dataframe(
                round_dataframe_for_display(summary_df),
                use_container_width=True,
                hide_index=True,
            )

        st.subheader("載重換算")
        load_rows = [
            {"項目": "橫料總風力", "數值": summary["橫料總風力 (kgf)"], "單位": "kgf"},
            {"項目": "上方／單片最大 tributary width", "數值": summary["上方／單片最大 tributary width (cm)"], "單位": "cm"},
        ]
        if summary.get("下方最大 tributary width (cm)") != "-":
            load_rows.append(
                {"項目": "下方最大 tributary width", "數值": summary["下方最大 tributary width (cm)"], "單位": "cm"}
            )
        load_rows.extend(
            [
                {"項目": "最大風力線載重", "數值": summary["最大線載重 qmax (kgf/cm)"], "單位": "kgf/cm"},
                {"項目": "玻璃總自重", "數值": summary["玻璃總自重 (kgf)"], "單位": "kgf"},
                {"項目": "單一 setting block 載重", "數值": summary["單一 setting block 集中力 (kgf)"], "單位": "kgf"},
            ]
        )
        load_df = pd.DataFrame(load_rows)
        st.dataframe(round_dataframe_for_display(load_df), use_container_width=True, hide_index=True)

        st.subheader("龜殼式 Tributary Area 示意圖")
        st.caption(str(result.get("wind_mode_description", "")))
        st.pyplot(figures["tributary_area"], use_container_width=False)

        st.subheader("風壓線載重")
        st.caption(f"風壓玻璃配置：{result.get('wind_mode_label', '')}。")
        st.pyplot(figures["tributary_load"], use_container_width=False)

        st.info(
            "雙軸應力採 |fb,strong| + |fb,weak| ≤ Fb_allow。"
            "上、下橫料分開時，風壓彎矩依各自 Istrong 比例分配；"
            "玻璃自重僅由上橫料承受，下橫料弱軸應力為 0。"
        )

        profile_stress_rows = stress.get("profiles", [])
        if isinstance(profile_stress_rows, list) and profile_stress_rows:
            st.subheader("斷面應力檢核")
            profile_stress_df = pd.DataFrame(profile_stress_rows)[
                ["name", "carries_gravity", "fb_strong", "fb_weak", "combined_stress", "combined_utilization", "result"]
            ].rename(
                columns={
                    "name": "斷面",
                    "carries_gravity": "承受玻璃自重",
                    "fb_strong": "強軸應力 (kgf/cm²)",
                    "fb_weak": "弱軸應力 (kgf/cm²)",
                    "combined_stress": "合成應力 (kgf/cm²)",
                    "combined_utilization": "使用率",
                    "result": "結果",
                }
            )
            profile_stress_df["承受玻璃自重"] = profile_stress_df["承受玻璃自重"].map(
                {True: "是", False: "否"}
            )
            st.dataframe(
                round_dataframe_for_display(profile_stress_df),
                use_container_width=True,
                hide_index=True,
            )

        with st.expander("下載數據"):
            render_dataframe_download(
                axis_df,
                "下載強軸與弱軸最大結果 CSV",
                "transom_max_results.csv",
                "download_axis_csv",
            )

    with tab_wind:
        metric1, metric2, metric3 = st.columns(3)
        metric1.metric("最大強軸變形", f"{float(strong['max_abs_displacement']):.3f} cm")
        metric2.metric("最大強軸彎矩", f"{float(strong['max_abs_moment']):.3f} kgf-cm")
        metric3.metric("最大強軸剪力", f"{float(strong['max_abs_shear']):.3f} kgf")

        st.caption(
            f"容許變形：{strong['allowable_formula']} = {float(strong['allowable_deflection']):.3f} cm；"
            f"使用率 = {float(strong['deflection_utilization']):.3f}。"
        )
        st.pyplot(figures["tributary_load"], use_container_width=False)
        with st.expander("顯示強軸分析模型"):
            st.pyplot(figures["strong_model"], use_container_width=False)
        st.pyplot(figures["strong_deformation"], use_container_width=False)
        st.pyplot(figures["strong_moment"], use_container_width=False)
        st.pyplot(figures["strong_shear"], use_container_width=False)

        st.subheader("強軸支承反力")
        strong_reaction_df = pd.DataFrame(
            [
                {
                    "支承": "左支承",
                    "反力 (kgf)": strong["left_reaction"],
                    "位置 X (cm)": 0.0,
                },
                {
                    "支承": "右支承",
                    "反力 (kgf)": strong["right_reaction"],
                    "位置 X (cm)": summary["橫料支承跨距 L (cm)"],
                },
            ]
        )
        st.dataframe(
            round_dataframe_for_display(strong_reaction_df),
            use_container_width=True,
            hide_index=True,
        )

    with tab_gravity:
        metric1, metric2, metric3 = st.columns(3)
        metric1.metric("最大弱軸變形", f"{float(weak['max_abs_displacement']):.3f} cm")
        metric2.metric("最大弱軸彎矩", f"{float(weak['max_abs_moment']):.3f} kgf-cm")
        metric3.metric("最大弱軸剪力", f"{float(weak['max_abs_shear']):.3f} kgf")

        if len(result.get("profiles", [])) > 1:
            st.caption("上、下橫料分開時，本頁自重分析僅採上橫料的弱軸斷面性質。")

        st.caption(
            f"容許變形：{weak['allowable_formula']}；"
            f"使用率 = {float(weak['deflection_utilization']):.3f}。"
        )
        st.pyplot(figures["weak_model"], use_container_width=False)
        st.pyplot(figures["weak_deformation"], use_container_width=False)
        st.pyplot(figures["weak_moment"], use_container_width=False)
        st.pyplot(figures["weak_shear"], use_container_width=False)

        st.subheader("弱軸支承反力")
        weak_reaction_df = pd.DataFrame(
            [
                {
                    "支承": "左支承",
                    "反力 (kgf)": weak["left_reaction"],
                    "位置 X (cm)": 0.0,
                },
                {
                    "支承": "右支承",
                    "反力 (kgf)": weak["right_reaction"],
                    "位置 X (cm)": summary["橫料支承跨距 L (cm)"],
                },
            ]
        )
        st.dataframe(
            round_dataframe_for_display(weak_reaction_df),
            use_container_width=True,
            hide_index=True,
        )



# ============================================================
# 6. Streamlit 主程式
# ============================================================


def main() -> None:
    st.set_page_config(
        page_title="帷幕橫料分析",
        page_icon="🏗️",
        layout="wide",
    )

    st.title("帷幕橫料分析")
    st.caption(
        "單位：kgf、cm。強軸承受風壓；"
        "弱軸承受玻璃自重兩個 setting block 集中載重。"
        "上、下橫料分開時，風壓由兩者共同承受，自重僅由上橫料承受。"
        "變形限制：風壓方向 L/240；垂直自重方向 3 mm。"
    )

    if OPENSEES_IMPORT_ERROR is not None:
        st.error("目前環境缺少必要分析套件，請先安裝後重新啟動。")
        st.code(
            "python -m pip install openseespy opsvis streamlit matplotlib numpy pandas python-docx\n"
            "python -m streamlit run curtain_wall_transom_opensees_loadshare_20260817.py",
            language="bash",
        )
        st.exception(OPENSEES_IMPORT_ERROR)
        st.stop()

    # --------------------------------------------------------
    # 左側：材料、斷面、報告
    # --------------------------------------------------------
    st.sidebar.header("材料、斷面與報告")

    report_name_input = st.sidebar.text_input(
        "Word 報告檔名",
        value="",
        placeholder="transom_analysis_report",
        help="未輸入時預設為 transom_analysis_report.docx。",
    )

    material_name = st.sidebar.selectbox(
        "材料種類",
        list(MATERIAL_DATABASE.keys()),
        on_change=clear_analysis_result,
    )
    selected_material = MATERIAL_DATABASE[material_name]
    grades = selected_material["grades"]
    assert isinstance(grades, dict)
    grade_name = st.sidebar.selectbox(
        "材料牌號",
        list(grades.keys()),
        on_change=clear_analysis_result,
    )
    selected_grade = grades[grade_name]
    assert isinstance(selected_grade, dict)
    st.sidebar.caption(
        f"E = {float(selected_material['E']):,.3f} kgf/cm²　｜　"
        f"Fb = {float(selected_grade['Fb_allow']):,.3f} kgf/cm²"
    )

    split_profiles = st.sidebar.checkbox(
        "上、下橫料分開",
        value=False,
        on_change=clear_analysis_result,
        help="勾選後，風壓方向分別輸入上、下橫料強軸資料；玻璃自重方向只需輸入上橫料弱軸資料。",
    )

    def section_inputs(
        prefix: str,
        display_name: str,
        defaults: Tuple[float, float, float, float, float],
        include_weak_axis: bool = True,
    ) -> Dict[str, object]:
        default_a, default_is, default_ss, default_iw, default_sw = defaults
        with st.sidebar.expander(display_name, expanded=True):
            area = st.number_input(
                "斷面積 A (cm²)",
                min_value=0.000001,
                value=default_a,
                step=1.0,
                format="%.3f",
                key=f"{prefix}_A",
                on_change=clear_analysis_result,
            )
            i_strong = st.number_input(
                "強軸慣性矩 Istrong (cm⁴)",
                min_value=0.000001,
                value=default_is,
                step=10.0,
                format="%.3f",
                key=f"{prefix}_Istrong",
                on_change=clear_analysis_result,
            )
            s_strong = st.number_input(
                "強軸斷面係數 Sstrong (cm³)",
                min_value=0.000001,
                value=default_ss,
                step=1.0,
                format="%.3f",
                key=f"{prefix}_Sstrong",
                on_change=clear_analysis_result,
            )

            if include_weak_axis:
                i_weak = st.number_input(
                    "弱軸慣性矩 Iweak (cm⁴)",
                    min_value=0.000001,
                    value=default_iw,
                    step=5.0,
                    format="%.3f",
                    key=f"{prefix}_Iweak",
                    on_change=clear_analysis_result,
                )
                s_weak = st.number_input(
                    "弱軸斷面係數 Sweak (cm³)",
                    min_value=0.000001,
                    value=default_sw,
                    step=1.0,
                    format="%.3f",
                    key=f"{prefix}_Sweak",
                    on_change=clear_analysis_result,
                )
            else:
                # 下橫料不承受玻璃自重，因此弱軸斷面性質不參與本分析。
                i_weak = None
                s_weak = None
                st.caption("玻璃自重不由下橫料承受，因此不需輸入弱軸 I、S。")

        return {
            "name": display_name,
            "A": float(area),
            "Istrong": float(i_strong),
            "Sstrong": float(s_strong),
            "Iweak": None if i_weak is None else float(i_weak),
            "Sweak": None if s_weak is None else float(s_weak),
        }

    if split_profiles:
        st.sidebar.caption(
            "風壓由上、下橫料共同承受；玻璃自重僅由上橫料承受。"
        )
        profiles = [
            section_inputs(
                "upper_transom",
                "上橫料",
                (10.0, 500.0, 50.0, 50.0, 10.0),
                include_weak_axis=True,
            ),
            section_inputs(
                "lower_transom",
                "下橫料",
                (10.0, 500.0, 50.0, 50.0, 10.0),
                include_weak_axis=False,
            ),
        ]
    else:
        profiles = [
            section_inputs("single", "橫料斷面", (20.0, 1000.0, 100.0, 100.0, 20.0))
        ]

    st.sidebar.caption(
        "變形限制：強軸風壓方向 = L/240；弱軸玻璃自重方向 = 3 mm（0.300 cm）。"
    )

    # --------------------------------------------------------
    # 主畫面：幾何與載重
    # --------------------------------------------------------
    st.header("1. 橫料與玻璃幾何")

    transom_length_cm = st.number_input(
        "橫料支承跨距 L (cm)",
        min_value=0.000001,
        value=150.0,
        step=1.0,
        format="%.3f",
        on_change=clear_analysis_result,
    )

    wind_mode_label_selected = st.radio(
        "風壓玻璃配置",
        list(WIND_GLASS_MODE_OPTIONS.keys()),
        index=1,
        horizontal=True,
        on_change=clear_analysis_result,
    )
    wind_glass_mode = WIND_GLASS_MODE_OPTIONS[wind_mode_label_selected]
    st.caption(WIND_GLASS_MODE_DESCRIPTIONS[wind_glass_mode])

    if wind_glass_mode == "separate":
        st.markdown("**上方玻璃**")
        up_col1, up_col2 = st.columns(2)
        with up_col1:
            upper_glass_width_cm = st.number_input(
                "上方玻璃水平寬度 Wu (cm)",
                min_value=0.000001,
                value=150.0,
                step=1.0,
                format="%.3f",
                on_change=clear_analysis_result,
            )
        with up_col2:
            upper_glass_height_cm = st.number_input(
                "上方玻璃垂直高度 Hu (cm)",
                min_value=0.000001,
                value=160.0,
                step=1.0,
                format="%.3f",
                on_change=clear_analysis_result,
            )

        st.markdown("**下方玻璃**")
        low_col1, low_col2 = st.columns(2)
        with low_col1:
            lower_glass_width_cm = st.number_input(
                "下方玻璃水平寬度 Wd (cm)",
                min_value=0.000001,
                value=150.0,
                step=1.0,
                format="%.3f",
                on_change=clear_analysis_result,
            )
        with low_col2:
            lower_glass_height_cm = st.number_input(
                "下方玻璃垂直高度 Hd (cm)",
                min_value=0.000001,
                value=120.0,
                step=1.0,
                format="%.3f",
                on_change=clear_analysis_result,
            )
    else:
        size_col1, size_col2 = st.columns(2)
        with size_col1:
            upper_glass_width_cm = st.number_input(
                "玻璃水平寬度 Wg (cm)",
                min_value=0.000001,
                value=140.0,
                step=1.0,
                format="%.3f",
                on_change=clear_analysis_result,
            )
        with size_col2:
            upper_glass_height_cm = st.number_input(
                "玻璃垂直高度 Hg (cm)",
                min_value=0.000001,
                value=160.0,
                step=1.0,
                format="%.3f",
                on_change=clear_analysis_result,
            )

        if wind_glass_mode == "same_double":
            lower_glass_width_cm = float(upper_glass_width_cm)
            lower_glass_height_cm = float(upper_glass_height_cm)
        else:
            lower_glass_width_cm = None
            lower_glass_height_cm = None

    st.caption(
        "玻璃水平寬度為與橫料平行方向；垂直高度為上下方向。"
        "上下玻璃獨立輸入時，風壓會分別計算後疊加，玻璃自重只以上方玻璃計算。"
    )

    thickness_col1, thickness_col2 = st.columns(2)
    with thickness_col1:
        total_glass_thickness_mm = st.number_input(
            "玻璃總實體厚度 (mm)",
            min_value=0.000001,
            value=12.0,
            step=1.0,
            format="%.3f",
            on_change=clear_analysis_result,
        )
    with thickness_col2:
        glass_density_kg_m3 = st.number_input(
            "玻璃密度 (kg/m³)",
            min_value=0.000001,
            value=DEFAULT_GLASS_DENSITY_KG_M3,
            step=10.0,
            format="%.3f",
            on_change=clear_analysis_result,
        )

    widths_to_check = [("上方／單片玻璃", float(upper_glass_width_cm))]
    if wind_glass_mode == "separate":
        widths_to_check.append(("下方玻璃", float(lower_glass_width_cm)))
    for glass_name, width_value in widths_to_check:
        if width_value > float(transom_length_cm):
            st.info(
                f"{glass_name}水平寬度大於橫料淨跨；程式會依 x/L 比例將 tributary load 映射到支承跨距。"
            )

    st.header("2. 風壓與玻璃自重")
    load_col1, load_col2 = st.columns(2)
    with load_col1:
        wind_pressure_kpa = st.number_input(
            "設計風壓 p (kPa)",
            min_value=0.000001,
            value=3.0,
            step=0.1,
            format="%.3f",
            on_change=clear_analysis_result,
        )
    with load_col2:
        setting_option = st.selectbox(
            "Setting block 位置",
            list(SETTING_BLOCK_OPTIONS.keys()),
            on_change=clear_analysis_result,
        )

    if SETTING_BLOCK_OPTIONS[setting_option] is None:
        setting_block_distance_cm = st.number_input(
            "Setting block 距支承距離 a (cm)",
            min_value=0.000001,
            max_value=max(float(transom_length_cm) / 2.0 - 1.0e-6, 0.000001),
            value=min(float(transom_length_cm) / 4.0, float(transom_length_cm) / 2.0 - 1.0e-6),
            step=1.0,
            format="%.3f",
            on_change=clear_analysis_result,
        )
    else:
        divisor = float(SETTING_BLOCK_OPTIONS[setting_option])
        setting_block_distance_cm = float(transom_length_cm) / divisor
        st.info(
            f"兩個 setting block 位置：X = {setting_block_distance_cm:.3f} cm 與 "
            f"X = {float(transom_length_cm) - setting_block_distance_cm:.3f} cm。"
        )

    preview_weight = glass_self_weight_kgf(
        float(upper_glass_width_cm),
        float(upper_glass_height_cm),
        float(total_glass_thickness_mm),
        float(glass_density_kg_m3),
    )
    preview_point = preview_weight / 2.0

    preview_x = np.linspace(0.0, float(transom_length_cm), 1201)
    preview_q_upper = exact_wind_line_load_kgf_per_cm(
        preview_x,
        float(wind_pressure_kpa),
        float(transom_length_cm),
        float(upper_glass_width_cm),
        float(upper_glass_height_cm),
        1.0,
    )
    if wind_glass_mode == "single":
        preview_q_lower = np.zeros_like(preview_q_upper)
    elif wind_glass_mode == "same_double":
        preview_q_lower = preview_q_upper.copy()
    else:
        preview_q_lower = exact_wind_line_load_kgf_per_cm(
            preview_x,
            float(wind_pressure_kpa),
            float(transom_length_cm),
            float(lower_glass_width_cm),
            float(lower_glass_height_cm),
            1.0,
        )
    preview_qmax = float(np.max(preview_q_upper + preview_q_lower))

    preview_col1, preview_col2, preview_col3 = st.columns(3)
    preview_col1.metric("玻璃總自重", f"{preview_weight:.3f} kgf")
    preview_col2.metric("每一集中載重 P", f"{preview_point:.3f} kgf")
    preview_col3.metric("最大風力線載重 qmax", f"{preview_qmax:.3f} kgf/cm")

    with st.expander("龜殼式 tributary area 的程式定義"):
        st.markdown(
            "每片玻璃都先把橫料淨跨位置 `x/L` 映射到該片玻璃的水平位置，"
            "再取 tributary width：`min(距左玻璃邊, 距右玻璃邊, H/2)`。"
            "單片玻璃只計一組；上下同尺寸玻璃將同一組載重疊加兩次；"
            "上下玻璃獨立輸入時，兩組 tributary load 各自計算後相加。"
        )

    st.divider()
    analyze = st.button("開始分析", type="primary", use_container_width=True)

    if analyze:
        try:
            with st.spinner("正在執行橫料強軸與弱軸靜力分析..."):
                result = run_analysis(
                    transom_length_cm=float(transom_length_cm),
                    wind_glass_mode=wind_glass_mode,
                    upper_glass_width_cm=float(upper_glass_width_cm),
                    upper_glass_height_cm=float(upper_glass_height_cm),
                    lower_glass_width_cm=(
                        float(lower_glass_width_cm)
                        if lower_glass_width_cm is not None
                        else None
                    ),
                    lower_glass_height_cm=(
                        float(lower_glass_height_cm)
                        if lower_glass_height_cm is not None
                        else None
                    ),
                    total_glass_thickness_mm=float(total_glass_thickness_mm),
                    glass_density_kg_m3=float(glass_density_kg_m3),
                    wind_pressure_kpa=float(wind_pressure_kpa),
                    setting_block_distance_cm=float(setting_block_distance_cm),
                    material_name=material_name,
                    grade_name=grade_name,
                    profiles=profiles,
                )
        except Exception as exc:
            try:
                if ops is not None:
                    ops.wipe()
            except Exception:
                pass
            st.error(f"分析失敗：{exc}")
            with st.expander("顯示完整錯誤資訊"):
                st.exception(exc)
        else:
            st.session_state["analysis_result"] = result

    stored_result = st.session_state.get("analysis_result")
    if isinstance(stored_result, dict):
        render_results(stored_result, report_name_input)


if __name__ == "__main__":
    main()
