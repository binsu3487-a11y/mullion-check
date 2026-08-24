from __future__ import annotations

import math
import re
from io import BytesIO
from datetime import datetime
from typing import Dict, Tuple

import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from mpl_toolkits.mplot3d.art3d import Poly3DCollection

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ============================================================
# 1. 基本設定
# ============================================================

st.set_page_config(
    page_title="帷幕牆設計風壓計算",
    page_icon="🌬️",
    layout="wide",
)

KGF_M2_TO_KPA = 0.00980665

TERRAIN = {
    "A": {"alpha": 0.32, "zg": 500.0},
    "B": {"alpha": 0.25, "zg": 400.0},
    "C": {"alpha": 0.15, "zg": 300.0},
}

GCPI = {
    "封閉式建築物": 0.375,
    "部分封閉式建築物": 1.146,
}

# ------------------------------------------------------------
# h > 18 m 外牆 GCp：工程討論後採用值
# 保守值為小面積端；非保守模式依單元面積在對數面積軸內插，
# 至 A >= 50 m² 採右端值。
# ------------------------------------------------------------
GCP_GT18_ENGINEERING = {
    "positive": {
        "conservative": 1.80,
        "large_area": 1.25,
    },
    "zone4_negative": {
        "conservative": -1.85,
        "large_area": -1.50,
    },
    "zone5_negative": {
        "conservative": -3.85,
        "large_area": -2.15,
    },
}

# ------------------------------------------------------------
# 規範 2.4 節：V10(C) 基本設計風速行政區分區
# 行政區名稱以現行常用名稱整理；風速值依規範分區。
# ------------------------------------------------------------

WIND_SPEED_BY_LOCATION: Dict[str, Dict[str, float]] = {
    "基隆市": {"全市": 42.5},
    "臺北市": {"全市": 42.5},
    "新北市": {
        **{x: 42.5 for x in [
            "貢寮區", "雙溪區", "坪林區", "瑞芳區", "平溪區", "石碇區",
            "深坑區", "汐止區", "萬里區", "金山區", "石門區", "三芝區", "淡水區"
        ]},
        **{x: 37.5 for x in [
            "烏來區", "新店區", "三峽區", "五股區", "蘆洲區", "三重區",
            "泰山區", "新莊區", "板橋區", "中和區", "永和區", "土城區",
            "樹林區", "鶯歌區", "林口區", "八里區"
        ]},
    },
    "桃園市": {"全市": 37.5},
    "新竹縣": {
        **{x: 37.5 for x in ["新豐鄉", "湖口鄉", "新埔鎮", "關西鎮", "橫山鄉", "尖石鄉"]},
        **{x: 32.5 for x in ["五峰鄉", "北埔鄉", "峨眉鄉", "竹東鎮", "寶山鄉", "芎林鄉", "竹北市"]},
    },
    "新竹市": {"全市": 32.5},
    "苗栗縣": {"全縣": 32.5},
    "臺中市": {
        "和平區": 37.5,
        **{x: 32.5 for x in [
            "東勢區", "新社區", "太平區", "石岡區", "豐原區", "潭子區", "神岡區",
            "大雅區", "大肚區", "龍井區", "沙鹿區", "梧棲區", "清水區", "后里區",
            "外埔區", "大安區", "大甲區", "中區", "東區", "南區", "西區", "北區",
            "北屯區", "西屯區", "南屯區"
        ]},
        **{x: 27.5 for x in ["烏日區", "霧峰區", "大里區"]},
    },
    "彰化縣": {
        **{x: 32.5 for x in ["伸港鄉", "線西鄉", "和美鎮"]},
        **{x: 27.5 for x in [
            "鹿港鎮", "福興鄉", "芳苑鄉", "大城鄉", "二林鎮", "埔鹽鄉",
            "竹塘鄉", "埤頭鄉", "溪湖鎮", "溪州鄉", "二水鄉", "彰化市",
            "花壇鄉", "芬園鄉", "秀水鄉", "大村鄉", "員林市", "社頭鄉",
            "埔心鄉", "永靖鄉", "田尾鄉", "北斗鎮", "田中鎮"
        ]},
    },
    "南投縣": {
        "信義鄉": 37.5,
        "仁愛鄉": 32.5,
        **{x: 27.5 for x in ["草屯鎮", "南投市", "名間鄉", "中寮鄉", "國姓鄉", "埔里鎮", "魚池鄉"]},
        **{x: 22.5 for x in ["竹山鎮", "水里鄉", "集集鎮", "鹿谷鄉"]},
    },
    "雲林縣": {
        **{x: 32.5 for x in ["口湖鄉", "水林鄉", "四湖鄉"]},
        **{x: 27.5 for x in [
            "麥寮鄉", "臺西鄉", "東勢鄉", "崙背鄉", "褒忠鄉", "元長鄉",
            "北港鎮", "土庫鎮", "二崙鄉", "西螺鎮", "虎尾鎮", "大埤鄉",
            "莿桐鄉", "斗六市", "斗南鎮", "古坑鄉", "林內鄉"
        ]},
    },
    "嘉義縣": {
        **{x: 32.5 for x in ["布袋鎮", "義竹鄉", "鹿草鄉", "太保市", "六腳鄉", "朴子市", "東石鄉"]},
        **{x: 27.5 for x in [
            "新港鄉", "水上鄉", "溪口鄉", "民雄鄉", "大林鎮", "梅山鄉",
            "竹崎鄉", "中埔鄉", "番路鄉", "大埔鄉", "阿里山鄉"
        ]},
    },
    "嘉義市": {"全市": 27.5},
    "臺南市": {
        **{x: 37.5 for x in ["七股區", "中西區", "東區", "南區", "北區", "安平區", "安南區"]},
        **{x: 32.5 for x in [
            "永康區", "歸仁區", "新化區", "左鎮區", "玉井區", "楠西區", "南化區",
            "仁德區", "關廟區", "龍崎區", "官田區", "麻豆區", "佳里區", "西港區",
            "將軍區", "學甲區", "北門區", "新營區", "後壁區", "東山區", "六甲區",
            "下營區", "柳營區", "鹽水區", "善化區", "大內區", "山上區", "新市區", "安定區"
        ]},
        "白河區": 27.5,
    },
    "高雄市": {
        **{x: 37.5 for x in [
            "林園區", "大寮區", "大樹區", "燕巢區", "大社區", "仁武區", "鳥松區",
            "鳳山區", "橋頭區", "岡山區", "梓官區", "彌陀區", "永安區", "茄萣區",
            "路竹區", "湖內區", "桃源區", "新興區", "前金區", "苓雅區", "鹽埕區",
            "鼓山區", "旗津區", "前鎮區", "三民區", "楠梓區", "小港區", "左營區"
        ]},
        **{x: 32.5 for x in [
            "阿蓮區", "田寮區", "旗山區", "美濃區", "內門區",
            "杉林區", "六龜區", "茂林區", "甲仙區", "那瑪夏區"
        ]},
    },
    "屏東縣": {
        **{x: 47.5 for x in ["恆春鎮", "滿州鄉"]},
        **{x: 42.5 for x in ["車城鄉", "牡丹鄉", "枋山鄉", "獅子鄉", "枋寮鄉", "春日鄉"]},
        **{x: 37.5 for x in [
            "佳冬鄉", "林邊鄉", "東港鎮", "新埤鄉", "來義鄉", "泰武鄉",
            "萬巒鄉", "潮州鎮", "竹田鄉", "崁頂鄉", "南州鄉", "萬丹鄉",
            "新園鄉", "麟洛鄉", "瑪家鄉", "內埔鄉", "長治鄉", "屏東市",
            "九如鄉", "鹽埔鄉", "里港鄉", "高樹鄉", "三地門鄉", "霧臺鄉"
        ]},
        "琉球鄉": 40.0,
    },
    "宜蘭縣": {
        **{x: 42.5 for x in ["南澳鄉", "蘇澳鎮", "冬山鄉", "五結鄉", "壯圍鄉", "頭城鎮"]},
        **{x: 37.5 for x in ["大同鄉", "三星鄉", "員山鄉", "羅東鎮", "宜蘭市", "礁溪鄉"]},
    },
    "花蓮縣": {
        **{x: 47.5 for x in ["花蓮市", "吉安鄉"]},
        **{x: 42.5 for x in [
            "玉里鎮", "瑞穗鄉", "豐濱鄉", "光復鄉", "鳳林鎮",
            "壽豐鄉", "新城鄉", "秀林鄉"
        ]},
        **{x: 37.5 for x in ["富里鄉", "卓溪鄉", "萬榮鄉"]},
    },
    "臺東縣": {
        **{x: 42.5 for x in ["達仁鄉", "大武鄉", "太麻里鄉", "長濱鄉"]},
        **{x: 37.5 for x in [
            "金峰鄉", "卑南鄉", "臺東市", "東河鄉", "鹿野鄉",
            "延平鄉", "關山鎮", "池上鄉", "海端鄉", "成功鎮"
        ]},
        "綠島鄉": 65.0,
        "蘭嶼鄉": 65.0,
    },
    "金門縣": {"全縣": 35.0},
    "連江縣（馬祖）": {"全縣": 42.0},
    "澎湖縣": {"一般地區": 33.0, "東吉島": 45.0},
    "其他特殊島嶼": {"彭佳嶼": 57.0},
}


# ============================================================
# 2. Kzt 地形係數資料（表 2.3）
# ============================================================

K1_RATIOS = np.array([0.20, 0.25, 0.30, 0.35, 0.40, 0.45, 0.50])

K1_TABLE = {
    ("A/B", "山脊"): np.array([0.26, 0.33, 0.39, 0.46, 0.52, 0.59, 0.65]),
    ("A/B", "懸崖"): np.array([0.15, 0.19, 0.23, 0.26, 0.33, 0.34, 0.38]),
    ("A/B", "山丘"): np.array([0.19, 0.24, 0.29, 0.33, 0.38, 0.43, 0.48]),
    ("C", "山脊"): np.array([0.29, 0.36, 0.43, 0.51, 0.58, 0.65, 0.72]),
    ("C", "懸崖"): np.array([0.17, 0.21, 0.26, 0.30, 0.34, 0.38, 0.43]),
    ("C", "山丘"): np.array([0.21, 0.26, 0.32, 0.37, 0.42, 0.47, 0.53]),
}

K2_X_RATIO = np.array([
    -4.0, -3.5, -3.0, -2.5, -2.0, -1.5, -1.0, -0.5,
     0.0,  0.5,  1.0,  1.5,  2.0,  2.5,  3.0,  3.5, 4.0
])
K2_CLIFF = np.array([
    0.00, 0.00, 0.00, 0.00, 0.00, 0.00, 0.33, 0.67,
    1.00, 0.88, 0.75, 0.63, 0.50, 0.38, 0.25, 0.13, 0.00
])
K2_RIDGE_HILL = np.array([
    0.00, 0.00, 0.00, 0.00, 0.00, 0.00, 0.33, 0.67,
    1.00, 0.67, 0.33, 0.00, 0.00, 0.00, 0.00, 0.00, 0.00
])

K3_Z_RATIO = np.array([
    0.0, 0.1, 0.2, 0.3, 0.4, 0.5, 0.6,
    0.7, 0.8, 0.9, 1.0, 1.5, 2.0
])
K3_TABLE = {
    "山脊": np.array([1.00, 0.74, 0.55, 0.41, 0.30, 0.22, 0.17, 0.12, 0.09, 0.07, 0.05, 0.01, 0.00]),
    "懸崖": np.array([1.00, 0.78, 0.61, 0.47, 0.37, 0.29, 0.22, 0.17, 0.14, 0.11, 0.08, 0.02, 0.00]),
    "山丘": np.array([1.00, 0.67, 0.45, 0.30, 0.20, 0.14, 0.09, 0.06, 0.04, 0.03, 0.02, 0.00, 0.00]),
}


# ============================================================
# 3. 計算函式
# ============================================================

def interp_clamped(x: float, xp: np.ndarray, fp: np.ndarray) -> float:
    return float(np.interp(float(x), xp, fp, left=float(fp[0]), right=float(fp[-1])))


def k_z(z_m: float, terrain: str) -> float:
    data = TERRAIN[terrain]
    z_eff = min(max(float(z_m), 5.0), data["zg"])
    return 2.774 * (z_eff / data["zg"]) ** (2.0 * data["alpha"])


def calculate_kzt(
    z_m: float,
    terrain: str,
    topography_type: str,
    H_m: float,
    Lh_m: float,
    x_m: float,
) -> Tuple[float, float, float, float]:
    """
    Kzt = (1 + K1*K2*K3)^2
    x < 0：上風側；x > 0：下風側；x = 0：頂點。
    """
    H = float(H_m)
    Lh_original = float(Lh_m)
    if H <= 0.0 or Lh_original <= 0.0:
        return 1.0, 0.0, 0.0, 0.0

    ratio_original = H / Lh_original

    # 表 2.3 註：H/Lh > 0.5 時，
    # K1 採 H/Lh = 0.5；K2、K3 採 Lh = 2H。
    k1_ratio = min(max(ratio_original, 0.20), 0.50)
    Lh_for_decay = 2.0 * H if ratio_original > 0.50 else Lh_original

    terrain_group = "C" if terrain == "C" else "A/B"
    K1 = interp_clamped(
        k1_ratio,
        K1_RATIOS,
        K1_TABLE[(terrain_group, topography_type)],
    )

    x_ratio = float(x_m) / Lh_for_decay
    k2_values = K2_CLIFF if topography_type == "懸崖" else K2_RIDGE_HILL
    K2 = interp_clamped(x_ratio, K2_X_RATIO, k2_values)

    z_ratio = max(float(z_m), 0.0) / Lh_for_decay
    K3 = interp_clamped(
        z_ratio,
        K3_Z_RATIO,
        K3_TABLE[topography_type],
    )

    Kzt = (1.0 + K1 * K2 * K3) ** 2
    return Kzt, K1, K2, K3


def log_interp_clamped(
    area: float,
    a1: float,
    y1: float,
    a2: float,
    y2: float,
) -> float:
    """GCp chart interpolation using log10(area)."""
    A = max(float(area), 1.0e-12)

    if A <= a1:
        return float(y1)

    if A >= a2:
        return float(y2)

    t = (
        math.log10(A) - math.log10(a1)
    ) / (
        math.log10(a2) - math.log10(a1)
    )

    return float(y1 + t * (y2 - y1))


def wall_gcp_by_unit_area(
    area: float,
    h: float,
    zone: int,
) -> Tuple[float, float]:
    """
    以單元面積 A = W*H 作為本工具的 GCp 查圖面積。

    h <= 18 m：沿用圖 3.1(a) 既有曲線設定。

    h > 18 m：
        Zone 4: +1.80 -> +1.25；-1.85 -> -1.50
        Zone 5: +1.80 -> +1.25；-3.85 -> -2.15
        A <= 2 m² 採保守端值；A >= 50 m² 採大面積端值；
        中間依 log10(A) 線性內插。
    """
    if h <= 18.0:
        positive = log_interp_clamped(area, 1.0, 2.1, 50.0, 1.5)
        if zone == 4:
            negative = log_interp_clamped(area, 1.0, -2.3, 50.0, -1.7)
        else:
            negative = log_interp_clamped(area, 1.0, -3.0, 50.0, -1.7)
    else:
        positive = log_interp_clamped(
            area, 2.0,
            GCP_GT18_ENGINEERING["positive"]["conservative"],
            50.0, GCP_GT18_ENGINEERING["positive"]["large_area"],
        )
        if zone == 4:
            negative = log_interp_clamped(
                area, 2.0,
                GCP_GT18_ENGINEERING["zone4_negative"]["conservative"],
                50.0, GCP_GT18_ENGINEERING["zone4_negative"]["large_area"],
            )
        else:
            negative = log_interp_clamped(
                area, 2.0,
                GCP_GT18_ENGINEERING["zone5_negative"]["conservative"],
                50.0, GCP_GT18_ENGINEERING["zone5_negative"]["large_area"],
            )
    return float(positive), float(negative)


def conservative_wall_gcp(
    h: float,
    zone: int,
) -> Tuple[float, float]:
    """外牆局部構材 / 外部被覆物之保守 GCp。"""
    if h <= 18.0:
        positive = 2.1
        negative = -2.3 if zone == 4 else -3.0
    else:
        positive = GCP_GT18_ENGINEERING["positive"]["conservative"]
        negative = (
            GCP_GT18_ENGINEERING["zone4_negative"]["conservative"]
            if zone == 4
            else GCP_GT18_ENGINEERING["zone5_negative"]["conservative"]
        )
    return float(positive), float(negative)
def wall_zone_width_a(h: float, bmin: float) -> float:
    if h <= 18.0:
        return max(min(0.4 * h, 0.10 * bmin), max(0.9, 0.04 * bmin))
    return max(0.10 * bmin, 0.9)


# ============================================================
# 4. UI－工址、地況、有效受風面積
# ============================================================

st.title("帷幕牆外牆設計風壓計算")
st.caption(
    "由工址基本設計風速、地況、地形與外牆風壓係數，"
    "計算 Zone 4 與 Zone 5 的正、負設計風壓。"
)

with st.sidebar:
    st.header("報告輸出")

    report_filename_input = st.text_input(
        "Word 報告檔名",
        value="curtain_wall_wind_pressure_report",
        help="可輸入檔名；未輸入 .docx 時會自動補上。",
    )

    report_download_slot = st.empty()

    st.divider()
    st.header("工址")

    # --------------------------------------------------------
    # 工址：縣市 / 行政區分開選擇
    # 選單顯示同時包含「臺 / 台」別名，
    # 因此鍵盤輸入任一寫法都可以篩選到相同地區。
    # --------------------------------------------------------
    def location_display_name(name: str) -> str:
        canonical = str(name)

        if "臺" in canonical:
            alias = canonical.replace("臺", "台")
            return f"{canonical}（{alias}）"

        if "台" in canonical:
            alias = canonical.replace("台", "臺")
            return f"{canonical}（{alias}）"

        return canonical

    city = st.selectbox(
        "縣市／地區",
        options=list(WIND_SPEED_BY_LOCATION.keys()),
        index=0,
        format_func=location_display_name,
        help=(
            "點開選單後可直接輸入文字搜尋。"
            "例如「臺中」或「台中」都可以找到臺中市。"
        ),
        key="city_selectbox",
    )

    districts = WIND_SPEED_BY_LOCATION[city]

    district = st.selectbox(
        "行政區",
        options=list(districts.keys()),
        index=0,
        format_func=location_display_name,
        help=(
            "點開選單後可直接輸入文字搜尋；"
            "名稱中的「臺 / 台」兩種寫法都支援。"
        ),
        key="district_selectbox",
    )

    v10c = float(districts[district])

    st.success(
        f"{city} {district} ｜ "
        f"V10(C) = {v10c:.3f} m/s"
    )

    allow_custom_v = st.checkbox("手動覆寫 V10(C)", value=False)
    if allow_custom_v:
        v10c = st.number_input(
            "自訂 V10(C) (m/s)",
            min_value=1.0,
            value=v10c,
            step=0.5,
            format="%.3f",
        )

    st.divider()
    st.header("建築物與地況")

    h = st.number_input(
        "建築物平均屋頂高度 h (m)",
        min_value=1.0,
        value=30.0,
        step=1.0,
        format="%.3f",
    )

    terrain_input = st.selectbox(
        "上風側地況",
        ["A", "B", "C"],
        index=1,
    )

    # 高度不超過 18 m 之普通建築物依規範採地況 C。
    terrain = "C" if h <= 18.0 else terrain_input
    if h <= 18.0:
        st.info("h ≤ 18 m：本計算自動採地況 C。")

    importance = st.selectbox(
        "用途係數 I",
        [1.1, 1.0, 0.9],
        index=1,
        format_func=lambda x: {
            1.1: "I = 1.1（重要／特定公眾使用建築）",
            1.0: "I = 1.0（一般建築物）",
            0.9: "I = 0.9（重要性較低建築物）",
        }[x],
    )

    enclosure = st.selectbox(
        "封閉類型",
        ["封閉式建築物", "部分封閉式建築物"],
    )

    st.divider()
    st.header("地形係數 Kzt")

    kzt_mode = st.radio(
        "Kzt 採用方式",
        [
            "規範一般地形 Kzt = 1.000",
            "保守計算 Kzt = 1.100",
            "特殊山丘／山脊／懸崖自動計算",
        ],
        index=0,
    )

    topo_type = None
    H = Lh = x_topo = None

    if kzt_mode == "保守計算 Kzt = 1.100":
        st.caption(
            "採保守計算：Kzt = 1.100。"
        )

    elif kzt_mode == "規範一般地形 Kzt = 1.000":
        st.caption(
            "無特殊地形局部加速效應：Kzt = 1.000。"
        )

    else:
        topo_type = st.selectbox(
            "特殊地形種類",
            ["山脊", "懸崖", "山丘"],
        )

        H = st.number_input(
            "地形高度 H (m)",
            min_value=0.01,
            value=30.0,
            step=1.0,
            format="%.3f",
        )

        Lh = st.number_input(
            "上風側半高處水平距離 Lh (m)",
            min_value=0.01,
            value=100.0,
            step=5.0,
            format="%.3f",
        )

        x_topo = st.number_input(
            "構址相對頂點水平距離 x (m)",
            value=0.0,
            step=5.0,
            format="%.3f",
            help="頂點 x=0；上風側輸入負值；下風側輸入正值。",
        )

        ratio = H / Lh
        threshold = 4.5 if terrain == "C" else 18.0

        st.caption(
            f"H/Lh = {ratio:.3f}；"
            f"目前地況 {terrain} 的高度門檻為 "
            f"H > {threshold:.3f} m。"
        )

        if ratio < 0.20:
            st.warning(
                "H/Lh < 0.20，請確認是否符合特殊地形條件。"
            )

        if H <= threshold:
            st.warning(
                f"H ≤ {threshold:.3f} m，"
                "請確認是否符合特殊地形條件。"
            )

    st.divider()
    st.header("帷幕單元尺寸")

    unit_width = st.number_input(
        "單元寬度 W (m)",
        min_value=0.01,
        value=1.5,
        step=0.1,
        format="%.3f",
    )

    unit_height = st.number_input(
        "單元高度 H (m)",
        min_value=0.01,
        value=3.0,
        step=0.1,
        format="%.3f",
    )

    unit_area = float(unit_width) * float(unit_height)

    st.caption(
        f"單元受風面積 A = W × H = {unit_area:.3f} m²。"
    )

    st.subheader("GCp 採用方式")

    gcp_mode = st.radio(
        "GCp",
        [
            "直接採保守值",
            "依單元面積 A = W × H 查圖",
        ],
        index=0,
    )

    if gcp_mode == "直接採保守值":
        st.caption(
            "直接採規範圖中較保守的 GCp，"
            "不以受風面積降低風壓係數。"
        )
    else:
        st.caption(
            f"以單元面積 A = {unit_area:.3f} m² "
            "作為 GCp 圖表查值面積，並按對數面積軸內插。"
        )

    st.subheader("建築物平面尺寸")

    building_x = st.number_input(
        "建築物平面 X 向尺寸 Bx (m)",
        min_value=0.1,
        value=40.0,
        step=1.0,
        format="%.3f",
    )

    building_y = st.number_input(
        "建築物平面 Y 向尺寸 By (m)",
        min_value=0.1,
        value=20.0,
        step=1.0,
        format="%.3f",
    )

    bmin = min(float(building_x), float(building_y))
    bmax = max(float(building_x), float(building_y))

    st.caption(
        f"Bmin = min(Bx, By) = {bmin:.3f} m ｜ "
        f"Bmax = max(Bx, By) = {bmax:.3f} m"
    )


# ============================================================
# 5. Kzt 與控制風速壓
# ============================================================

def get_kzt_at_height(height: float) -> Tuple[float, float, float, float]:
    if kzt_mode == "保守計算 Kzt = 1.100":
        return 1.1, 0.0, 0.0, 0.0

    if kzt_mode == "規範一般地形 Kzt = 1.000":
        return 1.0, 0.0, 0.0, 0.0

    return calculate_kzt(
        height,
        terrain,
        str(topo_type),
        float(H),
        float(Lh),
        float(x_topo),
    )


def q_at(height: float) -> Tuple[float, float, float, float, float, float]:
    Kzt, K1, K2, K3 = get_kzt_at_height(height)
    K = k_z(height, terrain)
    q = 0.06 * K * Kzt * (float(importance) * float(v10c)) ** 2
    return q, K, Kzt, K1, K2, K3


# 屋頂高度之風速壓
qh, Kh, Kzt_h, K1_h, K2_h, K3_h = q_at(h)

# ------------------------------------------------------------
# 整面帷幕牆採最大控制風速壓
# ------------------------------------------------------------
if kzt_mode != "特殊山丘／山脊／懸崖自動計算":
    control_height = float(h)
    q_control = float(qh)
    K_control = float(Kh)
    Kzt_control = float(Kzt_h)
    K1_control = float(K1_h)
    K2_control = float(K2_h)
    K3_control = float(K3_h)
else:
    # 特殊地形時 K(z) 隨高度增加，但 Kzt（控制高度） 可能下降，
    # 因此在建築物高度範圍內自動搜尋最大 q。
    z_scan_min = min(5.0, float(h))
    z_scan = np.linspace(z_scan_min, float(h), 1001)

    q_values = []
    detail_values = []

    for zz in z_scan:
        q_i, K_i, Kzt_i, K1_i, K2_i, K3_i = q_at(float(zz))
        q_values.append(q_i)
        detail_values.append((K_i, Kzt_i, K1_i, K2_i, K3_i))

    idx_control = int(np.argmax(q_values))
    control_height = float(z_scan[idx_control])
    q_control = float(q_values[idx_control])

    (
        K_control,
        Kzt_control,
        K1_control,
        K2_control,
        K3_control,
    ) = detail_values[idx_control]

# ============================================================
# 6. 內外風速壓設定
# ============================================================

# 本工具採整面帷幕牆之最大控制風速壓。
q_external = q_control
face_label = "整面帷幕牆採最大控制風速壓"

gcpi_abs = GCPI[enclosure]

# 為維持整面保守設計，內風壓所使用之風速壓亦採控制風速壓。
qi_negative_internal = q_control
qi_positive_internal = q_control

# ============================================================
# 7. Zone 4 / 5 同時計算
# ============================================================

zone_results = {}

for zone in (4, 5):
    if gcp_mode == "直接採保守值":
        gcp_pos, gcp_neg = conservative_wall_gcp(h, zone)
    else:
        gcp_pos, gcp_neg = wall_gcp_by_unit_area(
            unit_area, h, zone
        )

    p_positive_kgf_m2 = (
        q_external * gcp_pos
        - qi_negative_internal * (-gcpi_abs)
    )

    p_negative_kgf_m2 = (
        q_external * gcp_neg
        - qi_positive_internal * (+gcpi_abs)
    )

    p_positive_kpa = p_positive_kgf_m2 * KGF_M2_TO_KPA
    p_negative_kpa = p_negative_kgf_m2 * KGF_M2_TO_KPA

    zone_results[zone] = {
        "GCp+": gcp_pos,
        "GCp-": gcp_neg,
        "p+": p_positive_kgf_m2,
        "p-": p_negative_kgf_m2,
        "p+ kPa": p_positive_kpa,
        "p- kPa": p_negative_kpa,
    }

a_zone = wall_zone_width_a(h, bmin)


def create_gcp_chart_figure(
    area: float,
    building_height: float,
    gcp_mode: str,
):
    """
    Draw the exterior-wall GCp chart used by this tool.

    For exterior walls:
    - Zone 4 and Zone 5 share the same positive GCp curve.
    - Negative GCp is different for Zone 4 and Zone 5.
    """
    h_local = float(building_height)

    x_curve = np.logspace(
        math.log10(0.5),
        math.log10(100.0),
        400,
    )

    pos_curve = []
    neg4_curve = []
    neg5_curve = []

    for A in x_curve:
        pos4, neg4 = wall_gcp_by_unit_area(
            float(A), h_local, 4
        )
        pos5, neg5 = wall_gcp_by_unit_area(
            float(A), h_local, 5
        )

        # Positive curve is shared by Zone 4 and Zone 5.
        pos_curve.append(pos4)
        neg4_curve.append(neg4)
        neg5_curve.append(neg5)

        if abs(pos4 - pos5) > 1.0e-12:
            raise RuntimeError(
                "Zone 4 / Zone 5 positive GCp curves are inconsistent."
            )

    if gcp_mode == "依單元面積 A = W × H 查圖":
        marker_area = float(area)
        marker_mode_text = f"A = {marker_area:.3f} m²"
    else:
        # Conservative value corresponds to the left plateau
        # of the relevant GCp chart.
        marker_area = 1.0 if h_local <= 18.0 else 2.0
        marker_mode_text = "Conservative point"

    pos_marker, neg4_marker = wall_gcp_by_unit_area(
        marker_area, h_local, 4
    )
    _, neg5_marker = wall_gcp_by_unit_area(
        marker_area, h_local, 5
    )

    fig, ax = plt.subplots(figsize=(6.0, 4.6))

    ax.plot(
        x_curve,
        neg5_curve,
        linewidth=2.0,
        label="Zone 5 negative",
    )
    ax.plot(
        x_curve,
        neg4_curve,
        linewidth=2.0,
        label="Zone 4 negative",
    )
    ax.plot(
        x_curve,
        pos_curve,
        linewidth=2.0,
        label="Zone 4 & 5 positive",
    )

    ax.axvline(
        marker_area,
        linestyle="--",
        linewidth=1.2,
        alpha=0.75,
    )

    marker_data = [
        (neg5_marker, "Z5(-)"),
        (neg4_marker, "Z4(-)"),
        (pos_marker, "Z4&Z5(+)"),
    ]

    for y_value, label in marker_data:
        ax.scatter(
            [marker_area],
            [y_value],
            s=34,
            zorder=5,
        )
        ax.annotate(
            f"{label} = {y_value:.3f}",
            xy=(marker_area, y_value),
            xytext=(8, 5),
            textcoords="offset points",
            fontsize=8,
        )

    ax.set_xscale("log")
    ax.set_xlim(0.5, 100.0)

    # Match the convention of the regulation figure:
    # negative values visually above and positive values below.
    if h_local <= 18.0:
        ax.set_ylim(2.6, -3.5)
        title_prefix = "GCp chart | h ≤ 18 m"
    else:
        ax.set_ylim(2.6, -4.4)
        title_prefix = "GCp chart | h > 18 m"

    tick_values = [0.5, 1, 2, 5, 10, 20, 50, 100]
    ax.set_xticks(tick_values)
    ax.set_xticklabels(
        ["0.5", "1", "2", "5", "10", "20", "50", "100"]
    )

    ax.set_xlabel("Area A (m²)", fontsize=9)
    ax.set_ylabel("GCp", fontsize=9)

    ax.set_title(
        f"{title_prefix} | {marker_mode_text}",
        fontsize=10,
    )

    ax.grid(
        True,
        which="both",
        linestyle=":",
        linewidth=0.7,
        alpha=0.6,
    )

    ax.legend(
        loc="best",
        fontsize=8,
    )

    ax.tick_params(labelsize=8)
    fig.tight_layout()

    return fig


def create_zone_3d_figure(
    building_height: float,
    building_x: float,
    building_y: float,
    zone5_width: float,
    zone_results: dict,
    pressure_mode: str,
):
    """
    3D schematic for facade Zone 4 / Zone 5.

    Uses user-defined Bx x By as the schematic building footprint.
    The purpose is to visualize Zone 4 / Zone 5 extent and relative pressure magnitude.
    """
    H_plot = max(float(building_height), 1.0e-6)
    X_plot = max(float(building_x), 1.0e-6)
    Y_plot = max(float(building_y), 1.0e-6)

    # Zone 5 corner width is based on Bmin.
    Bmin_plot = min(X_plot, Y_plot)
    a_plot = min(
        max(float(zone5_width), 0.0),
        Bmin_plot / 2.0,
    )

    if pressure_mode == "負壓":
        p4 = float(zone_results[4]["p-"])
        p5 = float(zone_results[5]["p-"])
        p4_kpa = float(zone_results[4]["p- kPa"])
        p5_kpa = float(zone_results[5]["p- kPa"])
        pressure_key = "Negative pressure"
        direction_sign = -1.0
    else:
        p4 = float(zone_results[4]["p+"])
        p5 = float(zone_results[5]["p+"])
        p4_kpa = float(zone_results[4]["p+ kPa"])
        p5_kpa = float(zone_results[5]["p+ kPa"])
        pressure_key = "Positive pressure"
        direction_sign = 1.0

    max_p = max(abs(p4), abs(p5), 1.0e-12)

    fig = plt.figure(figsize=(6.2, 5.8))
    ax = fig.add_subplot(
        111,
        projection="3d",
        computed_zorder=False,
    )

    zone4_color = "#4C78A8"
    zone5_color = "#E45756"
    roof_color = "#B8B8B8"

    def add_panel(vertices, color, alpha=0.72):
        poly = Poly3DCollection(
            [vertices],
            facecolors=color,
            edgecolors="black",
            linewidths=0.6,
            alpha=alpha,
            zorder=1,
        )
        ax.add_collection3d(poly)

    # Roof
    add_panel(
        [
            (0, 0, H_plot),
            (X_plot, 0, H_plot),
            (X_plot, Y_plot, H_plot),
            (0, Y_plot, H_plot),
        ],
        roof_color,
        alpha=0.20,
    )

    # Split each of four facades into Zone 5 edge strips + Zone 4 center.
    # Front / back faces: x-direction width.
    x_segments = []
    if a_plot > 0:
        x_segments.append((0.0, a_plot, 5))
    if X_plot - 2.0 * a_plot > 1.0e-9:
        x_segments.append((a_plot, X_plot - a_plot, 4))
    if a_plot > 0:
        x_segments.append((X_plot - a_plot, X_plot, 5))

    for y in (0.0, Y_plot):
        for x1, x2, zone in x_segments:
            add_panel(
                [
                    (x1, y, 0.0),
                    (x2, y, 0.0),
                    (x2, y, H_plot),
                    (x1, y, H_plot),
                ],
                zone5_color if zone == 5 else zone4_color,
            )

    # Left / right faces: y-direction width.
    y_segments = []
    if a_plot > 0:
        y_segments.append((0.0, a_plot, 5))
    if Y_plot - 2.0 * a_plot > 1.0e-9:
        y_segments.append((a_plot, Y_plot - a_plot, 4))
    if a_plot > 0:
        y_segments.append((Y_plot - a_plot, Y_plot, 5))

    for x in (0.0, X_plot):
        for y1, y2, zone in y_segments:
            add_panel(
                [
                    (x, y1, 0.0),
                    (x, y2, 0.0),
                    (x, y2, H_plot),
                    (x, y1, H_plot),
                ],
                zone5_color if zone == 5 else zone4_color,
            )

    # --------------------------------------------------------
    # Zone widths
    # --------------------------------------------------------
    zone4_x_width = max(X_plot - 2.0 * a_plot, 0.0)
    zone4_y_width = max(Y_plot - 2.0 * a_plot, 0.0)

    ax.set_title(
        pressure_key,
        fontsize=10,
    )

    ax.set_xlabel("Width X (m)", fontsize=8, labelpad=0)
    ax.set_ylabel("Depth Y (m)", fontsize=8, labelpad=0)
    ax.set_zlabel("Height (m)", fontsize=8, labelpad=4)

    ax.set_xlim(-0.12 * X_plot, 1.20 * X_plot)
    ax.set_ylim(-0.20 * Y_plot, 1.12 * Y_plot)
    ax.set_zlim(-0.13 * H_plot, 1.08 * H_plot)

    try:
        ax.set_box_aspect(
            (
                1.0,
                max(Y_plot / X_plot, 0.25),
                max(H_plot / X_plot, 0.8),
            )
        )
    except Exception:
        pass

    ax.view_init(elev=21, azim=-50)
    ax.tick_params(axis="both", labelsize=7)
    ax.zaxis.set_tick_params(labelsize=7)

    legend_handles = [
        Patch(
            facecolor=zone4_color,
            edgecolor="black",
            label=f"Zone 4 | {p4:.3f} kgf/m²\n({p4_kpa:.3f} kPa)",
        ),
        Patch(
            facecolor=zone5_color,
            edgecolor="black",
            label=f"Zone 5 | {p5:.3f} kgf/m²\n({p5_kpa:.3f} kPa)",
        ),
    ]
    # --------------------------------------------------------
    # Engineering-style Zone dimension lines
    # --------------------------------------------------------
    # X facade dimension line outside the front face.
    x_dim_y = -0.13 * max(Y_plot, 1.0)
    x_dim_z = 0.0
    tick_h = 0.020 * max(H_plot, 1.0)

    x_segments_dim = [
        (0.0, a_plot, "Zone 5", a_plot),
        (a_plot, X_plot - a_plot, "Zone 4", zone4_x_width),
        (X_plot - a_plot, X_plot, "Zone 5", a_plot),
    ]

    for x1, x2, zone_name, seg_len in x_segments_dim:
        if x2 - x1 <= 1.0e-9:
            continue

        ax.plot(
            [x1, x2],
            [x_dim_y, x_dim_y],
            [x_dim_z, x_dim_z],
            linewidth=1.5,
            zorder=20,
        )

        for xx in (x1, x2):
            ax.plot(
                [xx, xx],
                [x_dim_y, x_dim_y],
                [x_dim_z, x_dim_z + tick_h],
                linewidth=1.0,
                zorder=20,
            )

        ax.text(
            0.5 * (x1 + x2),
            x_dim_y,
            x_dim_z + 1.35 * tick_h,
            f"{zone_name}\n{seg_len:.3f} m",
            ha="center",
            va="bottom",
            fontsize=6.5,
            zorder=50,
            clip_on=False,
            bbox=dict(
                boxstyle="round,pad=0.15",
                facecolor="white",
                edgecolor="none",
                alpha=0.88,
            ),
        )

    # Y facade dimension line outside the right face.
    y_dim_x = X_plot + 0.10 * max(X_plot, 1.0)
    y_dim_z = 0.0

    y_segments_dim = [
        (0.0, a_plot, "Zone 5", a_plot),
        (a_plot, Y_plot - a_plot, "Zone 4", zone4_y_width),
        (Y_plot - a_plot, Y_plot, "Zone 5", a_plot),
    ]

    for y1, y2, zone_name, seg_len in y_segments_dim:
        if y2 - y1 <= 1.0e-9:
            continue

        ax.plot(
            [y_dim_x, y_dim_x],
            [y1, y2],
            [y_dim_z, y_dim_z],
            linewidth=1.5,
            zorder=20,
        )

        for yy in (y1, y2):
            ax.plot(
                [y_dim_x, y_dim_x],
                [yy, yy],
                [y_dim_z, y_dim_z + tick_h],
                linewidth=1.0,
                zorder=20,
            )

        ax.text(
            y_dim_x,
            0.5 * (y1 + y2),
            y_dim_z + 1.35 * tick_h,
            f"{zone_name}\n{seg_len:.3f} m",
            ha="left",
            va="bottom",
            fontsize=6.5,
            zorder=50,
            clip_on=False,
            bbox=dict(
                boxstyle="round,pad=0.15",
                facecolor="white",
                edgecolor="none",
                alpha=0.88,
            ),
        )

    # Reserve the right side of the figure for clear annotations.
    fig.subplots_adjust(
        left=0.08,
        right=0.68,
        top=0.90,
        bottom=0.40,
    )

    fig.legend(
        handles=legend_handles,
        loc="upper left",
        bbox_to_anchor=(0.695, 0.82),
        fontsize=7,
        frameon=False,
    )

    # --------------------------------------------------------
    # Clear 2D side panel: zone ranges and design pressure.
    # This text is outside the 3D axes, so it cannot be hidden
    # behind the blue/red facade surfaces.
    # --------------------------------------------------------
    fig.text(
        0.72,
        0.66,
        "Zone widths",
        ha="left",
        va="top",
        fontsize=8,
        fontweight="bold",
    )

    fig.text(
        0.72,
        0.615,
        (
            f"X direction\n"
            f"Z5  {a_plot:.3f} m | "
            f"Z4  {zone4_x_width:.3f} m | "
            f"Z5  {a_plot:.3f} m\n\n"
            f"Y direction\n"
            f"Z5  {a_plot:.3f} m | "
            f"Z4  {zone4_y_width:.3f} m | "
            f"Z5  {a_plot:.3f} m"
        ),
        ha="left",
        va="top",
        fontsize=7,
    )

    fig.text(
        0.72,
        0.46,
        "Design pressure",
        ha="left",
        va="top",
        fontsize=8,
        fontweight="bold",
    )

    fig.text(
        0.72,
        0.415,
        (
            f"Zone 4  {p4:.3f} kgf/m²\n"
            f"        ({p4_kpa:.3f} kPa)\n\n"
            f"Zone 5  {p5:.3f} kgf/m²\n"
            f"        ({p5_kpa:.3f} kPa)"
        ),
        ha="left",
        va="top",
        fontsize=7,
    )

    return fig



def sanitize_docx_filename(name: str) -> str:
    filename = str(name or "").strip() or "curtain_wall_wind_pressure_report"
    filename = re.sub(r'[<>:"/\\|?*]+', "_", filename)
    if not filename.lower().endswith(".docx"):
        filename += ".docx"
    return filename


def fig_to_png_buffer(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    buf.seek(0)
    return buf


def _doc_add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = str(value)
            row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def build_word_report(ctx: dict) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("尚未安裝 python-docx，請在 requirements.txt 加入 python-docx。")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("帷幕牆外牆設計風壓計算報告")
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("報告產生時間：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    doc.add_paragraph(
        "說明：規範原始依據與本工具工程設定分開標示。"
        "保守計算或工程討論採值會明確註明，不視為規範表定固定值。"
    )

    h = ctx["h"]
    gcp_mode = ctx["gcp_mode"]
    kzt_mode = ctx["kzt_mode"]
    gcp_figure = "圖 3.1" if h <= 18.0 else "圖 3.2"
    pressure_eq = "式 (3.1)" if h <= 18.0 else "式 (3.2)"

    if h > 18.0:
        gcp_basis = (
            "規範第 3.3 節、圖 3.2；工程討論後保守取值"
            if gcp_mode == "直接採保守值"
            else "規範第 3.3 節、圖 3.2；A=W×H 於對數面積軸內插，端點採工程討論值"
        )
    else:
        gcp_basis = f"規範第 3.3 節、{gcp_figure}；本工具既有取值設定"

    if kzt_mode == "規範一般地形 Kzt = 1.000":
        kzt_basis = "規範第 2.6 節；一般地形本工具採 Kzt=1.000"
    elif kzt_mode == "保守計算 Kzt = 1.100":
        kzt_basis = "保守計算設定 Kzt=1.100；非規範表定一般地形固定值"
    else:
        kzt_basis = "規範表 2.3(a)、2.3(b)、2.3(c)；Kzt=(1+K1K2K3)^2"

    zone_basis = (
        "規範圖 3.1 外牆區域定義"
        if h <= 18.0
        else "規範圖 3.2 註 4：a 為最小寬度 10%，且不小於 0.9 m"
    )

    td = TERRAIN[ctx["terrain"]]
    zr = ctx["zone_results"]

    doc.add_heading("1. 輸入、規範取值與工程設定", level=1)
    input_rows = [
        ["工址", f'{ctx["city"]} {ctx["district"]}', "-", "使用者選擇"],
        ["基本設計風速 V10(C)", f'{ctx["v10c"]:.3f}', "m/s", "規範第 2.4 節：基本設計風速"],
        ["建築物高度 h", f'{h:.3f}', "m", "使用者輸入"],
        ["地況", ctx["terrain"], "-", "規範第 2.3 節：地況分類"],
        ["地況參數 α", f'{td["alpha"]:.3f}', "-", "規範表 2.2"],
        ["梯度高度 zg", f'{td["zg"]:.3f}', "m", "規範表 2.2"],
        ["用途係數 I", f'{ctx["importance"]:.3f}', "-", "規範第 2.5 節：用途係數"],
        ["封閉類型", ctx["enclosure"], "-", "使用者選擇；GCpi 依規範表 2.17"],
        ["內風壓係數 GCpi", f'±{ctx["gcpi_abs"]:.3f}', "-", "規範表 2.17"],
        ["Kzt 採用方式", kzt_mode, "-", kzt_basis],
        ["K(h)", f'{ctx["Kh"]:.3f}', "-", "規範式 (2.7)、表 2.2"],
        ["Kzt(h)", f'{ctx["Kzt_h"]:.3f}', "-", kzt_basis],
        ["q(h)", f'{ctx["qh"]:.3f}', "kgf/m²", "規範式 (2.6)"],
        ["控制高度", f'{ctx["control_height"]:.3f}', "m", "工程設定：整面帷幕牆採最大控制風速壓"],
        ["控制 K", f'{ctx["K_control"]:.3f}', "-", "規範式 (2.7)"],
        ["控制 Kzt", f'{ctx["Kzt_control"]:.3f}', "-", kzt_basis],
        ["控制風速壓 q", f'{ctx["q_control"]:.3f}', "kgf/m²", "規範式 (2.6)；整面採最大值為工程保守設定"],
        ["控制風速壓 q", f'{ctx["q_control"] * KGF_M2_TO_KPA:.3f}', "kPa", "由 kgf/m² 單位換算"],
        ["單元寬度 W", f'{ctx["unit_width"]:.3f}', "m", "使用者輸入"],
        ["單元高度 H", f'{ctx["unit_height"]:.3f}', "m", "使用者輸入"],
        ["單元面積 A=W×H", f'{ctx["unit_area"]:.3f}', "m²", "幾何計算；查圖模式作為 GCp 內插面積"],
        ["建築平面 Bx", f'{ctx["building_x"]:.3f}', "m", "使用者輸入"],
        ["建築平面 By", f'{ctx["building_y"]:.3f}', "m", "使用者輸入"],
        ["Bmin", f'{ctx["bmin"]:.3f}', "m", "Bmin=min(Bx,By)"],
        ["Bmax", f'{ctx["bmax"]:.3f}', "m", "Bmax=max(Bx,By)"],
        ["Zone 5 角隅寬度", f'{ctx["a_zone"]:.3f}', "m", zone_basis],
        ["GCp 採用方式", gcp_mode, "-", gcp_basis],
    ]
    if kzt_mode == "特殊山丘／山脊／懸崖自動計算":
        input_rows += [
            ["特殊地形種類", str(ctx["topo_type"]), "-", "規範表 2.3(a)～(c)"],
            ["地形高度 H", f'{float(ctx["H"]):.3f}', "m", "規範表 2.3(a)"],
            ["半高距離 Lh", f'{float(ctx["Lh"]):.3f}', "m", "規範表 2.3(a)～(c)"],
            ["相對頂點距離 x", f'{float(ctx["x_topo"]):.3f}', "m", "規範表 2.3(b)"],
        ]
    _doc_add_table(doc, ["項目", "數值", "單位", "規範依據／工程設定"], input_rows)

    doc.add_heading("2. Zone 4 / Zone 5 設計風壓結果", level=1)
    result_rows = []
    for zone in (4, 5):
        r = zr[zone]
        result_rows += [
            [f"Zone {zone} GCp(+)", f'{r["GCp+"]:.3f}', "-", gcp_basis],
            [f"Zone {zone} GCp(-)", f'{r["GCp-"]:.3f}', "-", gcp_basis],
            [f"Zone {zone} 正設計風壓", f'{r["p+"]:.3f}', "kgf/m²", f"規範第 3.2 節、{pressure_eq}；正外壓搭配負 GCpi"],
            [f"Zone {zone} 正設計風壓", f'{r["p+ kPa"]:.3f}', "kPa", f"規範第 3.2 節、{pressure_eq}；單位換算"],
            [f"Zone {zone} 負設計風壓", f'{r["p-"]:.3f}', "kgf/m²", f"規範第 3.2 節、{pressure_eq}；負外壓搭配正 GCpi"],
            [f"Zone {zone} 負設計風壓", f'{r["p- kPa"]:.3f}', "kPa", f"規範第 3.2 節、{pressure_eq}；單位換算"],
        ]
    _doc_add_table(doc, ["項目", "數值", "單位", "規範依據／工程設定"], result_rows)

    doc.add_heading("3. 主要計算式", level=1)
    _doc_add_table(doc, ["計算項目", "公式／採用方式", "依據"], [
        ["風速壓地況係數", "K(z)=2.774(z/zg)^(2α)，z≤5 m 時以 5 m 計", "規範式 (2.7)、表 2.2"],
        ["風速壓", "q(z)=0.06 K(z) Kzt [I V10(C)]²", "規範式 (2.6)"],
        ["設計風壓", "p=q(h)[GCp-GCpi]" if h <= 18.0 else "p=q(GCp)-qi(GCpi)", f"規範第 3.2 節、{pressure_eq}；表 2.1(a)"],
        ["內風壓係數", f'GCpi=±{ctx["gcpi_abs"]:.3f}', "規範表 2.17"],
        ["GCp", "保守取值" if gcp_mode == "直接採保守值" else "依 A=W×H 於對數面積軸內插", gcp_basis],
    ])

    doc.add_heading("4. GCp 查圖與採值", level=1)
    fig_gcp_report = create_gcp_chart_figure(ctx["unit_area"], h, gcp_mode)
    doc.add_picture(fig_to_png_buffer(fig_gcp_report), width=Inches(5.8))
    plt.close(fig_gcp_report)
    cap = doc.add_paragraph(f"GCp 曲線與本案採值位置；依據規範第 3.3 節、{gcp_figure}。")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h > 18.0:
        doc.add_paragraph(
            "h>18 m 的 GCp 工程採用值：Zone 4 為 +1.800→+1.250、-1.850→-1.500；"
            "Zone 5 為 +1.800→+1.250、-3.850→-2.150。"
            "此段為工程討論採值，係參照圖 3.2 使用。"
        )

    doc.add_heading("5. Zone 4 / Zone 5 視覺化", level=1)
    for pmode in ("負壓", "正壓"):
        fig_zone = create_zone_3d_figure(h, ctx["building_x"], ctx["building_y"], ctx["a_zone"], zr, pmode)
        doc.add_picture(fig_to_png_buffer(fig_zone), width=Inches(6.1))
        plt.close(fig_zone)
        c = doc.add_paragraph("Zone 4 / Zone 5 負壓分區示意" if pmode == "負壓" else "Zone 4 / Zone 5 正壓分區示意")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("6. 規範依據摘要", level=1)
    _doc_add_table(doc, ["規範章節／表／圖／公式", "用途"], [
        ["第 2.3 節、表 2.2", "地況分類、α、zg"],
        ["第 2.4 節", "基本設計風速 V10(C)"],
        ["第 2.5 節", "用途係數 I"],
        ["式 (2.6)", "風速壓 q(z)"],
        ["式 (2.7)", "風速壓地況係數 K(z)"],
        ["表 2.3(a)～(c)", "特殊地形 Kzt 計算所需 K1、K2、K3"],
        ["表 2.17", "內風壓係數 GCpi"],
        ["表 2.1(a)、第 3.2 節", "局部構材及外部被覆物設計風壓計算式"],
        [gcp_figure, "外牆 GCp 與 Zone 區域"],
    ])

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ============================================================
# 8. 結果
# ============================================================

st.subheader("設計風壓結果")

q_col1, q_col2, q_col3 = st.columns(3)

with q_col1:
    st.metric(
        "控制風速壓 q",
        f"{q_control:.3f} kgf/m²",
        f"{q_control * KGF_M2_TO_KPA:.3f} kPa",
    )

with q_col2:
    st.metric(
        "控制高度",
        f"{control_height:.3f} m",
    )

with q_col3:
    st.metric(
        "Kzt（控制高度）",
        f"{Kzt_control:.3f}",
    )

col4, col5 = st.columns(2, gap="large")

for col, zone in [(col4, 4), (col5, 5)]:
    r = zone_results[zone]
    with col:
        with st.container(border=True):
            st.subheader(
                "Zone 4－一般外牆區"
                if zone == 4
                else "Zone 5－角隅區"
            )

            m1, m2 = st.columns(2)

            with m1:
                st.metric(
                    "正壓",
                    f"{r['p+']:.3f} kgf/m²",
                    f"{r['p+ kPa']:.3f} kPa",
                )

            with m2:
                st.metric(
                    "負壓",
                    f"{r['p-']:.3f} kgf/m²",
                    f"{r['p- kPa']:.3f} kPa",
                )

            st.caption(
                f"GCp(+) = {r['GCp+']:.3f} ｜ "
                f"GCp(-) = {r['GCp-']:.3f}"
            )

st.caption(
    "Zone 4 與 Zone 5 的外牆正壓採同一條 GCp(+) 曲線，"
    "因此在相同內壓條件下正設計風壓相同；"
    "Zone 4 / Zone 5 的差異主要出現在負壓。"
)

st.caption(
    f"內風壓係數：{enclosure} → GCpi = ±{gcpi_abs:.3f}"
)

st.info(
    f"{city} {district} ｜ V10(C) = {v10c:.3f} m/s ｜ "
    f"控制風速壓 q = {q_control:.3f} kgf/m² "
    f"({q_control * KGF_M2_TO_KPA:.3f} kPa) ｜ "
    f"單元 A = {unit_area:.3f} m² ｜ "
    f"建築平面 = {building_x:.3f} × {building_y:.3f} m ｜ "
    f"Zone 5 角隅區寬度 a = {a_zone:.3f} m"
)

if DOCX_AVAILABLE:
    try:
        report_ctx = {
            "city": city, "district": district, "v10c": v10c, "h": h,
            "terrain": terrain, "importance": importance, "enclosure": enclosure,
            "gcpi_abs": gcpi_abs, "kzt_mode": kzt_mode, "topo_type": topo_type,
            "H": H, "Lh": Lh, "x_topo": x_topo, "Kh": Kh, "Kzt_h": Kzt_h,
            "qh": qh, "control_height": control_height, "K_control": K_control,
            "Kzt_control": Kzt_control, "q_control": q_control, "gcp_mode": gcp_mode,
            "unit_width": unit_width, "unit_height": unit_height, "unit_area": unit_area,
            "building_x": building_x, "building_y": building_y, "bmin": bmin, "bmax": bmax,
            "a_zone": a_zone, "zone_results": zone_results,
        }
        report_bytes = build_word_report(report_ctx)
        report_download_slot.download_button(
            "匯出 Word 計算報告",
            data=report_bytes,
            file_name=sanitize_docx_filename(report_filename_input),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
            key="download_wind_pressure_word_report",
        )
    except Exception as report_error:
        report_download_slot.error(f"Word 報告建立失敗：{report_error}")
else:
    report_download_slot.warning(
        "Word 匯出需要 python-docx。請在 requirements.txt 加入：python-docx"
    )


# ============================================================
# 9. 詳細資料
# ============================================================

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    [
        "計算摘要",
        "GCp 與單元受風面積",
        "地形係數 Kzt",
        "3D Zone 視覺化",
        "公式與中間值",
    ]
)

with tab1:
    rows = [
        ["工址", f"{city} {district}", "-"],
        ["V10(C)", v10c, "m/s"],
        ["計算採用地況", terrain, "-"],
        ["建築物高度 h", h, "m"],
        ["用途係數 I", importance, "-"],
        ["Kzt 採用方式", kzt_mode, "-"],
        ["GCp 採用方式", gcp_mode, "-"],
        ["K(h)", Kh, "-"],
        ["Kzt(h)", Kzt_h, "-"],
        ["q(h)", qh, "kgf/m²"],
        ["控制高度", control_height, "m"],
        ["控制 K", K_control, "-"],
        ["控制 Kzt", Kzt_control, "-"],
        ["控制風速壓 q", q_control, "kgf/m²"],
        ["控制風速壓 q", q_control * KGF_M2_TO_KPA, "kPa"],
        ["單元寬度 W", unit_width, "m"],
        ["單元高度 H", unit_height, "m"],
        ["單元受風面積 A", unit_area, "m²"],
        ["建築物平面 X 向尺寸 Bx", building_x, "m"],
        ["建築物平面 Y 向尺寸 By", building_y, "m"],
        ["建築物較短邊 Bmin", bmin, "m"],
        ["建築物較長邊 Bmax", bmax, "m"],
        ["Zone 4 正壓", zone_results[4]["p+"], "kgf/m²"],
        ["Zone 4 正壓（換算）", zone_results[4]["p+ kPa"], "kPa"],
        ["Zone 4 負壓", zone_results[4]["p-"], "kgf/m²"],
        ["Zone 4 負壓（換算）", zone_results[4]["p- kPa"], "kPa"],
        ["Zone 5 正壓", zone_results[5]["p+"], "kgf/m²"],
        ["Zone 5 正壓（換算）", zone_results[5]["p+ kPa"], "kPa"],
        ["Zone 5 負壓", zone_results[5]["p-"], "kgf/m²"],
        ["Zone 5 負壓（換算）", zone_results[5]["p- kPa"], "kPa"],
    ]

    summary_df = pd.DataFrame(rows, columns=["項目", "數值", "單位"])

    def format_display_value(value):
        if isinstance(value, (int, float, np.integer, np.floating)):
            return f"{float(value):.3f}"
        return value

    summary_df["數值"] = summary_df["數值"].map(format_display_value)

    st.dataframe(
        summary_df,
        use_container_width=True,
        hide_index=True,
    )

with tab2:
    st.subheader("GCp 採用值")

    if gcp_mode == "直接採保守值":
        st.write(
            "本案 GCp 採規範圖中的保守值，"
            "不依面積折減。"
        )
    else:
        st.write(
            f"本案以單元面積 A = W × H = "
            f"{unit_area:.3f} m² 查 GCp 圖，"
            "並依圖表的對數面積軸進行內插。"
        )

    gcps = pd.DataFrame(
        [
            {
                "區域": "Zone 4",
                "GCp 正值": zone_results[4]["GCp+"],
                "GCp 負值": zone_results[4]["GCp-"],
                "設計正壓 (kgf/m²)": zone_results[4]["p+"],
                "設計負壓 (kgf/m²)": zone_results[4]["p-"],
                "正壓換算 (kPa)": zone_results[4]["p+ kPa"],
                "負壓換算 (kPa)": zone_results[4]["p- kPa"],
            },
            {
                "區域": "Zone 5",
                "GCp 正值": zone_results[5]["GCp+"],
                "GCp 負值": zone_results[5]["GCp-"],
                "設計正壓 (kgf/m²)": zone_results[5]["p+"],
                "設計負壓 (kgf/m²)": zone_results[5]["p-"],
                "正壓換算 (kPa)": zone_results[5]["p+ kPa"],
                "負壓換算 (kPa)": zone_results[5]["p- kPa"],
            },
        ]
    )

    for col in gcps.columns:
        if col != "區域":
            gcps[col] = gcps[col].map(lambda x: f"{float(x):.3f}")

    st.dataframe(
        gcps,
        use_container_width=True,
        hide_index=True,
    )

    st.caption(
        "外牆 Zone 4 與 Zone 5 的正壓共用同一條 GCp(+) 曲線；"
        "負壓 GCp(-) 才分為 Zone 4 與 Zone 5。"
    )

    st.subheader("GCp 查圖")

    fig_gcp = create_gcp_chart_figure(
        area=unit_area,
        building_height=h,
        gcp_mode=gcp_mode,
    )

    gcp_left, gcp_center, gcp_right = st.columns([0.7, 2.6, 0.7])

    with gcp_center:
        st.pyplot(
            fig_gcp,
            use_container_width=True,
        )

    plt.close(fig_gcp)

    if gcp_mode == "依單元面積 A = W × H 查圖":
        st.caption(
            f"目前查值面積 A = {unit_area:.3f} m²；"
            f"Zone 4 / 5 正壓 GCp = {zone_results[4]['GCp+']:.3f}；"
            f"Zone 4 負壓 GCp = {zone_results[4]['GCp-']:.3f}；"
            f"Zone 5 負壓 GCp = {zone_results[5]['GCp-']:.3f}。"
        )

        if h > 18.0:
            st.info(
                "h > 18 m 時，本工具依工程討論後採用圖 3.2 取值："
                "Zone 4：GCp(+) 由 +1.800 內插至 +1.250，"
                "GCp(-) 由 -1.850 內插至 -1.500；"
                "Zone 5：GCp(+) 由 +1.800 內插至 +1.250，"
                "GCp(-) 由 -3.850 內插至 -2.150。"
                "A ≤ 2 m² 採保守端值，A ≥ 50 m² 採大面積端值。"
            )
        else:
            st.info(
                "圖 3.1 的外牆正壓與 Zone 4 負壓亦為不同曲線，"
                "不可直接假設兩者只差正負號。"
            )
    else:
        st.caption(
            "目前採保守值模式；圖上的標記位於規範曲線左側控制平台，"
            "單元面積不參與 GCp 折減。"
        )

    st.subheader("單元受風面積")

    st.latex(r"A=W\times H")
    st.write(
        f"W = **{unit_width:.3f} m**，"
        f"H = **{unit_height:.3f} m**，"
        f"因此 A = **{unit_area:.3f} m²**。"
    )

    if gcp_mode == "直接採保守值":
        st.caption(
            "本模式採保守 GCp，因此單元面積不參與 GCp 折減。"
        )
    else:
        st.caption(
            "本模式下，單元面積用於 GCp 圖表查值與內插。"
        )

    st.write(
        f"Bmin = **{bmin:.3f} m**，Zone 5 角隅區寬度 **a = {a_zone:.3f} m**。"
    )

with tab3:
    st.subheader("Kzt 採用方式")
    st.write(f"本案採用：**{kzt_mode}**")

    if kzt_mode == "保守計算 Kzt = 1.100":
        st.success("保守計算：Kzt = 1.100。")
        st.caption(
            "此值為額外保守計算設定；規範一般地形預設仍採 Kzt = 1.000。"
        )

    elif kzt_mode == "規範一般地形 Kzt = 1.000":
        st.success("一般地形：Kzt = 1.000。")

    else:
        st.latex(r"K_{zt}=(1+K_1K_2K_3)^2")

        kzt_df = pd.DataFrame(
            [
                {
                    "高度": "屋頂高度 h",
                    "z (m)": h,
                    "K1": K1_h,
                    "K2": K2_h,
                    "K3": K3_h,
                    "Kzt": Kzt_h,
                },
                {
                    "高度": "控制高度",
                    "z (m)": control_height,
                    "K1": K1_control,
                    "K2": K2_control,
                    "K3": K3_control,
                    "Kzt": Kzt_control,
                },
            ]
        )

        for col in ["z (m)", "K1", "K2", "K3", "Kzt"]:
            kzt_df[col] = kzt_df[col].map(
                lambda x: f"{float(x):.3f}"
            )

        st.dataframe(
            kzt_df,
            use_container_width=True,
            hide_index=True,
        )

        st.write(
            f"地形：**{topo_type}** ｜ "
            f"H = **{H:.3f} m** ｜ "
            f"Lh = **{Lh:.3f} m** ｜ "
            f"x = **{x_topo:.3f} m**"
        )

        st.caption(
            "K1 由表 2.3(a) 依 H/Lh 內插；"
            "K2 由表 2.3(b) 依 x/Lh 內插；"
            "K3 由表 2.3(c) 依 z/Lh 內插。"
        )


with tab4:
    st.subheader("Zone 4 / Zone 5 3D 視覺化")

    pressure_mode = st.radio(
        "顯示風壓",
        ["負壓", "正壓"],
        horizontal=True,
        key="zone_3d_pressure_mode",
    )

    fig_zone = create_zone_3d_figure(
        building_height=h,
        building_x=building_x,
        building_y=building_y,
        zone5_width=a_zone,
        zone_results=zone_results,
        pressure_mode=pressure_mode,
    )

    zone_left, zone_center, zone_right = st.columns([0.6, 3.0, 0.6])

    with zone_center:
        st.pyplot(
            fig_zone,
            use_container_width=True,
        )

    plt.close(fig_zone)

    st.caption(
        f"建築物平面：Bx = {building_x:.3f} m，"
        f"By = {building_y:.3f} m。"
        "藍色為 Zone 4、紅色為 Zone 5；"
        "尺寸線直接標示各 Zone 4 / Zone 5 的實際寬度，"
        "右側圖例顯示目前設計風壓。"
    )


with tab5:
    terrain_data = TERRAIN[terrain]

    st.subheader("內風壓係數 GCpi")

    gcpi_table = pd.DataFrame(
        [
            {
                "建築物類型": "開放式建築物",
                "GCpi": "0.000",
            },
            {
                "建築物類型": "部分封閉式建築物",
                "GCpi": "±1.146",
            },
            {
                "建築物類型": "封閉式建築物",
                "GCpi": "±0.375",
            },
        ]
    )

    st.dataframe(
        gcpi_table,
        use_container_width=True,
        hide_index=True,
    )

    st.write(
        f"本案選擇：**{enclosure}**，因此採用 "
        f"**GCpi = ±{gcpi_abs:.3f}**。"
    )

    st.caption(
        "GCpi 是內風壓係數，來自規範表 2.17。"
        "設計時正、負內壓兩種情況都必須分別考慮。"
    )

    st.divider()

    st.latex(
        r"K(z)=2.774\left(\frac{z}{z_g}\right)^{2\alpha}"
    )
    st.write(
        f"地況 {terrain}：α = {terrain_data['alpha']:.3f}，"
        f"zg = {terrain_data['zg']:.3f} m"
    )

    st.latex(
        r"q(z)=0.06K(z)K_{zt}(z)[I\,V_{10}(C)]^2"
    )

    st.write(
        f"本案整面帷幕牆採控制風速壓："
        f" **q = {q_control:.3f} kgf/m² "
        f"({q_control * KGF_M2_TO_KPA:.3f} kPa)**，"
        f"控制高度 = **{control_height:.3f} m**。"
    )

    for zone in (4, 5):
        r = zone_results[zone]

        st.markdown(f"**Zone {zone}**")

        st.write("控制正壓：正外壓搭配負內壓")
        st.code(
            f"p(+) = {q_external:.3f} × ({r['GCp+']:.3f}) "
            f"- {qi_negative_internal:.3f} × ({-gcpi_abs:.3f}) "
            f"= {r['p+']:.3f} kgf/m² "
            f"= {r['p+ kPa']:.3f} kPa"
        )

        st.write("控制負壓：負外壓搭配正內壓")
        st.code(
            f"p(-) = {q_external:.3f} × ({r['GCp-']:.3f}) "
            f"- {qi_positive_internal:.3f} × ({gcpi_abs:.3f}) "
            f"= {r['p-']:.3f} kgf/m² "
            f"= {r['p- kPa']:.3f} kPa"
        )
