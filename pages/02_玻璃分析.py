import math
import io
import re
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================
# 玻璃面材分析
# 對應原始 Excel：玻璃面材及背擋V2.2a.xlsm
# 工作表：4-S GLASS jp + 參考資料
#
# 原則：
# 1. 不修正原 Excel 的儲存格引用與公式邏輯。
# 2. 4-S GLASS jp 的輸入欄位、常數、查表及計算順序照原檔。
# 3. 參考資料工作表的查表數據內建於程式，不在畫面顯示公式或程式碼。
# 4. 原始 Excel 內已經是 #REF! 的公式不自行猜測或補值。
# ============================================================

st.set_page_config(
    page_title="玻璃面材分析",
    page_icon="🪟",
    layout="wide",
)

# ------------------------------------------------------------
# Excel 參考資料：4-S GLASS jp!B16:G21
# ------------------------------------------------------------
GLASS_TYPES = [
    "普通玻璃(膠合)",
    "熱硬化玻璃(膠合)",
    "強化玻璃(膠合)",
    "普通玻璃(非膠合)",
    "熱硬化玻璃(非膠合)",
    "強化玻璃(非膠合)",
]

K1_TABLE = {
    "普通玻璃(膠合)": 0.75,
    "熱硬化玻璃(膠合)": 2.00,
    "強化玻璃(膠合)": 3.50,
    "普通玻璃(非膠合)": 0.75,
    "熱硬化玻璃(非膠合)": 2.00,
    "強化玻璃(非膠合)": 3.50,
}

# ------------------------------------------------------------
# 參考資料!J33:K45
# 非膠合玻璃厚度查表（VLOOKUP TRUE）
# ------------------------------------------------------------
MONOLITHIC_EFFECTIVE_THICKNESS = [
    (0.0, 0.0),
    (2.5, 2.16),
    (2.7, 2.59),
    (3.0, 2.92),
    (4.0, 3.78),
    (5.0, 4.57),
    (6.0, 5.56),
    (8.0, 7.42),
    (10.0, 9.02),
    (12.0, 11.91),
    (16.0, 15.09),
    (19.0, 18.26),
    (22.0, 21.44),
]

# ------------------------------------------------------------
# 參考資料!J47:K54
# 膠合玻璃厚度查表（VLOOKUP TRUE）
# K48:K54 在原 Excel 中由公式得到；此處使用其相同計算結果。
# ------------------------------------------------------------
LAMINATED_EFFECTIVE_THICKNESS = [
    (0.0, 0.0),
    (6.0, 2.0 * 2.92),
    (8.0, 2.0 * 3.78),
    (10.0, 2.0 * 4.57),
    (12.0, 2.0 * 5.56),
    (14.0, 5.56 + 7.42),
    (16.0, 2.0 * 7.42),
    (20.0, 2.0 * 9.02),
]

# ------------------------------------------------------------
# 參考資料其餘表格亦保留於程式內，確保資料基準與 Excel 一致。
# 這些資料目前未直接被 4-S GLASS jp 的有效公式引用。
# ------------------------------------------------------------
BETA_G_TABLE = [
    (1.0, 3.1, 6.3),
    (0.5, 4.1, 8.2),
    (0.0, 5.5, 11.1),
    (-0.5, 7.3, 14.0),
    (-1.0, 8.0, 14.0),
]

SUPPORT_TYPE_TABLE = {
    "四邊固定": 0,
    "四點支承": 1,
    "六點支承": 2,
    "一邊點支承一邊有槽": 3,
    "角落點支承對角兩邊有槽": 4,
    "角落四點支承": 5,
}

G2_G3_TABLE = [
    ("簡支承受均佈力", 3.6, 1.4, 2, "w"),
    ("簡支承受均佈力，弱軸抗彎", 6.1, 1.8, 2, "w"),
    ("兩端固定受均佈力", 4.1, 4.9, 2, "w"),
    ("兩端固定受均佈力，弱軸抗彎", 5.4, 5.2, 2, "w"),
    ("簡支承受集中力", 4.2, 1.7, 2, "P"),
    ("簡支承受集中力，弱軸抗彎", 6.7, 2.6, 2, "P"),
    ("兩端固定受集中力", 5.3, 4.5, 2, "P"),
    ("兩端固定受集中力，弱軸抗彎", 6.5, 5.3, 2, "P"),
    ("簡支承受兩集中力", 3.3, 1.3, 2, "P"),
    ("懸臂受集中力", 4.0, 2.0, 1, "P"),
    ("懸臂受均佈力", 6.4, 2.0, 1, "w"),
]

MATERIAL_REFERENCE = {
    "花崗石": {"M": 84.0, "N": 84.0, "O": 84.0, "E": 490000.0, "nu": 0.28},
    "普通玻璃": {"M": 168.0, "N": 236.6, "O": 168.0, "E": 734000.0, "nu": 0.22},
    "熱硬化玻璃": {"M": 336.0, "N": 472.5, "O": 336.0, "E": 734000.0, "nu": 0.22},
    "強化玻璃": {"M": 672.0, "N": 945.0, "O": 672.0, "E": 734000.0, "nu": 0.22},
    "不鏽鋼": {"M": 1136.0, "N": None, "O": None, "E": 2040000.0, "nu": 0.33},
    "A36": {"M": 1512.0, "N": None, "O": None, "E": 2040000.0, "nu": 0.33},
    "A572": {"M": 1980.0, "N": None, "O": None, "E": 2040000.0, "nu": 0.33},
}

PLATE_TABLE_1 = [
    (0.00, 0.12500, 0.0130200),
    (0.05, 0.12425, 0.0130075),
    (0.10, 0.12350, 0.0129950),
    (0.15, 0.12275, 0.0129825),
    (0.20, 0.12200, 0.0129700),
    (0.25, 0.12300, 0.0128200),
    (0.29, 0.11900, 0.0125250),
    (0.33, 0.11800, 0.0122300),
    (0.365, 0.11400, 0.0117050),
    (0.40, 0.11150, 0.0111800),
    (0.45, 0.10500, 0.0106550),
    (0.50, 0.10000, 0.0101300),
    (0.55, 0.09340, 0.0094000),
    (0.60, 0.08658, 0.0086700),
    (0.65, 0.08040, 0.0079600),
    (0.70, 0.07420, 0.0072700),
    (0.75, 0.06830, 0.0066300),
    (0.80, 0.06280, 0.0060300),
    (0.85, 0.05760, 0.0054700),
    (0.90, 0.05280, 0.0049600),
    (0.95, 0.04830, 0.0044900),
    (1.00, 0.04420, 0.0040600),
]

THETA_ETA_TABLE = [
    (0, 1.000),
    (5, 1.000),
    (10, 0.960),
    (15, 0.940),
    (20, 0.920),
    (30, 0.880),
    (40, 0.840),
    (50, 0.810),
    (60, 0.780),
    (70, 0.755),
    (80, 0.730),
    (90, 0.705),
    (100, 0.680),
    (110, 0.665),
    (120, 0.650),
    (135, 0.630),
    (150, 0.610),
    (175, 0.590),
    (200, 0.570),
    (225, 0.555),
    (250, 0.540),
    (275, 0.530),
    (300, 0.520),
    (325, 0.515),
    (350, 0.510),
    (375, 0.505),
    (400, 0.500),
]

PLATE_TABLE_2 = [
    (0.00, 0.12500, 0.0130200),
    (0.05, 0.12525, 0.0130575),
    (0.10, 0.12550, 0.0130950),
    (0.15, 0.12575, 0.0131325),
    (0.20, 0.12600, 0.0131700),
    (0.25, 0.12650, 0.0132600),
    (0.30, 0.12700, 0.0133500),
    (0.35, 0.12800, 0.0135100),
    (0.40, 0.12900, 0.0136700),
    (0.45, 0.12950, 0.0139200),
    (0.50, 0.13000, 0.0141700),
    (0.55, 0.13200, 0.0145100),
    (0.60, 0.13400, 0.0149600),
    (0.65, 0.13600, 0.0155500),
    (0.70, 0.13800, 0.0163000),
    (0.75, 0.14000, 0.0172500),
    (0.80, 0.14200, 0.0184200),
    (0.85, 0.14500, 0.0198400),
    (0.90, 0.14800, 0.0215700),
    (0.95, 0.15100, 0.0236300),
    (1.00, 0.15400, 0.0260300),
]

KM_KF_TABLE = [
    (0.10, 0.1303, 0.01417),
    (0.20, 0.1303, 0.01417),
    (0.30, 0.1303, 0.01417),
    (0.40, 0.1303, 0.01417),
    (0.50, 0.1303, 0.01417),
    (0.55, 0.1317, 0.01451),
    (0.60, 0.1335, 0.01496),
    (0.65, 0.1355, 0.01555),
    (0.70, 0.1376, 0.01630),
    (0.75, 0.1398, 0.01725),
    (0.80, 0.1423, 0.01842),
    (0.85, 0.1449, 0.01984),
    (0.90, 0.1477, 0.02157),
    (0.95, 0.1506, 0.02363),
    (1.00, 0.1536, 0.02603),
]

# 參考資料!C61:M86
REFERENCE_C61_M86 = [
    (0.10, 2.175, 0.938, 0.904, 0.888, 0.5000, 0.754, 0.990, 0.3100, 0.3080, 0.1050),
    (0.20, 2.182, 0.938, 0.904, 0.888, 0.6250, 0.758, 0.990, 0.3100, 0.3080, 0.1500),
    (0.30, 2.194, 0.938, 0.904, 0.888, 0.7000, 0.764, 0.990, 0.3100, 0.3080, 0.1950),
    (0.40, 2.213, 0.980, 0.904, 0.888, 0.7800, 0.771, 1.140, 0.3100, 0.3080, 0.2400),
    (0.50, 2.268, 1.063, 0.904, 0.888, 0.8750, 0.780, 1.290, 0.3100, 0.3080, 0.2700),
    (0.55, 2.363, 1.1255, 1.0445, 1.015, 0.9375, 0.793, 1.380, 0.3560, 0.3525, 0.2925),
    (0.60, 2.458, 1.188, 1.185, 1.142, 1.0000, 0.807, 1.470, 0.4020, 0.3970, 0.3150),
    (0.65, 2.569, 1.313, 1.3185, 1.253, 1.0625, 0.821, 1.5525, 0.4475, 0.4375, 0.3450),
    (0.70, 2.713, 1.438, 1.452, 1.364, 1.1250, 0.837, 1.635, 0.4930, 0.4780, 0.3750),
    (0.75, 2.903, 1.6255, 1.543, 1.467, 1.2500, 0.855, 1.725, 0.5390, 0.5145, 0.4050),
    (0.80, 3.125, 1.813, 1.634, 1.570, 1.3750, 0.875, 1.815, 0.5850, 0.5510, 0.4350),
    (0.85, 3.394, 2.063, 1.9505, 1.662, 1.5000, 0.896, 1.8975, 0.6370, 0.5840, 0.4725),
    (0.90, 3.696, 2.313, 2.267, 1.754, 1.6250, 0.919, 1.980, 0.6890, 0.6170, 0.5100),
    (0.95, 4.044, 2.688, 2.664, 1.8605, 1.8125, 0.941, 2.070, 0.7625, 0.6470, 0.5475),
    (1.00, 4.472, 3.063, 3.061, 1.967, 2.0000, 0.962, 2.160, 0.8360, 0.6770, 0.5850),
    (1.10, 4.472, 4.063, 3.395, 2.588, 2.0000, 0.962, 2.325, 0.9960, 0.7530, 0.5850),
    (1.20, 4.472, 5.313, 5.453, 3.370, 2.0000, 0.962, 2.505, 1.1910, 0.8290, 0.5850),
    (1.30, 4.472, 6.938, 7.129, 4.348, 2.0000, 0.962, 2.670, 1.3570, 0.8940, 0.5850),
    (1.40, 4.472, 9.063, 9.189, 5.367, 2.0000, 0.962, 2.835, 1.5580, 0.9580, 0.5850),
    (1.50, 4.472, 11.688, 11.760, 6.524, 2.0000, 0.962, 3.000, 1.7730, 1.0180, 0.5850),
    (1.60, 4.472, 11.688, 14.870, 7.835, 2.0000, 0.962, 3.000, 2.0040, 0.8580, 0.5850),
    (1.70, 4.472, 11.688, 18.620, 9.254, 2.0000, 0.962, 3.000, 2.4300, 1.1220, 0.5850),
    (1.80, 4.472, 11.688, 23.180, 10.728, 2.0000, 0.962, 3.000, 2.5080, 1.1650, 0.5850),
    (1.90, 4.472, 11.688, 29.540, 12.475, 2.0000, 0.962, 3.000, 2.7800, 1.2020, 0.5850),
    (2.00, 4.472, 11.688, 34.780, 14.211, 2.0000, 0.962, 3.000, 3.0680, 1.2310, 0.5850),
]


# ------------------------------------------------------------
# Excel 型錯誤
# ------------------------------------------------------------
class ExcelError(Exception):
    def __init__(self, code):
        super().__init__(code)
        self.code = code


def is_error(x):
    return isinstance(x, ExcelError)


def excel_text(x):
    if isinstance(x, ExcelError):
        return x.code
    return x


def exact_vlookup(value, rows):
    for key, result in rows:
        if value == key:
            return result
    raise ExcelError("#N/A")


def approx_vlookup(value, rows):
    if not rows:
        raise ExcelError("#N/A")
    if value < rows[0][0]:
        raise ExcelError("#N/A")
    result = rows[0][1]
    for key, candidate in rows:
        if value >= key:
            result = candidate
        else:
            break
    return result


def excel_ln(x):
    if isinstance(x, ExcelError):
        raise x
    if x <= 0:
        raise ExcelError("#NUM!")
    return math.log(x)


def excel_exp(x):
    if isinstance(x, ExcelError):
        raise x
    try:
        return math.exp(x)
    except OverflowError:
        raise ExcelError("#NUM!")


def fmt(x, digits=3):
    if isinstance(x, ExcelError):
        return x.code
    if x is None:
        return "—"
    if isinstance(x, str):
        return x
    return f"{x:,.{digits}f}"


def calc_excel_model(t1, t2, glass_type_1, glass_type_2, a, b, q):
    """
    按 4-S GLASS jp 的儲存格公式順序建立結果。
    儲存格名稱保留在內部，畫面不顯示公式。
    """
    c = {}

    # Excel 固定值
    c["C2"] = t1
    c["C3"] = t2
    c["E2"] = glass_type_1
    c["E3"] = glass_type_2
    c["C5"] = a
    c["C6"] = b
    c["C7"] = q
    c["F4"] = 0.22
    c["F6"] = 734000.0

    # B16:G21 依原 Excel 公式建立動態列
    try:
        mono_t1 = approx_vlookup(t1, MONOLITHIC_EFFECTIVE_THICKNESS)
        mono_t2 = approx_vlookup(t2, MONOLITHIC_EFFECTIVE_THICKNESS)
        lam_t1 = approx_vlookup(t1, LAMINATED_EFFECTIVE_THICKNESS)
        lam_t2 = approx_vlookup(t2, LAMINATED_EFFECTIVE_THICKNESS)
    except ExcelError as e:
        mono_t1 = mono_t2 = lam_t1 = lam_t2 = e

    if t2 == 0:
        laminated_k2_t2 = 0.0
        nonlam_k2_t1 = 1.0
        nonlam_k2_t2 = 0.0
    else:
        laminated_k2_t2 = 0.75
        try:
            nonlam_k2_t1 = 0.75 * (1.0 + (t2 / t1) ** 3)
            nonlam_k2_t2 = 0.75 * (1.0 + (t1 / t2) ** 3)
        except ZeroDivisionError:
            nonlam_k2_t1 = ExcelError("#DIV/0!")
            nonlam_k2_t2 = ExcelError("#DIV/0!")

    rows = {
        "普通玻璃(膠合)": [0.75, 0.75, laminated_k2_t2, lam_t1, lam_t2],
        "熱硬化玻璃(膠合)": [2.00, 0.75, laminated_k2_t2, lam_t1, lam_t2],
        "強化玻璃(膠合)": [3.50, 0.75, laminated_k2_t2, lam_t1, lam_t2],
        "普通玻璃(非膠合)": [0.75, nonlam_k2_t1, nonlam_k2_t2, mono_t1, mono_t2],
        "熱硬化玻璃(非膠合)": [2.00, nonlam_k2_t1, nonlam_k2_t2, mono_t1, mono_t2],
        "強化玻璃(非膠合)": [3.50, nonlam_k2_t1, nonlam_k2_t2, mono_t1, mono_t2],
    }

    # F2 = VLOOKUP(E2,B16:G21,5,FALSE)
    c["F2"] = rows[glass_type_1][3]

    # F3 = VLOOKUP(E2,B16:G21,6,FALSE)
    # 注意：原 Excel 確實引用 E2，不是 E3；此處完全照原檔。
    c["F3"] = rows[glass_type_1][4]

    # 輔助儲存格
    try:
        denom = t1**3 + t2**3
        if denom == 0:
            raise ZeroDivisionError
        c["L3"] = t1**3 / denom
        c["L4"] = (1.0 if t2 == 0 else 1.1) * t1**3 / denom
        c["N4"] = t2**3 / denom
    except ZeroDivisionError:
        c["L3"] = c["L4"] = c["N4"] = ExcelError("#DIV/0!")

    # C4
    if t2 == 0:
        c["C4"] = t1
    else:
        c["C4"] = 0.95 * ((t1**3 + t2**3) ** (1.0 / 3.0))

    # F5 = C6/C5
    try:
        c["F5"] = b / a
    except ZeroDivisionError:
        c["F5"] = ExcelError("#DIV/0!")

    # F7
    try:
        c["F7"] = c["F6"] * (c["C4"] / 10.0) ** 3 / 12.0 / (1.0 - c["F4"] ** 2)
    except Exception:
        c["F7"] = ExcelError("#VALUE!")

    # C8 / D8 / F8
    thickness_factor_1 = 0.75 if t1 >= 20 else (0.8 if t1 >= 12 else (1.0 if t1 <= 8 else 0.9))
    c["C8"] = max(thickness_factor_1, rows[glass_type_1][0])
    c["D8"] = rows[glass_type_1][1]
    try:
        if is_error(c["D8"]):
            raise c["D8"]
        c["F8"] = (
            300.0 * c["C8"] * c["D8"]
            / (a / 1000.0 * b / 1000.0)
            * (t1 + (t1**2) / 4.0)
            * 0.102
        )
    except ExcelError as e:
        c["F8"] = e
    except ZeroDivisionError:
        c["F8"] = ExcelError("#DIV/0!")

    # C9 / D9 / F9
    thickness_factor_2 = 0.75 if t2 >= 20 else (0.8 if t2 >= 12 else (1.0 if t2 <= 8 else 0.9))
    c["C9"] = max(thickness_factor_2, rows[glass_type_2][0])
    c["D9"] = rows[glass_type_2][2]
    try:
        if is_error(c["D9"]):
            raise c["D9"]
        c["F9"] = (
            300.0 * c["C9"] * c["D9"]
            / (a / 1000.0 * b / 1000.0)
            * (t2 + (t2**2) / 4.0)
            * 0.102
        )
    except ExcelError as e:
        c["F9"] = e
    except ZeroDivisionError:
        c["F9"] = ExcelError("#DIV/0!")

    # ASTM coefficients
    if is_error(c["F5"]):
        c["D10"] = c["G10"] = c["D11"] = c["F5"]
    else:
        alpha = c["F5"]
        c["D10"] = 0.553 - 3.83 * alpha + 1.11 * alpha * alpha - 0.0969 * alpha * alpha * alpha
        c["G10"] = -2.29 + 5.83 * alpha - 2.17 * alpha * alpha + 0.2067 * alpha * alpha * alpha
        c["D11"] = 1.485 - 1.908 * alpha + 0.815 * alpha * alpha - 0.0822 * alpha * alpha * alpha

    # D12
    try:
        if is_error(c["L3"]):
            raise c["L3"]
        if is_error(c["F2"]):
            raise c["F2"]
        z1 = (c["L3"] * q / 10000.0) * (a * b) ** 2 / c["F6"] / (c["F2"] ** 4)
        c["D12"] = excel_ln(excel_ln(z1))
    except ExcelError as e:
        c["D12"] = e
    except ZeroDivisionError:
        c["D12"] = ExcelError("#DIV/0!")

    # G12
    if t2 == 0:
        c["G12"] = 0.0
    else:
        try:
            if is_error(c["N4"]):
                raise c["N4"]
            if is_error(c["F3"]):
                raise c["F3"]
            z2 = (c["N4"] * q / 10000.0) * (a * b) ** 2 / c["F6"] / (c["F3"] ** 4)
            c["G12"] = excel_ln(excel_ln(z2))
        except ExcelError as e:
            c["G12"] = e
        except ZeroDivisionError:
            c["G12"] = ExcelError("#DIV/0!")

    # F13 = MAX(...)
    try:
        for key in ["F2", "F3", "D10", "G10", "D11", "D12", "G12"]:
            if is_error(c[key]):
                raise c[key]

        d_out = c["F2"] * excel_exp(
            c["D10"] + c["G10"] * c["D12"] + c["D11"] * c["D12"] * c["D12"]
        ) / 10.0

        d_in = c["F3"] * excel_exp(
            c["D10"] + c["G10"] * c["G12"] + c["D11"] * c["G12"] * c["G12"]
        ) / 10.0

        c["F13"] = max(d_out, d_in)
        c["_d_out"] = d_out
        c["_d_in"] = d_in
    except ExcelError as e:
        c["F13"] = e
        c["_d_out"] = e
        c["_d_in"] = e

    return c


# ------------------------------------------------------------
# Word 報告
# ------------------------------------------------------------
def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_doc_font(doc):
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(9.5)
    if "Heading 1" in styles:
        styles["Heading 1"].font.size = Pt(12)
        styles["Heading 1"].paragraph_format.space_before = Pt(6)
        styles["Heading 1"].paragraph_format.space_after = Pt(3)


def _add_report_table(doc, rows, widths=None, header=None):
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    if header:
        cells = table.add_row().cells
        for i, text in enumerate(header):
            cells[i].text = text
            _set_cell_shading(cells[i], "D9EAF7")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[i].paragraphs[0].runs:
                run.bold = True

    for name, value, unit in rows:
        cells = table.add_row().cells
        cells[0].text = str(name)
        cells[1].text = fmt(value)
        cells[2].text = str(unit)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    return table


def build_word_report(
    t1,
    t2,
    glass_type_1,
    glass_type_2,
    a,
    b,
    q,
    result,
    sealant_result,
):
    doc = Document()
    _set_doc_font(doc)

    section = doc.sections[0]
    section.top_margin = Pt(30)
    section.bottom_margin = Pt(30)
    section.left_margin = Pt(40)
    section.right_margin = Pt(40)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("玻璃面材分析報告")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    doc.add_heading("一、輸入資料", level=1)
    _add_report_table(
        doc,
        [
            ("室外側玻璃種類", glass_type_1, ""),
            ("室外側玻璃厚度 t1", t1, "mm"),
            ("室內側玻璃種類", glass_type_2, ""),
            ("室內側玻璃厚度 t2", t2, "mm"),
            ("玻璃短邊尺寸 a", a, "mm"),
            ("玻璃長邊尺寸 b", b, "mm"),
            ("設計風壓 q", q, "kgf/m²"),
        ],
        header=("項目", "數值", "單位"),
    )

    doc.add_heading("二、主要結果", level=1)
    key_table = doc.add_table(rows=1, cols=3)
    key_table.style = "Table Grid"
    key_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = key_table.rows[0].cells
    for i, text in enumerate(["項目", "結果", "單位"]):
        hdr[i].text = text
        _set_cell_shading(hdr[i], "D9EAF7")
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    key_rows = [
        ("室外側最大容許風壓", result["F8"], "kgf/m²"),
        ("室內側最大容許風壓", result["F9"], "kgf/m²"),
        ("設計風壓作用下玻璃變形量 (ASTM)", result["F13"], "cm"),
    ]
    for name, value, unit in key_rows:
        cells = key_table.add_row().cells
        cells[0].text = name
        cells[1].text = fmt(value)
        cells[2].text = unit
        _set_cell_shading(cells[0], "FFF2CC")
        _set_cell_shading(cells[1], "FFF2CC")
        _set_cell_shading(cells[2], "FFF2CC")
        for run in cells[1].paragraphs[0].runs:
            run.bold = True

    doc.add_heading("三、基本資料與玻璃性質", level=1)
    _add_report_table(
        doc,
        [
            ("玻璃等效厚度 t", result["C4"], "mm"),
            ("波松比 ν", result["F4"], ""),
            ("短邊/長邊比 α", result["F5"], ""),
            ("楊氏係數 E", result["F6"], "kgf/cm²"),
            ("玻璃剛度 D", result["F7"], "kgf-cm"),
        ],
        header=("項目", "數值", "單位"),
    )

    doc.add_heading("四、k1 / k2 與容許風壓", level=1)
    _add_report_table(
        doc,
        [
            ("室外側 k1", result["C8"], ""),
            ("室外側 k2", result["D8"], ""),
            ("室外側最大容許風壓", result["F8"], "kgf/m²"),
            ("室內側 k1", result["C9"], ""),
            ("室內側 k2", result["D9"], ""),
            ("室內側最大容許風壓", result["F9"], "kgf/m²"),
        ],
        header=("項目", "數值", "單位"),
    )

    doc.add_heading("五、ASTM 計算結果", level=1)
    _add_report_table(
        doc,
        [
            ("ASTM coefficient r0", result["D10"], ""),
            ("ASTM coefficient r1", result["G10"], ""),
            ("ASTM coefficient r2", result["D11"], ""),
            ("ASTM coefficient x1", result["D12"], ""),
            ("ASTM coefficient x2", result["G12"], ""),
            ("設計風壓作用下玻璃變形量 (ASTM)", result["F13"], "cm"),
        ],
        header=("項目", "數值", "單位"),
    )

    doc.add_heading("六、結構膠計算", level=1)
    sealant_table = _add_report_table(
        doc,
        [
            ("設計風壓", sealant_result["q_kpa"], "kPa"),
            ("玻璃面積", sealant_result["glass_area_m2"], "m²"),
            ("玻璃重量", sealant_result["glass_weight_kg"], "kg"),
            ("結構膠施膠總長度", sealant_result["sealant_length_m"], "m"),
            ("風壓所需 Bite", sealant_result["wind_bite_mm"], "mm"),
            (
                "自重所需 Bite",
                sealant_result["dead_bite_mm"] if sealant_result["include_dead_load"] else "未計入",
                "mm" if sealant_result["include_dead_load"] else "",
            ),
            ("控制 Bite（理論值）", sealant_result["controlling_theoretical_bite_mm"], "mm"),
            ("採用 Bite", sealant_result["adopted_bite_mm"], "mm"),
            ("Glueline thickness", sealant_result["glueline_thickness_mm"], "mm"),
            ("Bite / Glueline thickness", sealant_result["bite_to_glueline_ratio"], ""),
            ("結構膠用量", sealant_result["sealant_volume_ml"], "mL"),
            ("結構膠用量", sealant_result["sealant_volume_l"], "L"),
            ("尺寸檢核", "PASS" if sealant_result["overall_ok"] else "FAIL", ""),
        ],
        header=("項目", "數值", "單位"),
    )

    # 將會直接影響結構膠用量與最終採用結果的項目特別標示
    sealant_mark_items = {
        "結構膠施膠總長度",
        "採用 Bite",
        "Glueline thickness",
        "Bite / Glueline thickness",
        "結構膠用量",
        "尺寸檢核",
    }

    # 第 0 列為表頭，因此由第 1 列開始檢查。
    for row in sealant_table.rows[1:]:
        item_name = row.cells[0].text.strip()
        if item_name in sealant_mark_items:
            for cell in row.cells:
                _set_cell_shading(cell, "FFF2CC")
            for run in row.cells[1].paragraphs[0].runs:
                run.bold = True

    # 實際用量（mL / L）再加粗項目名稱，讓報告中更容易辨識。
    for row in sealant_table.rows[1:]:
        if row.cells[0].text.strip() == "結構膠用量":
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def normalize_docx_filename(name):
    name = (name or "glass_panel_analysis_report").strip()
    name = re.sub(r'[\\/:*?"<>|]+', "_", name)
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name


# ------------------------------------------------------------
# 結構膠計算
# ------------------------------------------------------------
def calc_structural_silicone(
    t1,
    t2,
    a,
    b,
    q,
    glueline_thickness,
    sealant_length_m,
    include_dead_load=True,
):
    """
    Dow 結構膠尺寸計算：
    - 風壓所需 Bite：
      0.5 × 玻璃短邊(mm) × 風壓(kPa) / 138(kPa)
    - 自重所需 Bite：
      玻璃重量(kg) / [施膠總長度(m) × 700(kg/m²)]，再換算為 mm
    - Bite 最小 6 mm
    - Glueline thickness 最小 6 mm
    - Bite : Glueline thickness 介於 1:1 ~ 3:1
    """
    q_kpa = q * 0.00980665
    glass_area_m2 = (a / 1000.0) * (b / 1000.0)
    total_glass_thickness_m = (t1 + t2) / 1000.0
    glass_weight_kg = 2500.0 * total_glass_thickness_m * glass_area_m2

    wind_bite_mm = 0.0
    if a > 0:
        wind_bite_mm = 0.5 * a * q_kpa / 138.0

    dead_bite_mm = 0.0
    if include_dead_load and sealant_length_m > 0:
        dead_bite_mm = glass_weight_kg / (sealant_length_m * 700.0) * 1000.0

    controlling_theoretical_bite_mm = max(
        wind_bite_mm,
        dead_bite_mm if include_dead_load else 0.0,
        6.0,
    )

    # 手冊範例以整 mm 採用，因此向上取整。
    adopted_bite_mm = float(math.ceil(controlling_theoretical_bite_mm))

    if glueline_thickness > 0:
        bite_to_glueline_ratio = adopted_bite_mm / glueline_thickness
    else:
        bite_to_glueline_ratio = math.inf

    bite_min_ok = adopted_bite_mm >= 6.0
    glueline_min_ok = glueline_thickness >= 6.0
    ratio_ok = (
        glueline_thickness > 0
        and 1.0 <= bite_to_glueline_ratio <= 3.0
    )
    overall_ok = bite_min_ok and glueline_min_ok and ratio_ok

    # 1 mL = 1000 mm³；長度(m)×1000 = 長度(mm)
    # 因此體積(mL) = Bite(mm) × 膠厚(mm) × 長度(m)
    sealant_volume_ml = (
        adopted_bite_mm * glueline_thickness * sealant_length_m
        if sealant_length_m > 0 and glueline_thickness > 0
        else 0.0
    )

    return {
        "q_kpa": q_kpa,
        "glass_area_m2": glass_area_m2,
        "glass_weight_kg": glass_weight_kg,
        "sealant_length_m": sealant_length_m,
        "wind_bite_mm": wind_bite_mm,
        "dead_bite_mm": dead_bite_mm,
        "include_dead_load": include_dead_load,
        "controlling_theoretical_bite_mm": controlling_theoretical_bite_mm,
        "adopted_bite_mm": adopted_bite_mm,
        "glueline_thickness_mm": glueline_thickness,
        "bite_to_glueline_ratio": bite_to_glueline_ratio,
        "bite_min_ok": bite_min_ok,
        "glueline_min_ok": glueline_min_ok,
        "ratio_ok": ratio_ok,
        "overall_ok": overall_ok,
        "sealant_volume_ml": sealant_volume_ml,
        "sealant_volume_l": sealant_volume_ml / 1000.0,
    }


# ------------------------------------------------------------
# 頁面外觀
# ------------------------------------------------------------
st.markdown(
    """
    <style>
    .block-container {
        padding-top: 1.25rem;
        padding-bottom: 3rem;
        max-width: 1400px;
    }
    .app-subtitle {
        color: #6b7280;
        margin-top: -0.6rem;
        margin-bottom: 1.4rem;
    }
    .section-card {
        border: 1px solid rgba(128,128,128,.25);
        border-radius: 12px;
        padding: 1rem 1.1rem;
        margin-bottom: .9rem;
    }
    .section-heading {
        font-size: 1.08rem;
        font-weight: 700;
        margin-bottom: .6rem;
    }
    .result-row {
        display: grid;
        grid-template-columns: 1.55fr 1fr .7fr;
        gap: .8rem;
        padding: .48rem 0;
        border-bottom: 1px solid rgba(128,128,128,.14);
        align-items: baseline;
    }
    .result-row:last-child {
        border-bottom: 0;
    }
    .result-name {
        font-weight: 600;
    }
    .result-value {
        text-align: right;
        font-variant-numeric: tabular-nums;
        font-weight: 700;
    }
    .result-unit {
        color: #6b7280;
    }
    .excel-note {
        font-size: .86rem;
        color: #6b7280;
    }
    .key-result-card {
        border: 2px solid #c89b2b;
        border-radius: 14px;
        padding: 1rem 1.1rem;
        background: rgba(255, 193, 7, 0.08);
        min-height: 128px;
    }
    .key-result-title {
        font-size: .92rem;
        color: #ffffff !important;
        margin-bottom: .45rem;
        font-weight: 700;
    }
    .key-result-value {
        font-size: 1.75rem;
        line-height: 1.15;
        font-weight: 800;
        font-variant-numeric: tabular-nums;
        color: #ffffff !important;
    }
    .key-result-unit {
        font-size: .9rem;
        color: #ffffff !important;
        margin-top: .35rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("玻璃面材分析")

# ------------------------------------------------------------
# 輸入：只提供原 Excel 4-S GLASS jp 中的輸入儲存格
# ------------------------------------------------------------
with st.sidebar:
    st.markdown("#### Word 報告")
    report_filename_input = st.text_input(
        "報告檔名",
        value="glass_panel_analysis_report",
    )
    word_download_placeholder = st.empty()

    st.markdown("---")
    st.header("輸入資料")

    st.markdown("#### 室外側玻璃")
    t1 = st.number_input(
        "室外側玻璃厚度 t1 (mm)",
        min_value=0.0,
        value=5.0,
        step=0.1,
        format="%.3f",
    )
    glass_type_1 = st.selectbox(
        "室外側玻璃種類",
        GLASS_TYPES,
        index=4,
    )

    st.markdown("#### 室內側玻璃")
    t2 = st.number_input(
        "室內側玻璃厚度 t2 (mm)",
        min_value=0.0,
        value=5.0,
        step=0.1,
        format="%.3f",
    )
    glass_type_2 = st.selectbox(
        "室內側玻璃種類",
        GLASS_TYPES,
        index=4,
    )

    st.markdown("#### 玻璃尺寸")
    a = st.number_input(
        "玻璃短邊尺寸 a (mm)",
        min_value=0.0,
        value=1200.0,
        step=10.0,
        format="%.3f",
    )
    b = st.number_input(
        "玻璃長邊尺寸 b (mm)",
        min_value=0.0,
        value=3000.0,
        step=10.0,
        format="%.3f",
    )

    st.markdown("#### 設計風壓")
    q = st.number_input(
        "設計風壓 q (kgf/m²)",
        min_value=0.0,
        value=300.0,
        step=10.0,
        format="%.3f",
    )

    st.markdown("---")
    st.markdown("#### 結構膠")

    include_dead_load = st.checkbox(
        "結構膠承受玻璃自重",
        value=True,
    )

    glueline_thickness = st.number_input(
        "Glueline thickness 膠厚 (mm)",
        min_value=0.0,
        value=6.0,
        step=1.0,
        format="%.3f",
    )

    use_full_perimeter = st.checkbox(
        "使用玻璃四邊周長作為施膠總長度",
        value=True,
    )

    full_perimeter_m = 2.0 * (a + b) / 1000.0
    if use_full_perimeter:
        sealant_length_m = full_perimeter_m
        st.caption(f"施膠總長度：{sealant_length_m:.3f} m")
    else:
        sealant_length_m = st.number_input(
            "結構膠施膠總長度 (m)",
            min_value=0.0,
            value=max(full_perimeter_m, 0.0),
            step=0.1,
            format="%.3f",
        )

result = calc_excel_model(t1, t2, glass_type_1, glass_type_2, a, b, q)

sealant_result = calc_structural_silicone(
    t1=t1,
    t2=t2,
    a=a,
    b=b,
    q=q,
    glueline_thickness=glueline_thickness,
    sealant_length_m=sealant_length_m,
    include_dead_load=include_dead_load,
)

# ------------------------------------------------------------
# 主要結果
# ------------------------------------------------------------
st.markdown("### 主要結果")
k1, k2, k3 = st.columns(3, gap="large")
with k1:
    st.markdown(
        f"""
        <div class="key-result-card">
            <div class="key-result-title">室外側最大容許風壓</div>
            <div class="key-result-value">{fmt(result['F8'])}</div>
            <div class="key-result-unit">kgf/m²</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with k2:
    st.markdown(
        f"""
        <div class="key-result-card">
            <div class="key-result-title">室內側最大容許風壓</div>
            <div class="key-result-value">{fmt(result['F9'])}</div>
            <div class="key-result-unit">kgf/m²</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with k3:
    st.markdown(
        f"""
        <div class="key-result-card">
            <div class="key-result-title">設計風壓作用下玻璃變形量</div>
            <div class="key-result-value">{fmt(result['F13'])}</div>
            <div class="key-result-unit">cm</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ------------------------------------------------------------
# 結構膠主要結果
# ------------------------------------------------------------
st.markdown("### 結構膠主要結果")
s1, s2, s3 = st.columns(3, gap="large")

with s1:
    st.markdown(
        f"""
        <div class="key-result-card">
            <div class="key-result-title">結構膠採用 Bite</div>
            <div class="key-result-value">{fmt(sealant_result['adopted_bite_mm'])}</div>
            <div class="key-result-unit">mm</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with s2:
    st.markdown(
        f"""
        <div class="key-result-card">
            <div class="key-result-title">結構膠用量</div>
            <div class="key-result-value">{fmt(sealant_result['sealant_volume_ml'])}</div>
            <div class="key-result-unit">mL / 片</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

with s3:
    sealant_status_text = "PASS" if sealant_result["overall_ok"] else "FAIL"
    st.markdown(
        f"""
        <div class="key-result-card">
            <div class="key-result-title">結構膠尺寸檢核</div>
            <div class="key-result-value">{sealant_status_text}</div>
            <div class="key-result-unit">Bite / 膠厚 = {fmt(sealant_result['bite_to_glueline_ratio'])}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ------------------------------------------------------------
# 輸出：對應 Excel B2:G13
# ------------------------------------------------------------
left, right = st.columns([1, 1], gap="large")

with left:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">基本資料與玻璃性質</div>', unsafe_allow_html=True)

    rows_left = [
        ("室外側玻璃厚度 (t1)", t1, "mm"),
        ("室內側玻璃厚度 (t2)", t2, "mm"),
        ("玻璃等效厚度 (t)", result["C4"], "mm"),
        ("玻璃短邊尺寸 (a)", a, "mm"),
        ("玻璃長邊尺寸 (b)", b, "mm"),
        ("設計風壓 (q)", q, "kgf/m²"),
    ]

    for name, value, unit in rows_left:
        st.markdown(
            f"""
            <div class="result-row">
                <div class="result-name">{name}</div>
                <div class="result-value">{fmt(value)}</div>
                <div class="result-unit">{unit}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">k1 / k2 與容許風壓</div>', unsafe_allow_html=True)

    rows_strength = [
        ("室外側 k1", result["C8"], ""),
        ("室外側 k2", result["D8"], ""),
        ("室外側最大容許風壓", result["F8"], "kgf/m²"),
        ("室內側 k1", result["C9"], ""),
        ("室內側 k2", result["D9"], ""),
        ("室內側最大容許風壓", result["F9"], "kgf/m²"),
    ]

    for name, value, unit in rows_strength:
        st.markdown(
            f"""
            <div class="result-row">
                <div class="result-name">{name}</div>
                <div class="result-value">{fmt(value)}</div>
                <div class="result-unit">{unit}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

with right:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">彈性與幾何參數</div>', unsafe_allow_html=True)

    rows_right = [
        ("波松比 (ν)", result["F4"], ""),
        ("短邊/長邊比 (α)", result["F5"], ""),
        ("楊氏係數 (E)", result["F6"], "kgf/cm²"),
        ("玻璃剛度 (D)", result["F7"], "kgf-cm"),
    ]

    for name, value, unit in rows_right:
        st.markdown(
            f"""
            <div class="result-row">
                <div class="result-name">{name}</div>
                <div class="result-value">{fmt(value)}</div>
                <div class="result-unit">{unit}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">ASTM 計算結果</div>', unsafe_allow_html=True)

    rows_astm = [
        ("ASTM coefficient r0", result["D10"], ""),
        ("ASTM coefficient r1", result["G10"], ""),
        ("ASTM coefficient r2", result["D11"], ""),
        ("ASTM coefficient x1", result["D12"], ""),
        ("ASTM coefficient x2", result["G12"], ""),
        ("設計風壓作用下玻璃變形量 (ASTM)", result["F13"], "cm"),
    ]

    for name, value, unit in rows_astm:
        st.markdown(
            f"""
            <div class="result-row">
                <div class="result-name">{name}</div>
                <div class="result-value">{fmt(value)}</div>
                <div class="result-unit">{unit}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

# ------------------------------------------------------------
# 玻璃種類摘要
# ------------------------------------------------------------
st.markdown("### 玻璃配置")
c1, c2 = st.columns(2)
c1.info(f"室外側：{glass_type_1}，{t1:.3f} mm")
c2.info(f"室內側：{glass_type_2}，{t2:.3f} mm")

st.markdown("### 結構膠計算")
sg_left, sg_right = st.columns(2, gap="large")

with sg_left:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Bite 計算</div>', unsafe_allow_html=True)

    sealant_rows_left = [
        ("設計風壓", sealant_result["q_kpa"], "kPa"),
        ("風壓所需 Bite", sealant_result["wind_bite_mm"], "mm"),
        (
            "自重所需 Bite",
            sealant_result["dead_bite_mm"] if include_dead_load else "未計入",
            "mm" if include_dead_load else "",
        ),
        ("控制 Bite（理論值）", sealant_result["controlling_theoretical_bite_mm"], "mm"),
        ("採用 Bite", sealant_result["adopted_bite_mm"], "mm"),
    ]

    for name, value, unit in sealant_rows_left:
        st.markdown(
            f"""
            <div class="result-row">
                <div class="result-name">{name}</div>
                <div class="result-value">{fmt(value)}</div>
                <div class="result-unit">{unit}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

with sg_right:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">用量與尺寸檢核</div>', unsafe_allow_html=True)

    sealant_rows_right = [
        ("玻璃面積", sealant_result["glass_area_m2"], "m²"),
        ("玻璃重量", sealant_result["glass_weight_kg"], "kg"),
        ("施膠總長度", sealant_result["sealant_length_m"], "m"),
        ("Glueline thickness", sealant_result["glueline_thickness_mm"], "mm"),
        ("Bite / Glueline thickness", sealant_result["bite_to_glueline_ratio"], ""),
        ("結構膠用量", sealant_result["sealant_volume_ml"], "mL"),
        ("結構膠用量", sealant_result["sealant_volume_l"], "L"),
    ]

    for name, value, unit in sealant_rows_right:
        st.markdown(
            f"""
            <div class="result-row">
                <div class="result-name">{name}</div>
                <div class="result-value">{fmt(value)}</div>
                <div class="result-unit">{unit}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.markdown("</div>", unsafe_allow_html=True)

if sealant_result["overall_ok"]:
    st.success("結構膠尺寸檢核：PASS")
else:
    issues = []
    if not sealant_result["glueline_min_ok"]:
        issues.append("Glueline thickness 小於 6 mm")
    if not sealant_result["bite_min_ok"]:
        issues.append("Bite 小於 6 mm")
    if not sealant_result["ratio_ok"]:
        issues.append("Bite / Glueline thickness 不在 1.0～3.0 範圍內")
    st.error("結構膠尺寸檢核：FAIL" + ("；" + "；".join(issues) if issues else ""))

# ------------------------------------------------------------
# Word 匯出
# ------------------------------------------------------------
word_bytes = build_word_report(
    t1=t1,
    t2=t2,
    glass_type_1=glass_type_1,
    glass_type_2=glass_type_2,
    a=a,
    b=b,
    q=q,
    result=result,
    sealant_result=sealant_result,
)
report_filename = normalize_docx_filename(report_filename_input)

with word_download_placeholder:
    st.download_button(
        label="匯出 Word 報告",
        data=word_bytes,
        file_name=report_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

# ------------------------------------------------------------
# 原 Excel 錯誤狀態提示
# ------------------------------------------------------------
error_cells = []
for key in ["F2", "F3", "F8", "F9", "D12", "G12", "F13"]:
    if is_error(result.get(key)):
        error_cells.append(excel_text(result[key]))

if error_cells:
    st.warning(
        "目前輸入條件依原 Excel 公式會產生 Excel 計算錯誤："
        + "、".join(sorted(set(error_cells)))
        + "。本工具不自行修改原公式來消除該錯誤。"
    )
