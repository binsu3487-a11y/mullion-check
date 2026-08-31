from __future__ import annotations
from datetime import datetime
from io import BytesIO
from typing import Dict, List, Optional, Sequence, Tuple
from pathlib import Path
import re
import matplotlib
matplotlib.use("Agg")
import matplotlib.patches as patches
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
import pandas as pd
import streamlit as st

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt
except ImportError as exc:
    Document = None  # type: ignore[assignment]
    DOCX_IMPORT_ERROR: Optional[Exception] = exc
else:
    DOCX_IMPORT_ERROR = None

try:
    import openseespy.opensees as ops
    import opsvis as opsv
except ImportError as exc:  # 讓缺少套件時仍能顯示安裝說明
    ops = None  # type: ignore[assignment]
    opsv = None  # type: ignore[assignment]
    OPENSEES_IMPORT_ERROR: Optional[Exception] = exc
else:
    OPENSEES_IMPORT_ERROR = None


# ============================================================
# 1. 固定設定（kgf、cm）
# ============================================================

ELEMENTS_PER_INPUT_SEGMENT = 50
POSTPROCESS_POINTS_PER_ELEMENT = 50
COORDINATE_TOLERANCE = 1.0e-8
KPA_CM_TO_KGF_PER_CM = 1.0 / 98.0665

# Streamlit 網頁中的圖面顯示寬度；只影響網頁，不影響分析與下載檔。
MODEL_DISPLAY_WIDTH_PX = 400
DEFORMATION_DISPLAY_WIDTH_PX = 560
WORD_FIGURE_HEIGHT_CM = 10.2

# 材料資料庫：未來新增鋼材或其他材料時，可直接增加新的材料分類與牌號。
MATERIAL_DATABASE: Dict[str, Dict[str, object]] = {
    "鋁合金": {
        "E": 710100.0,
        "grades": {
            "6105-T5": {
                "Fty": 2450.0,
                "Ftu": 2660.0,
                "Fb_allow": 1485.0,
            },

            "6061-T6": {
                "Fty": 2450.0,
                "Ftu": 2660.0,
                "Fb_allow": 1485.0,
            },

            "6063-T5": {
                "Fty": 1050.0,
                "Ftu": 1470.0,
                "Fb_allow": 682.0,
            },

            "6063-T6": {
                "Fty": 1750.0,
                "Ftu": 2100.0,
                "Fb_allow": 1060.0,
            },

            "3003-H14": {
                "Fty": 1195.218,
                "Ftu": 1406.139,
                "Fb_allow": 717.131,
            },
        },
    },

    "碳鋼": {
        "E": 2039432.426,
        "grades": {
            "A36": {
                "Fty": 2531.050,
                "Ftu": 4077.804,
                "Fb_allow": 1670.493,
            },
        },
    },
}

STATION_LABELS = {
    "hinge": "Hinge 鉸支承",
    "pin": "Pin 內部鉸接",
    "free": "Free end 自由端",
}


# ============================================================
# 2. 共用函式
# ============================================================


def status_text(passed: bool) -> str:
    return "PASS（通過）" if passed else "FAIL（不通過）"


def y_key(y: float) -> int:
    return int(round(float(y) / COORDINATE_TOLERANCE))


def same_y(y1: float, y2: float) -> bool:
    return abs(float(y1) - float(y2)) <= COORDINATE_TOLERANCE


def unique_sorted_coordinates(values: Sequence[float]) -> List[float]:
    unique: Dict[int, float] = {}
    for value in values:
        unique[y_key(value)] = float(value)
    return sorted(unique.values())


def station_type_text(station_type: str) -> str:
    return STATION_LABELS[station_type]


def build_station_geometry(
    bottom_type: str,
    segment_lengths: Sequence[float],
    upper_types: Sequence[str],
) -> Tuple[List[Dict[str, object]], List[Dict[str, object]]]:
    """由 Streamlit 輸入資料建立控制點與輸入段資料。"""
    if bottom_type not in {"hinge", "free"}:
        raise ValueError("最下端只能是 Hinge 或 Free end。")

    if len(segment_lengths) == 0:
        raise ValueError("至少需要一段梁。")

    if len(segment_lengths) != len(upper_types):
        raise ValueError("段長與上端類型數量不一致。")

    if any((not np.isfinite(length)) or length <= 0.0 for length in segment_lengths):
        raise ValueError("所有段長都必須是大於 0 的有限數值。")

    stations: List[Dict[str, object]] = [
        {
            "station_number": 1,
            "y": 0.0,
            "type": bottom_type,
        }
    ]
    segments: List[Dict[str, object]] = []
    current_y = 0.0

    for index, (length, upper_type) in enumerate(
        zip(segment_lengths, upper_types),
        start=1,
    ):
        lower_station = stations[-1]
        lower_type = str(lower_station["type"])
        is_last = index == len(segment_lengths)

        if upper_type not in {"hinge", "pin", "free"}:
            raise ValueError(f"第 {index} 段上端類型無效。")
        if not is_last and upper_type == "free":
            raise ValueError("Free end 只能位於整根直料最上端或最下端。")
        if is_last and upper_type == "pin":
            raise ValueError("Pin 不能作為整根直料最上端。")
        if lower_type == "free" and upper_type != "hinge":
            raise ValueError("由下部 Free end 起始的第一段，上端必須先接 Hinge。")
        if is_last and lower_type == "pin" and upper_type == "free":
            raise ValueError("不可建立 Pin–Free 的不穩定頂部懸臂。")

        start_y = current_y
        current_y += float(length)

        upper_station = {
            "station_number": index + 1,
            "y": current_y,
            "type": upper_type,
        }
        stations.append(upper_station)

        segments.append(
            {
                "segment_number": index,
                "start_y": start_y,
                "end_y": current_y,
                "length": float(length),
                "lower_type": lower_type,
                "upper_type": upper_type,
            }
        )

    hinge_count = sum(str(item["type"]) == "hinge" for item in stations)
    if hinge_count < 2:
        raise ValueError(
            "模型至少需要兩個不同高程的 Hinge 支承，"
            "否則可能產生剛體轉動與矩陣奇異。"
        )

    hinge_number = 0
    pin_number = 0
    free_number = 0

    for station in stations:
        station_type = str(station["type"])
        if station_type == "hinge":
            hinge_number += 1
            station["name"] = f"H{hinge_number}"
            station["type_text"] = "鉸支承"
        elif station_type == "pin":
            pin_number += 1
            station["name"] = f"P{pin_number}"
            station["type_text"] = "內部 Pin"
        else:
            free_number += 1
            station["name"] = f"F{free_number}"
            station["type_text"] = "自由端"

    for segment in segments:
        segment_number = int(segment["segment_number"])
        lower_station = stations[segment_number - 1]
        upper_station = stations[segment_number]
        segment["lower_name"] = str(lower_station["name"])
        segment["upper_name"] = str(upper_station["name"])
        segment["name"] = f"{lower_station['name']}–{upper_station['name']}"
        segment["connection_type"] = (
            f"{str(lower_station['type']).title()}–"
            f"{str(upper_station['type']).title()}"
        )

    return stations, segments


def span_allowable_deflection(length_cm: float) -> Tuple[float, str]:
    if length_cm * 10.0 < 4115.0:
        return length_cm / 175.0, "L/175"
    return length_cm / 240.0 + 0.635, "L/240 + 0.635 cm"


def recommended_ix_for_simple_span(
    uniform_load_kgf_per_cm: float,
    span_length_cm: float,
    elastic_modulus_kgf_cm2: float,
) -> Tuple[float, float, str]:
    """以全跨均布載重簡支梁反算滿足容許變形所需的最低 Ix。

    簡支梁均布載重最大變形：
        delta_max = 5 w L^4 / (384 E I)

    因此：
        I_required = 5 w L^4 / (384 E delta_allow)
    """
    L = float(span_length_cm)
    w_ref = float(uniform_load_kgf_per_cm)
    E_ref = float(elastic_modulus_kgf_cm2)
    allowable, formula = span_allowable_deflection(L)

    if w_ref <= 0.0 or E_ref <= 0.0 or allowable <= 0.0:
        return 0.0, allowable, formula

    required_ix = 5.0 * w_ref * L**4 / (384.0 * E_ref * allowable)
    return required_ix, allowable, formula


def find_check_segment_key(
    y_mid: float,
    check_segments: Sequence[Dict[str, object]],
) -> Tuple[str, int]:
    for segment in check_segments:
        start_y = float(segment["start_y"])
        end_y = float(segment["end_y"])
        if start_y - COORDINATE_TOLERANCE <= y_mid <= end_y + COORDINATE_TOLERANCE:
            return str(segment["type"]), int(segment["index"])
    raise RuntimeError(f"無法判斷 Y={y_mid:.8f} cm 所屬的檢核區段。")


def build_check_segments(
    all_support_specs: Sequence[Dict[str, object]],
    total_height: float,
) -> List[Dict[str, object]]:
    ordered_supports = sorted(
        all_support_specs,
        key=lambda item: float(item["y"]),
    )
    if len(ordered_supports) < 2:
        raise ValueError("至少需要兩個 Hinge 支承。")

    segments: List[Dict[str, object]] = []
    first_support_y = float(ordered_supports[0]["y"])

    if first_support_y > COORDINATE_TOLERANCE:
        segments.append(
            {
                "type": "lower_cantilever",
                "index": 1,
                "name": "下部懸臂",
                "start_y": 0.0,
                "end_y": first_support_y,
                "length": first_support_y,
                "lower_support": None,
                "upper_support": str(ordered_supports[0]["name"]),
            }
        )

    span_index = 1
    for lower_support, upper_support in zip(
        ordered_supports[:-1],
        ordered_supports[1:],
    ):
        start_y = float(lower_support["y"])
        end_y = float(upper_support["y"])
        segments.append(
            {
                "type": "support_span",
                "index": span_index,
                "name": f"{lower_support['name']}–{upper_support['name']}",
                "start_y": start_y,
                "end_y": end_y,
                "length": end_y - start_y,
                "lower_support": str(lower_support["name"]),
                "upper_support": str(upper_support["name"]),
            }
        )
        span_index += 1

    last_support_y = float(ordered_supports[-1]["y"])
    if total_height - last_support_y > COORDINATE_TOLERANCE:
        segments.append(
            {
                "type": "upper_cantilever",
                "index": 1,
                "name": "頂部懸臂",
                "start_y": last_support_y,
                "end_y": total_height,
                "length": total_height - last_support_y,
                "lower_support": str(ordered_supports[-1]["name"]),
                "upper_support": None,
            }
        )

    return segments


def build_mesh_coordinates(
    input_segments: Sequence[Dict[str, object]],
    event_y_coordinates: Sequence[float],
) -> List[float]:
    coordinates: List[float] = [0.0]
    for segment in input_segments:
        start_y = float(segment["start_y"])
        length = float(segment["length"])
        for local_index in range(1, ELEMENTS_PER_INPUT_SEGMENT + 1):
            coordinates.append(
                start_y + length * local_index / ELEMENTS_PER_INPUT_SEGMENT
            )
    coordinates.extend(float(value) for value in event_y_coordinates)
    return unique_sorted_coordinates(coordinates)


def figure_to_png_bytes(figure: plt.Figure) -> bytes:
    buffer = BytesIO()
    figure.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    buffer.seek(0)
    return buffer.getvalue()


def dataframe_to_csv_bytes(dataframe: pd.DataFrame) -> bytes:
    return dataframe.round(3).to_csv(index=False, float_format="%.3f").encode("utf-8-sig")


def round_dataframe_for_display(dataframe: pd.DataFrame) -> pd.DataFrame:
    """只調整顯示與匯出精度，不改變分析內部原始數值。"""
    rounded = dataframe.copy()
    numeric_columns = rounded.select_dtypes(include=[np.number]).columns
    rounded[numeric_columns] = rounded[numeric_columns].round(3)
    return rounded


def sanitize_report_filename(raw_name: str) -> str:
    """清理 Windows 不允許的檔名字元，並補上 .docx。"""
    name = (raw_name or "").strip()
    if name.lower().endswith(".docx"):
        name = name[:-5]
    name = re.sub(r'[<>:"/\\|?*]+', "_", name)
    name = name.strip(" .")
    if not name:
        name = "mullion_analysis_report"
    return f"{name}.docx"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(
    cell,
    top: int = 8,
    start: int = 28,
    bottom: int = 8,
    end: int = 28,
) -> None:
    """設定 Word 表格儲存格內距；單位為 twip。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def _set_row_cant_split(row, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def _format_docx_number(value: object) -> str:
    """Word 報告中的數值統一顯示至小數點後 3 位，並移除 nan／科學記號。"""
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        number = float(value)
        if not np.isfinite(number):
            return "-"
        if abs(number) < 0.0005:
            number = 0.0
        return f"{number:.3f}"
    if isinstance(value, str):
        text = value.strip()
        if text.lower() in {"nan", "none", "null", "inf", "+inf", "-inf"}:
            return "-"
        compact = text.replace(",", "")
        if re.fullmatch(r"[+-]?(?:\d+(?:\.\d*)?|\.\d+)(?:[eE][+-]?\d+)?", compact):
            number = float(compact)
            if not np.isfinite(number):
                return "-"
            if abs(number) < 0.0005:
                number = 0.0
            # 純整數字串保留整數；含小數或科學記號者固定三位小數。
            if "." not in compact and "e" not in compact.lower():
                return str(int(number))
            return f"{number:.3f}"
        return text
    return str(value)


def _format_docx_paragraph(paragraph, font_size: float, bold: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(font_size + 0.6)
    for run in paragraph.runs:
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(font_size)
        run.bold = bold


def _set_cell_width(cell, width_cm: float) -> None:
    """固定 Word 儲存格寬度，避免表格自動撐滿整頁。"""
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def _add_docx_table(
    document,
    dataframe: pd.DataFrame,
    column_widths_cm: Optional[Sequence[float]] = None,
) -> None:
    """建立直向頁面用的窄版 Word 表格。

    寬欄表格會在呼叫端拆成兩張，這裡不再讓 Word 自動撐滿頁面。
    """
    display_df = dataframe.copy()
    column_count = len(display_df.columns)

    if column_widths_cm is None:
        total_width = 15.8
        column_widths_cm = [total_width / column_count] * column_count
    elif len(column_widths_cm) != column_count:
        raise ValueError("欄寬數量必須與表格欄數一致。")

    table = document.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    _set_row_cant_split(table.rows[0], repeat_header=True)
    header_cells = table.rows[0].cells
    for index, column in enumerate(display_df.columns):
        _set_cell_width(header_cells[index], float(column_widths_cm[index]))
        header_cells[index].text = str(column)
        _set_cell_shading(header_cells[index], "D9EAF7")
        _set_cell_margins(header_cells[index], top=18, start=38, bottom=18, end=38)
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[index].paragraphs:
            _format_docx_paragraph(paragraph, font_size=11.5, bold=True)

    for row in display_df.itertuples(index=False, name=None):
        added_row = table.add_row()
        _set_row_cant_split(added_row)
        cells = added_row.cells
        for index, value in enumerate(row):
            _set_cell_width(cells[index], float(column_widths_cm[index]))
            cells[index].text = _format_docx_number(value)
            _set_cell_margins(cells[index], top=14, start=38, bottom=14, end=38)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                _format_docx_paragraph(paragraph, font_size=11.0)

    # 保留極小的段落作為下一標題的定位點，避免預設空白過大。
    following = document.add_paragraph()
    following.paragraph_format.space_before = Pt(0)
    following.paragraph_format.space_after = Pt(0)
    following.paragraph_format.line_spacing = Pt(1)
    following.add_run("").font.size = Pt(1)




def build_word_report(result: Dict[str, object]) -> bytes:
    """將分析輸入、結果表格及四張圖輸出為 A4 直向 Word 報告。"""
    if Document is None:
        raise RuntimeError("尚未安裝 python-docx，無法建立 Word 報告。")

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft JhengHei"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = Pt(12.2)

    heading1_style = document.styles["Heading 1"]
    heading1_style.font.name = "Microsoft JhengHei"
    heading1_style.font.size = Pt(14)
    heading1_style.font.bold = True
    heading1_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    heading1_style.paragraph_format.space_before = Pt(2)
    heading1_style.paragraph_format.space_after = Pt(1)
    heading1_style.paragraph_format.keep_with_next = True

    title = document.add_heading("直立多段梁／帷幕直料分析報告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated = document.add_paragraph(
        f"報告產生時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.paragraph_format.space_before = Pt(0)
    generated.paragraph_format.space_after = Pt(2)

    overall_pass = bool(result["overall_pass"])
    stress_pass = bool(result["stress_pass"])
    deflection_pass = bool(result["deflection_pass"])
    verdict = document.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verdict.paragraph_format.space_before = Pt(0)
    verdict.paragraph_format.space_after = Pt(3)
    run = verdict.add_run(
        f"整體判定：{status_text(overall_pass)}　｜　"
        f"強軸應力：{status_text(stress_pass)}　｜　"
        f"容許變形：{status_text(deflection_pass)}"
    )
    run.bold = True
    run.font.size = Pt(12)

    summary = result["model_summary"]
    maximum = result["maximum_results"]
    assert isinstance(summary, dict)
    assert isinstance(maximum, dict)

    document.add_heading("1. 模型與輸入摘要", level=1)
    summary_df = pd.DataFrame(
        [{"項目": key, "數值": value} for key, value in summary.items()]
    )
    _add_docx_table(document, summary_df, [8.6, 5.6])

    document.add_heading("2. 最大分析結果", level=1)
    maximum_rows = [
        {
            "項目": "最大水平位移 |UX|",
            "數值": maximum["max_abs_displacement"],
            "單位": "cm",
            "位置 Y (cm)": maximum["max_displacement_y"],
        },
        {
            "項目": "最大彎矩 |M|",
            "數值": maximum["max_abs_moment"],
            "單位": "kgf-cm",
            "位置 Y (cm)": maximum["max_moment_y"],
        },
        {
            "項目": "最大剪力 |V|",
            "數值": maximum["max_abs_shear"],
            "單位": "kgf",
            "位置 Y (cm)": maximum["max_shear_y"],
        },
    ]
    lower_tip = maximum.get("lower_tip_ux")
    if lower_tip is not None and np.isfinite(float(lower_tip)):
        maximum_rows.append(
            {
                "項目": "下部懸臂端 UX",
                "數值": lower_tip,
                "單位": "cm",
                "位置 Y (cm)": 0.0,
            }
        )
    upper_tip = maximum.get("upper_tip_ux")
    if upper_tip is not None and np.isfinite(float(upper_tip)):
        maximum_rows.append(
            {
                "項目": "頂部懸臂端 UX",
                "數值": upper_tip,
                "單位": "cm",
                "位置 Y (cm)": summary.get("直料總高度 (cm)"),
            }
        )
    _add_docx_table(document, pd.DataFrame(maximum_rows), [5.4, 3.2, 2.3, 3.3])

    document.add_heading("3. 控制點配置", level=1)
    stations_df = pd.DataFrame(result["stations"])[
        ["station_number", "name", "type_text", "y"]
    ].rename(
        columns={
            "station_number": "控制點編號",
            "name": "名稱",
            "type_text": "類型",
            "y": "Y (cm)",
        }
    )
    _add_docx_table(document, stations_df, [3.5, 2.7, 4.5, 3.5])

    document.add_heading("4. 各輸入段", level=1)
    segment_source = pd.DataFrame(result["input_segments"])
    segments_basic_df = segment_source[
        ["segment_number", "name", "connection_type", "length"]
    ].rename(
        columns={
            "segment_number": "段號",
            "name": "控制點",
            "connection_type": "連接類型",
            "length": "長度 (cm)",
        }
    )
    _add_docx_table(document, segments_basic_df, [2.0, 3.5, 5.0, 3.7])

    segments_y_df = segment_source[
        ["segment_number", "start_y", "end_y"]
    ].rename(
        columns={
            "segment_number": "段號",
            "start_y": "起點 Y (cm)",
            "end_y": "終點 Y (cm)",
        }
    )
    _add_docx_table(document, segments_y_df, [2.8, 5.7, 5.7])

    document.add_heading("5. 各支承間距與懸臂變形檢核", level=1)
    span_source = pd.DataFrame(result["span_check_results"])
    span_response_df = span_source[
        ["name", "length", "control_y", "control_disp"]
    ].rename(
        columns={
            "name": "區段",
            "length": "長度 (cm)",
            "control_y": "控制位置 Y (cm)",
            "control_disp": "實際 |UX| (cm)",
        }
    )
    _add_docx_table(document, span_response_df, [3.0, 3.2, 4.2, 3.8])

    span_check_df = span_source[
        ["name", "formula", "allowable", "utilization", "result"]
    ].rename(
        columns={
            "name": "區段",
            "formula": "容許公式",
            "allowable": "容許值 (cm)",
            "utilization": "使用率",
            "result": "結果",
        }
    )
    _add_docx_table(document, span_check_df, [2.6, 3.2, 3.2, 2.2, 3.0])

    document.add_heading("6. 強軸應力檢核", level=1)
    stress_source = pd.DataFrame(result["stress_results"])
    stress_section_df = stress_source[["name", "A", "Ix", "Sx"]].rename(
        columns={
            "name": "斷面",
            "A": "A (cm²)",
            "Ix": "Ix (cm⁴)",
            "Sx": "Sx (cm³)",
        }
    )
    _add_docx_table(document, stress_section_df, [4.1, 3.2, 3.4, 3.5])

    stress_check_df = stress_source[
        ["name", "moment_share", "fb", "utilization", "result"]
    ].rename(
        columns={
            "name": "斷面",
            "moment_share": "分配彎矩 (kgf-cm)",
            "fb": "fb (kgf/cm²)",
            "utilization": "使用率",
            "result": "結果",
        }
    )
    _add_docx_table(document, stress_check_df, [3.0, 4.1, 3.1, 2.0, 2.0])

    document.add_heading("7. 各支承反力", level=1)
    reaction_source = pd.DataFrame(result["support_reactions"])
    reaction_location_df = reaction_source[
        ["name", "type_text", "node_tag", "y"]
    ].rename(
        columns={
            "name": "名稱",
            "type_text": "類型",
            "node_tag": "節點",
            "y": "Y (cm)",
        }
    )
    _add_docx_table(document, reaction_location_df, [2.8, 4.0, 2.8, 4.6])

    reaction_force_df = reaction_source[["name", "Rx", "Ry", "Mz"]].rename(
        columns={
            "name": "名稱",
            "Rx": "RX (kgf)",
            "Ry": "RY (kgf)",
            "Mz": "MZ (kgf-cm)",
        }
    )
    _add_docx_table(document, reaction_force_df, [2.8, 3.8, 3.8, 3.8])

    document.add_page_break()
    figures = result["figures"]
    assert isinstance(figures, dict)

    heading = document.add_heading("8. 圖表", level=1)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)

    figure_items = [
        ("模型與均布載重", "model"),
        ("變形圖", "deformation"),
        ("彎矩圖", "moment"),
        ("剪力圖", "shear"),
    ]

    for figure_index, (figure_title, figure_key) in enumerate(figure_items):
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_before = Pt(1)
        title_paragraph.paragraph_format.space_after = Pt(1)
        title_paragraph.paragraph_format.keep_with_next = True
        title_run = title_paragraph.add_run(figure_title)
        title_run.bold = True
        title_run.font.name = "Microsoft JhengHei"
        title_run.font.size = Pt(11.5)
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

        figure_stream = BytesIO(figure_to_png_bytes(figures[figure_key]))
        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_before = Pt(0)
        picture_paragraph.paragraph_format.space_after = Pt(2)
        picture_paragraph.paragraph_format.line_spacing = 1.0
        picture_paragraph.add_run().add_picture(
            figure_stream,
            height=Cm(WORD_FIGURE_HEIGHT_CM),
        )

        # 每頁放兩張獨立圖片；圖片本身不合併、不放入表格。
        if figure_index == 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


# ============================================================
# 3. 分析主函式
# ============================================================


def run_analysis(
    w: float,
    load_metadata: Dict[str, object],
    stations: Sequence[Dict[str, object]],
    input_segments: Sequence[Dict[str, object]],
    material_name: str,
    grade_name: str,
    profiles: Sequence[Dict[str, object]],
) -> Dict[str, object]:
    if ops is None or opsv is None:
        raise RuntimeError("必要的分析套件尚未安裝。")
    if not np.isfinite(w) or w <= 0.0:
        raise ValueError("均布風載 w 必須大於 0。")
    if not profiles:
        raise ValueError("至少需要一組斷面資料。")

    for profile in profiles:
        for field in ("A", "Ix", "Sx"):
            value = float(profile[field])
            if not np.isfinite(value) or value <= 0.0:
                raise ValueError(f"{profile['name']} 的 {field} 必須大於 0。")

    stations = [dict(item) for item in stations]
    input_segments = [dict(item) for item in input_segments]
    profiles = [dict(item) for item in profiles]

    number_of_input_segments = len(input_segments)
    segment_lengths = [float(item["length"]) for item in input_segments]
    total_height = float(stations[-1]["y"])

    all_support_specs = [
        dict(station) for station in stations if str(station["type"]) == "hinge"
    ]
    pin_specs = [
        dict(station) for station in stations if str(station["type"]) == "pin"
    ]
    free_end_specs = [
        dict(station) for station in stations if str(station["type"]) == "free"
    ]

    first_support_y = min(float(item["y"]) for item in all_support_specs)
    last_support_y = max(float(item["y"]) for item in all_support_specs)
    bottom_cantilever_length = first_support_y
    top_cantilever_length = total_height - last_support_y
    supported_start_y = first_support_y
    supported_end_y = last_support_y
    supported_height = last_support_y - first_support_y

    if material_name not in MATERIAL_DATABASE:
        raise ValueError("材料種類不存在於材料資料庫中。")
    material_data = MATERIAL_DATABASE[material_name]
    grades = material_data["grades"]
    assert isinstance(grades, dict)
    if grade_name not in grades:
        raise ValueError("材料牌號不存在於材料資料庫中。")
    grade = grades[grade_name]
    assert isinstance(grade, dict)
    elastic_modulus = float(material_data["E"])

    area_mullion = sum(float(item["A"]) for item in profiles)
    i_mullion = sum(float(item["Ix"]) for item in profiles)
    points_per_element = POSTPROCESS_POINTS_PER_ELEMENT

    support_specs_by_y: Dict[int, List[Dict[str, object]]] = {}
    for support in all_support_specs:
        support_specs_by_y.setdefault(y_key(float(support["y"])), []).append(support)

    pin_specs_by_y: Dict[int, Dict[str, object]] = {
        y_key(float(pin["y"])): pin for pin in pin_specs
    }

    check_segments = build_check_segments(all_support_specs, total_height)
    event_y_coordinates = [float(item["y"]) for item in stations]
    mesh_y_coordinates = build_mesh_coordinates(input_segments, event_y_coordinates)

    if not same_y(mesh_y_coordinates[0], 0.0):
        raise RuntimeError("網格最下端不是 Y=0。")
    if not same_y(mesh_y_coordinates[-1], total_height):
        raise RuntimeError("網格最上端不等於直料總高度。")

    # --------------------------------------------------------
    # 建立模型
    # --------------------------------------------------------
    ops.wipe()
    ops.model("basic", "-ndm", 2, "-ndf", 3)
    ops.geomTransf("Linear", 1)

    node_tags: List[int] = []
    element_tags: List[int] = []
    element_data: Dict[int, Dict[str, object]] = {}
    pin_model_data: List[Dict[str, object]] = []
    support_model_data: List[Dict[str, object]] = []

    next_node_tag = 1
    next_element_tag = 1

    ops.node(next_node_tag, 0.0, 0.0)
    node_tags.append(next_node_tag)
    current_active_node = next_node_tag
    next_node_tag += 1

    for support in support_specs_by_y.get(y_key(0.0), []):
        ops.fix(current_active_node, 1, 1, 0)
        support["node_tag"] = current_active_node
        support_model_data.append(support)

    lower_free_tip_node_tag = (
        current_active_node if str(stations[0]["type"]) == "free" else None
    )
    upper_free_tip_node_tag: Optional[int] = None
    last_created_element: Optional[int] = None

    for station_index in range(1, len(mesh_y_coordinates)):
        y_i = float(mesh_y_coordinates[station_index - 1])
        y_j = float(mesh_y_coordinates[station_index])
        length = y_j - y_i

        if length <= COORDINATE_TOLERANCE:
            raise RuntimeError(
                f"發現零長度或負長度梁區段：Y={y_i} 至 Y={y_j}。"
            )

        arrival_node = next_node_tag
        ops.node(arrival_node, 0.0, y_j)
        node_tags.append(arrival_node)
        next_node_tag += 1

        y_mid = 0.5 * (y_i + y_j)
        check_key = find_check_segment_key(y_mid, check_segments)

        ops.element(
            "elasticBeamColumn",
            next_element_tag,
            current_active_node,
            arrival_node,
            area_mullion,
            elastic_modulus,
            i_mullion,
            1,
        )

        element_tags.append(next_element_tag)
        element_data[next_element_tag] = {
            "node_i": current_active_node,
            "node_j": arrival_node,
            "y_i": y_i,
            "y_j": y_j,
            "length": length,
            "check_segment_key": check_key,
        }

        last_created_element = next_element_tag
        next_element_tag += 1
        station_key = y_key(y_j)

        for support in support_specs_by_y.get(station_key, []):
            ops.fix(arrival_node, 1, 1, 0)
            support["node_tag"] = arrival_node
            support_model_data.append(support)

        pin_spec = pin_specs_by_y.get(station_key)
        if pin_spec is not None:
            upper_node = next_node_tag
            ops.node(upper_node, 0.0, y_j)
            node_tags.append(upper_node)
            next_node_tag += 1

            ops.equalDOF(arrival_node, upper_node, 1, 2)

            pin_record = dict(pin_spec)
            pin_record.update(
                {
                    "lower_node": arrival_node,
                    "upper_node": upper_node,
                    "lower_element": last_created_element,
                    "upper_element": None,
                }
            )
            pin_model_data.append(pin_record)
            current_active_node = upper_node
        else:
            current_active_node = arrival_node

    for pin_record in pin_model_data:
        pin_y = float(pin_record["y"])
        upper_node = int(pin_record["upper_node"])
        candidates = [
            element_tag
            for element_tag, data in element_data.items()
            if int(data["node_i"]) == upper_node
            and same_y(float(data["y_i"]), pin_y)
        ]
        if len(candidates) != 1:
            raise RuntimeError(
                f"Pin {pin_record['name']} 無法唯一找到上方第一支梁元素。"
            )
        pin_record["upper_element"] = candidates[0]

    if str(stations[-1]["type"]) == "free":
        upper_free_tip_node_tag = current_active_node

    if len(support_model_data) != len(all_support_specs):
        raise RuntimeError("部分支承未成功建立節點。")

    # --------------------------------------------------------
    # 載重與分析
    # --------------------------------------------------------
    ops.timeSeries("Linear", 1)
    ops.pattern("Plain", 1, 1)
    q_local_y = -w

    for element_tag in element_tags:
        ops.eleLoad(
            "-ele",
            element_tag,
            "-type",
            "-beamUniform",
            q_local_y,
            0.0,
        )

    ops.constraints("Transformation")
    ops.numberer("RCM")
    ops.system("BandGeneral")
    ops.algorithm("Linear")
    ops.integrator("LoadControl", 1.0)
    ops.analysis("Static")

    analysis_result = ops.analyze(1)
    if analysis_result != 0:
        raise RuntimeError(f"靜力分析失敗，分析代碼={analysis_result}。")

    # --------------------------------------------------------
    # 位移、彎矩與剪力後處理
    # --------------------------------------------------------
    ux = {tag: float(ops.nodeDisp(tag, 1)) for tag in node_tags}
    rz = {tag: float(ops.nodeDisp(tag, 3)) for tag in node_tags}

    all_y: List[float] = []
    all_ux: List[float] = []
    all_moment: List[float] = []
    all_shear: List[float] = []

    segment_response: Dict[Tuple[str, int], Dict[str, List[float]]] = {
        (str(item["type"]), int(item["index"])): {
            "y": [],
            "ux": [],
            "moment": [],
            "shear": [],
        }
        for item in check_segments
    }

    for element_sequence, element_tag in enumerate(element_tags):
        data = element_data[element_tag]
        node_i = int(data["node_i"])
        node_j = int(data["node_j"])
        y_i = float(data["y_i"])
        y_j = float(data["y_j"])
        length = float(data["length"])

        local_v_i = -ux[node_i]
        local_v_j = -ux[node_j]
        theta_i = rz[node_i]
        theta_j = rz[node_j]
        xi = np.linspace(0.0, 1.0, points_per_element)

        n1 = 1.0 - 3.0 * xi**2 + 2.0 * xi**3
        n2 = xi - 2.0 * xi**2 + xi**3
        n3 = 3.0 * xi**2 - 2.0 * xi**3
        n4 = -xi**2 + xi**3

        load_bubble = (
            q_local_y
            * length**4
            / (24.0 * elastic_modulus * i_mullion)
            * xi**2
            * (1.0 - xi) ** 2
        )

        local_v = (
            n1 * local_v_i
            + n2 * length * theta_i
            + n3 * local_v_j
            + n4 * length * theta_j
            + load_bubble
        )
        global_ux = -local_v
        y_values = y_i + xi * length

        curvature = (
            (-6.0 + 12.0 * xi) / length**2 * local_v_i
            + (-4.0 + 6.0 * xi) / length * theta_i
            + (6.0 - 12.0 * xi) / length**2 * local_v_j
            + (-2.0 + 6.0 * xi) / length * theta_j
            + q_local_y
            * length**2
            / (24.0 * elastic_modulus * i_mullion)
            * (2.0 - 12.0 * xi + 12.0 * xi**2)
        )
        moment_values = elastic_modulus * i_mullion * curvature

        third_derivative = (
            12.0 / length**3 * local_v_i
            + 6.0 / length**2 * theta_i
            - 12.0 / length**3 * local_v_j
            + 6.0 / length**2 * theta_j
            + q_local_y
            * length
            / (24.0 * elastic_modulus * i_mullion)
            * (-12.0 + 24.0 * xi)
        )
        shear_values = elastic_modulus * i_mullion * third_derivative

        keep_start = element_sequence == 0
        if not keep_start:
            previous_element = element_data[element_tags[element_sequence - 1]]
            previous_node_j = int(previous_element["node_j"])
            keep_start = previous_node_j != node_i

        start_index = 0 if keep_start else 1
        all_y.extend(y_values[start_index:].tolist())
        all_ux.extend(global_ux[start_index:].tolist())
        all_moment.extend(moment_values[start_index:].tolist())
        all_shear.extend(shear_values[start_index:].tolist())

        key = tuple(data["check_segment_key"])
        response = segment_response[key]  # type: ignore[index]
        segment_start_index = 0 if len(response["y"]) == 0 or keep_start else 1
        response["y"].extend(y_values[segment_start_index:].tolist())
        response["ux"].extend(global_ux[segment_start_index:].tolist())
        response["moment"].extend(moment_values[segment_start_index:].tolist())
        response["shear"].extend(shear_values[segment_start_index:].tolist())

    all_y_array = np.asarray(all_y, dtype=float)
    all_ux_array = np.asarray(all_ux, dtype=float)
    all_moment_array = np.asarray(all_moment, dtype=float)
    all_shear_array = np.asarray(all_shear, dtype=float)

    max_displacement_index = int(np.argmax(np.abs(all_ux_array)))
    max_displacement = float(all_ux_array[max_displacement_index])
    max_displacement_y = float(all_y_array[max_displacement_index])
    max_abs_displacement = abs(max_displacement)

    max_moment_index = int(np.argmax(np.abs(all_moment_array)))
    max_moment_signed = float(all_moment_array[max_moment_index])
    max_moment_y = float(all_y_array[max_moment_index])
    max_abs_moment = abs(max_moment_signed)

    max_shear_index = int(np.argmax(np.abs(all_shear_array)))
    max_shear_signed = float(all_shear_array[max_shear_index])
    max_shear_y = float(all_y_array[max_shear_index])
    max_abs_shear = abs(max_shear_signed)

    # --------------------------------------------------------
    # 變形檢核
    # --------------------------------------------------------
    span_check_results: List[Dict[str, object]] = []
    for segment in check_segments:
        key = (str(segment["type"]), int(segment["index"]))
        response = segment_response[key]
        y_values = np.asarray(response["y"], dtype=float)
        ux_values = np.asarray(response["ux"], dtype=float)

        if len(y_values) == 0:
            raise RuntimeError(f"檢核區段 {segment['name']} 沒有後處理資料。")

        control_index = int(np.argmax(np.abs(ux_values)))
        control_disp_signed = float(ux_values[control_index])
        control_disp = abs(control_disp_signed)
        control_y = float(y_values[control_index])

        if segment["type"] == "support_span":
            allowable, formula = span_allowable_deflection(float(segment["length"]))
        else:
            allowable = 2.0 * float(segment["length"]) / 175.0
            formula = "2Lc/175"

        utilization = control_disp / allowable
        passed = control_disp <= allowable

        span_check_results.append(
            {
                **segment,
                "control_disp_signed": control_disp_signed,
                "control_disp": control_disp,
                "control_y": control_y,
                "allowable": allowable,
                "formula": formula,
                "utilization": utilization,
                "passed": passed,
                "result": status_text(passed),
            }
        )

    deflection_pass = all(bool(item["passed"]) for item in span_check_results)

    # --------------------------------------------------------
    # 變形控制之建議總 Ix
    # --------------------------------------------------------
    # 線彈性範圍內，在載重、E、幾何與支承條件固定時：
    #     δ ∝ 1 / I
    # 因此可由目前最大變形使用率反算滿足變形限制所需之總 Ix。
    governing_deflection_utilization = max(
        float(item["utilization"]) for item in span_check_results
    )
    recommended_ix_deflection = (
        i_mullion * governing_deflection_utilization
    )

    governing_deflection_segment = max(
        span_check_results,
        key=lambda item: float(item["utilization"]),
    )

    # --------------------------------------------------------
    # 應力檢核
    # --------------------------------------------------------
    stress_results: List[Dict[str, object]] = []
    for profile in profiles:
        moment_share = max_abs_moment * float(profile["Ix"]) / i_mullion
        fb = moment_share / float(profile["Sx"])
        utilization = fb / float(grade["Fb_allow"])
        passed = utilization <= 1.0
        stress_results.append(
            {
                **profile,
                "moment_share": moment_share,
                "fb": fb,
                "utilization": utilization,
                "passed": passed,
                "result": status_text(passed),
            }
        )

    stress_control = max(
        stress_results,
        key=lambda item: float(item["utilization"]),
    )
    stress_pass = all(bool(item["passed"]) for item in stress_results)
    overall_pass = stress_pass and deflection_pass

    # --------------------------------------------------------
    # 反力與 Pin 檢查
    # --------------------------------------------------------
    ops.reactions()

    support_reactions: List[Dict[str, object]] = []
    for support in sorted(support_model_data, key=lambda item: float(item["y"])):
        node_tag = int(support["node_tag"])
        support_reactions.append(
            {
                "name": str(support["name"]),
                "type_text": str(support["type_text"]),
                "node_tag": node_tag,
                "y": float(support["y"]),
                "Rx": float(ops.nodeReaction(node_tag, 1)),
                "Ry": float(ops.nodeReaction(node_tag, 2)),
                "Mz": float(ops.nodeReaction(node_tag, 3)),
            }
        )

    external_force_x = w * total_height

    lower_tip_ux = (
        ux[lower_free_tip_node_tag]
        if lower_free_tip_node_tag is not None
        else None
    )
    upper_tip_ux = (
        ux[upper_free_tip_node_tag]
        if upper_free_tip_node_tag is not None
        else None
    )

    # --------------------------------------------------------
    # 圖形
    # --------------------------------------------------------
    target_force_width = max(total_height * 0.20, 1.0)
    target_deformation_width = max(total_height * 0.08, 1.0)

    sfac_defo = (
        target_deformation_width / max_abs_displacement
        if max_abs_displacement > 1.0e-12
        else 1.0
    )
    sfac_moment = (
        target_force_width / max_abs_moment if max_abs_moment > 1.0e-12 else 1.0
    )
    sfac_shear = (
        target_force_width / max_abs_shear if max_abs_shear > 1.0e-12 else 1.0
    )

    fmt_force_main = {
        "color": "blue",
        "linestyle": "solid",
        "linewidth": 1.8,
        "marker": "",
        "markersize": 0,
    }
    fmt_force_reference = {
        "color": "blue",
        "linestyle": "solid",
        "linewidth": 0.6,
        "marker": "",
        "markersize": 0,
    }

    fig1, ax1 = plt.subplots(figsize=(5.2, 7.4))
    fig2, ax2 = plt.subplots(figsize=(5.2, 7.4))
    fig3, ax3 = plt.subplots(figsize=(7, 10))
    fig4, ax4 = plt.subplots(figsize=(7, 10))

    load_width = max(total_height * 0.08, 1.0)

    ax1.set_title("Vertical Multi-Segment Beam Model")
    opsv.plot_model(
        ax=ax1,
        node_supports=True,
        node_labels=0,
        element_labels=0,
    )

    for support in sorted(all_support_specs, key=lambda item: float(item["y"])):
        y_value = float(support["y"])
        ax1.annotate(
            str(support["name"]),
            xy=(0.0, y_value),
            xytext=(8, 0),
            textcoords="offset points",
            ha="left",
            va="center",
            fontsize=9,
            bbox={
                "boxstyle": "round,pad=0.18",
                "facecolor": "white",
                "alpha": 0.82,
            },
        )

    for pin in sorted(pin_model_data, key=lambda item: float(item["y"])):
        pin_y = float(pin["y"])
        ax1.scatter(
            0.0,
            pin_y,
            s=38,
            facecolors="white",
            edgecolors="black",
            zorder=7,
        )
        ax1.annotate(
            str(pin["name"]),
            xy=(0.0, pin_y),
            xytext=(-10, 0),
            textcoords="offset points",
            ha="right",
            va="center",
            fontsize=9,
            bbox={
                "boxstyle": "round,pad=0.15",
                "facecolor": "white",
                "alpha": 0.82,
            },
        )

    if lower_free_tip_node_tag is not None:
        ax1.annotate(
            "Lower free tip",
            xy=(0.0, 0.0),
            xytext=(8, -8),
            textcoords="offset points",
            ha="left",
            va="top",
            fontsize=9,
        )
    if upper_free_tip_node_tag is not None:
        ax1.annotate(
            "Upper free tip",
            xy=(0.0, total_height),
            xytext=(8, 8),
            textcoords="offset points",
            ha="left",
            va="bottom",
            fontsize=9,
        )

    rect = patches.Rectangle(
        (-load_width, 0.0),
        load_width,
        total_height,
        linewidth=1,
        edgecolor="steelblue",
        facecolor="dodgerblue",
        alpha=0.20,
    )
    ax1.add_patch(rect)

    for y_position in np.linspace(0.0, total_height, 18):
        ax1.annotate(
            "",
            xy=(0.0, y_position),
            xytext=(-load_width, y_position),
            arrowprops={
                "arrowstyle": "->",
                "color": "steelblue",
                "lw": 1.4,
            },
        )

    if bottom_cantilever_length > COORDINATE_TOLERANCE:
        ax1.axhline(supported_start_y, linestyle="--", linewidth=1.0, color="gray")
    if top_cantilever_length > COORDINATE_TOLERANCE:
        ax1.axhline(supported_end_y, linestyle="--", linewidth=1.0, color="gray")

    ax1.set_xlabel("Global X Coordinate (cm)")
    ax1.set_ylabel("Global Y Coordinate (cm)")
    ax1.set_aspect("equal", adjustable="box")
    ax1.grid(True, linestyle="--", alpha=0.3)
    ax1.set_xlim(-1.6 * load_width, 1.8 * load_width)
    ax1.text(
        0.03,
        0.985,
        f"Uniform load w = {w:.3f} kgf/cm",
        transform=ax1.transAxes,
        ha="left",
        va="top",
        fontsize=10,
        bbox={
            "boxstyle": "round,pad=0.25",
            "facecolor": "white",
            "alpha": 0.90,
        },
    )

    ax2.set_title(
        "Deformed Shape\n"
        f"Maximum |UX| = {max_abs_displacement:.3f} cm"
    )
    opsv.plot_defo(
        sfac=sfac_defo,
        unDefoFlag=1,
        node_supports=True,
        ax=ax2,
    )
    ax2.plot(
        max_displacement * sfac_defo,
        max_displacement_y,
        marker="o",
        markersize=6,
        zorder=6,
    )
    ax2.annotate(
        f"Max |UX|\nUX={max_displacement:.3f} cm\nY={max_displacement_y:.3f} cm",
        xy=(max_displacement * sfac_defo, max_displacement_y),
        xytext=(10, 8),
        textcoords="offset points",
        arrowprops={"arrowstyle": "->", "linewidth": 0.8},
        fontsize=9,
        bbox={
            "boxstyle": "round,pad=0.20",
            "facecolor": "white",
            "alpha": 0.85,
        },
    )
    if lower_tip_ux is not None:
        lower_tip_x = float(lower_tip_ux) * sfac_defo
        ax2.plot(lower_tip_x, 0.0, marker="s", markersize=6, zorder=7)
        ax2.annotate(
            f"Lower cantilever tip\nUX={float(lower_tip_ux):.3f} cm",
            xy=(lower_tip_x, 0.0),
            xytext=(10, 18),
            textcoords="offset points",
            arrowprops={"arrowstyle": "->", "linewidth": 0.8},
            fontsize=9,
            bbox={"boxstyle": "round,pad=0.20", "facecolor": "white", "alpha": 0.85},
        )
    if upper_tip_ux is not None:
        upper_tip_x = float(upper_tip_ux) * sfac_defo
        ax2.plot(upper_tip_x, total_height, marker="s", markersize=6, zorder=7)
        ax2.annotate(
            f"Upper cantilever tip\nUX={float(upper_tip_ux):.3f} cm",
            xy=(upper_tip_x, total_height),
            xytext=(10, -32),
            textcoords="offset points",
            arrowprops={"arrowstyle": "->", "linewidth": 0.8},
            fontsize=9,
            bbox={"boxstyle": "round,pad=0.20", "facecolor": "white", "alpha": 0.85},
        )

    ax2.xaxis.set_major_formatter(
        ticker.FuncFormatter(lambda x, position: f"{x / sfac_defo:.3f}")
    )
    ax2.set_xlabel("True Horizontal Displacement UX (cm)")
    ax2.set_ylabel("Global Y Coordinate (cm)")
    ax2.set_aspect("auto")
    ax2.grid(True, linestyle="--", alpha=0.3)

    ax3.set_title(
        "Bending Moment Diagram\n"
        f"Maximum |M| = {max_abs_moment:.3f} kgf-cm"
    )
    opsv.section_force_diagram_2d(
        "M",
        sfac=sfac_moment,
        nep=17,
        ax=ax3,
        fmt_secforce1=fmt_force_main,
        fmt_secforce2=fmt_force_reference,
        ref_vert_lines=False,
        end_max_values=False,
        node_supports=True,
        number_format=".3f",
    )
    ax3.axvline(0.0, linewidth=1.0)
    ax3.axhline(max_moment_y, linewidth=0.8, linestyle=":", alpha=0.65)
    ax3.plot(0.0, max_moment_y, marker="o", markersize=6, zorder=6)
    ax3.annotate(
        f"Max |M|\nM={max_moment_signed:.3f} kgf-cm\nY={max_moment_y:.3f} cm",
        xy=(0.0, max_moment_y),
        xytext=(10, 8),
        textcoords="offset points",
        arrowprops={"arrowstyle": "->", "linewidth": 0.8},
        fontsize=9,
        bbox={
            "boxstyle": "round,pad=0.20",
            "facecolor": "white",
            "alpha": 0.85,
        },
    )
    ax3.xaxis.set_major_formatter(
        ticker.FuncFormatter(lambda x, position: f"{x / sfac_moment:.3f}")
    )
    ax3.set_xlabel("True Bending Moment M (kgf-cm)")
    ax3.set_ylabel("Global Y Coordinate (cm)")
    ax3.set_aspect("auto")
    ax3.grid(True, linestyle="--", alpha=0.3)

    ax4.set_title(
        "Shear Force Diagram\n"
        f"Maximum |V| = {max_abs_shear:.3f} kgf"
    )
    opsv.section_force_diagram_2d(
        "V",
        sfac=sfac_shear,
        nep=17,
        ax=ax4,
        fmt_secforce1=fmt_force_main,
        fmt_secforce2=fmt_force_reference,
        ref_vert_lines=False,
        end_max_values=False,
        node_supports=True,
        number_format=".3f",
    )
    ax4.axvline(0.0, linewidth=1.0)
    ax4.axhline(max_shear_y, linewidth=0.8, linestyle=":", alpha=0.65)
    ax4.plot(0.0, max_shear_y, marker="o", markersize=6, zorder=6)
    ax4.annotate(
        f"Max |V|\nV={max_shear_signed:.3f} kgf\nY={max_shear_y:.3f} cm",
        xy=(0.0, max_shear_y),
        xytext=(10, 8),
        textcoords="offset points",
        arrowprops={"arrowstyle": "->", "linewidth": 0.8},
        fontsize=9,
        bbox={
            "boxstyle": "round,pad=0.20",
            "facecolor": "white",
            "alpha": 0.85,
        },
    )
    ax4.xaxis.set_major_formatter(
        ticker.FuncFormatter(lambda x, position: f"{x / sfac_shear:.3f}")
    )
    ax4.set_xlabel("True Shear Force V (kgf)")
    ax4.set_ylabel("Global Y Coordinate (cm)")
    ax4.set_aspect("auto")
    ax4.grid(True, linestyle="--", alpha=0.3)

    for figure in (fig1, fig2, fig3, fig4):
        figure.tight_layout()

    model_summary = {
        "使用者輸入段數": number_of_input_segments,
        "Hinge 鉸支承數": len(all_support_specs),
        "內部 Pin 數": len(pin_specs),
        "Free end 數": len(free_end_specs),
        "實際支承間檢核區段數": len(check_segments),
        "最下與最上支承間高度 (cm)": supported_height,
        "下部懸臂長度 (cm)": bottom_cantilever_length,
        "頂部懸臂長度 (cm)": top_cantilever_length,
        "直料總高度 (cm)": total_height,
        "分析節點數": len(node_tags),
        "分析元素數": len(element_tags),
        "材料種類": material_name,
        "材料牌號": grade_name,
        "彈性模數 E (kgf/cm²)": elastic_modulus,
        "容許強軸撓曲應力 Fb (kgf/cm²)": float(grade["Fb_allow"]),
        "總 A (cm²)": area_mullion,
        "總 Ix (cm⁴)": i_mullion,
        **load_metadata,
        "等效風力線載重 w (kgf/cm)": w,
        "風載水平合力 (kgf)": external_force_x,
    }

    maximum_results = {
        "max_displacement": max_displacement,
        "max_abs_displacement": max_abs_displacement,
        "max_displacement_y": max_displacement_y,
        "max_moment_signed": max_moment_signed,
        "max_abs_moment": max_abs_moment,
        "max_moment_y": max_moment_y,
        "max_shear_signed": max_shear_signed,
        "max_abs_shear": max_abs_shear,
        "max_shear_y": max_shear_y,
        "lower_tip_ux": lower_tip_ux,
        "upper_tip_ux": upper_tip_ux,
    }


    return {
        "stations": stations,
        "input_segments": input_segments,
        "segment_lengths": segment_lengths,
        "model_summary": model_summary,
        "maximum_results": maximum_results,
        "span_check_results": span_check_results,
        "stress_results": stress_results,
        "stress_control": stress_control,
        "support_reactions": support_reactions,
        "deflection_pass": deflection_pass,
        "stress_pass": stress_pass,
        "overall_pass": overall_pass,
        "recommended_ix_deflection": recommended_ix_deflection,
        "governing_deflection_utilization": governing_deflection_utilization,
        "governing_deflection_segment": str(governing_deflection_segment["name"]),
        "figures": {
            "model": fig1,
            "deformation": fig2,
            "moment": fig3,
            "shear": fig4,
        },
    }


# ============================================================
# 4. Streamlit 介面
# ============================================================


def clear_analysis_result() -> None:
    old_result = st.session_state.pop("mullion_analysis_result", None)
    if isinstance(old_result, dict):
        figures = old_result.get("figures", {})
        if isinstance(figures, dict):
            for figure in figures.values():
                try:
                    plt.close(figure)
                except Exception:
                    pass


def station_format(station_type: str) -> str:
    return STATION_LABELS[station_type]


def render_dataframe_download(
    dataframe: pd.DataFrame,
    label: str,
    filename: str,
    key: str,
) -> None:
    st.download_button(
        label=label,
        data=dataframe_to_csv_bytes(dataframe),
        file_name=filename,
        mime="text/csv",
        key=key,
    )


def render_results(result: Dict[str, object], report_name_input: str) -> None:
    overall_pass = bool(result["overall_pass"])
    stress_pass = bool(result["stress_pass"])
    deflection_pass = bool(result["deflection_pass"])
    maximum = result["maximum_results"]
    assert isinstance(maximum, dict)

    st.divider()
    if overall_pass:
        st.success("整體判定：PASS（應力與變形檢核皆通過）")
    else:
        st.error("整體判定：FAIL（至少一項應力或變形檢核未通過）")

    status_col1, status_col2, status_col3 = st.columns(3)
    status_col1.metric("強軸應力檢核", status_text(stress_pass))
    status_col2.metric("容許變形檢核", status_text(deflection_pass))
    status_col3.metric("整體判定", status_text(overall_pass))

    metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)
    metric_col1.metric(
        "最大 |UX|",
        f"{float(maximum['max_abs_displacement']):.3f} cm",
        f"Y={float(maximum['max_displacement_y']):.3f} cm",
    )
    metric_col2.metric(
        "最大 |M|",
        f"{float(maximum['max_abs_moment']):.3f} kgf-cm",
        f"Y={float(maximum['max_moment_y']):.3f} cm",
    )
    metric_col3.metric(
        "最大 |V|",
        f"{float(maximum['max_abs_shear']):.3f} kgf",
        f"Y={float(maximum['max_shear_y']):.3f} cm",
    )
    summary = result["model_summary"]
    assert isinstance(summary, dict)
    metric_col4.metric(
        "直料總高度",
        f"{float(summary['直料總高度 (cm)']):.3f} cm",
        None,
    )

    current_ix = float(summary["總 Ix (cm⁴)"])
    recommended_ix = float(result["recommended_ix_deflection"])
    governing_segment = str(result["governing_deflection_segment"])
    governing_utilization = float(result["governing_deflection_utilization"])

    st.info(
        f"依實際模型反算 Ix（變形控制）≈ {recommended_ix:.3f} cm⁴ ｜ "
        f"目前總 Ix = {current_ix:.3f} cm⁴ ｜ "
        f"控制區段：{governing_segment} ｜ "
        f"變形使用率 = {governing_utilization:.3f}"
    )

    report_filename = sanitize_report_filename(report_name_input)
    if DOCX_IMPORT_ERROR is None:
        try:
            report_bytes = build_word_report(result)
        except Exception as exc:
            st.warning(f"Word 報告建立失敗：{exc}")
        else:
            st.download_button(
                "匯出所有數據與圖表的 Word 報告",
                data=report_bytes,
                file_name=report_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
                key="download_full_word_report",
            )
    else:
        st.warning("未安裝 python-docx，因此目前無法匯出 Word 報告。")

    tab_model, tab_defo, tab_force = st.tabs(
        ["模型與幾何", "變形檢核", "應力、內力與反力"]
    )

    with tab_model:
        left, right = st.columns([1.05, 1.0])
        with left:
            st.subheader("模型圖")
            figures = result["figures"]
            assert isinstance(figures, dict)
            model_png = figure_to_png_bytes(figures["model"])
            model_space_left, model_image_col, model_space_right = st.columns(
                [0.15, 1.0, 0.15]
            )
            with model_image_col:
                st.image(
                    model_png,
                    width=MODEL_DISPLAY_WIDTH_PX,
                )
            st.download_button(
                "下載模型圖 PNG",
                data=model_png,
                file_name="vertical_beam_model.png",
                mime="image/png",
                key="download_model_png",
            )
        with right:
            st.subheader("模型摘要")
            summary_df = pd.DataFrame(
                [{"項目": key, "數值": value} for key, value in summary.items()]
            )
            st.dataframe(round_dataframe_for_display(summary_df), use_container_width=True, hide_index=True)

        st.subheader("控制點配置")
        stations_df = pd.DataFrame(result["stations"])[
            ["name", "type_text", "y", "station_number"]
        ].rename(
            columns={
                "name": "名稱",
                "type_text": "類型",
                "y": "Y (cm)",
                "station_number": "控制點編號",
            }
        )
        st.dataframe(round_dataframe_for_display(stations_df), use_container_width=True, hide_index=True)

        st.subheader("各輸入段")
        segments_df = pd.DataFrame(result["input_segments"])[
            [
                "segment_number",
                "name",
                "connection_type",
                "length",
                "start_y",
                "end_y",
            ]
        ].rename(
            columns={
                "segment_number": "段號",
                "name": "控制點",
                "connection_type": "連接類型",
                "length": "長度 (cm)",
                "start_y": "起點 Y (cm)",
                "end_y": "終點 Y (cm)",
            }
        )
        st.dataframe(round_dataframe_for_display(segments_df), use_container_width=True, hide_index=True)
        render_dataframe_download(
            segments_df,
            "下載幾何配置 CSV",
            "geometry_segments.csv",
            "download_geometry_csv",
        )

    with tab_defo:
        figures = result["figures"]
        assert isinstance(figures, dict)

        st.subheader("各支承間距與懸臂變形檢核")

        span_df = pd.DataFrame(result["span_check_results"])[
            [
                "name",
                "length",
                "control_y",
                "control_disp",
                "formula",
                "allowable",
                "utilization",
                "result",
            ]
        ].rename(
            columns={
                "name": "區段",
                "length": "長度 (cm)",
                "control_y": "控制位置 Y (cm)",
                "control_disp": "實際 |UX| (cm)",
                "formula": "容許公式",
                "allowable": "容許值 (cm)",
                "utilization": "使用率",
                "result": "結果",
            }
        )
        st.dataframe(
            span_df.style.format(
                {
                    "長度 (cm)": "{:.3f}",
                    "控制位置 Y (cm)": "{:.3f}",
                    "實際 |UX| (cm)": "{:.3f}",
                    "容許值 (cm)": "{:.3f}",
                    "使用率": "{:.3f}",
                }
            ),
            use_container_width=True,
            hide_index=True,
        )
        render_dataframe_download(
            span_df,
            "下載變形檢核 CSV",
            "deflection_checks.csv",
            "download_defo_csv",
        )

        lower_tip_ux = maximum.get("lower_tip_ux")
        upper_tip_ux = maximum.get("upper_tip_ux")
        tip_cols = st.columns(2)
        if lower_tip_ux is not None:
            tip_cols[0].metric("下部懸臂端 UX", f"{float(lower_tip_ux):.3f} cm")
        if upper_tip_ux is not None:
            tip_cols[1].metric("頂部懸臂端 UX", f"{float(upper_tip_ux):.3f} cm")

        st.subheader("變形圖")
        deformation_png = figure_to_png_bytes(figures["deformation"])
        defo_space_left, defo_image_col, defo_space_right = st.columns(
            [1.0, 1.35, 1.0]
        )
        with defo_image_col:
            st.image(
                deformation_png,
                width=DEFORMATION_DISPLAY_WIDTH_PX,
            )
        st.download_button(
            "下載變形圖 PNG",
            data=deformation_png,
            file_name="deformed_shape.png",
            mime="image/png",
            key="download_defo_png",
        )

    with tab_force:
        st.subheader("強軸應力檢核")
        stress_df = pd.DataFrame(result["stress_results"])[
            [
                "name",
                "A",
                "Ix",
                "Sx",
                "moment_share",
                "fb",
                "utilization",
                "result",
            ]
        ].rename(
            columns={
                "name": "斷面",
                "A": "A (cm²)",
                "Ix": "Ix (cm⁴)",
                "Sx": "Sx (cm³)",
                "moment_share": "分配彎矩 (kgf-cm)",
                "fb": "fb (kgf/cm²)",
                "utilization": "使用率",
                "result": "結果",
            }
        )
        st.dataframe(
            stress_df.style.format(
                {
                    "A (cm²)": "{:.3f}",
                    "Ix (cm⁴)": "{:.3f}",
                    "Sx (cm³)": "{:.3f}",
                    "分配彎矩 (kgf-cm)": "{:.3f}",
                    "fb (kgf/cm²)": "{:.3f}",
                    "使用率": "{:.3f}",
                }
            ),
            use_container_width=True,
            hide_index=True,
        )

        st.subheader("內力圖")
        moment_col, shear_col = st.columns(2)
        with moment_col:
            st.pyplot(figures["moment"], use_container_width=True)
            st.download_button(
                "下載彎矩圖 PNG",
                data=figure_to_png_bytes(figures["moment"]),
                file_name="bending_moment.png",
                mime="image/png",
                key="download_moment_png",
            )
        with shear_col:
            st.pyplot(figures["shear"], use_container_width=True)
            st.download_button(
                "下載剪力圖 PNG",
                data=figure_to_png_bytes(figures["shear"]),
                file_name="shear_force.png",
                mime="image/png",
                key="download_shear_png",
            )

        st.subheader("各支承反力")
        reaction_df = pd.DataFrame(result["support_reactions"])[
            ["name", "type_text", "node_tag", "y", "Rx", "Ry", "Mz"]
        ].rename(
            columns={
                "name": "名稱",
                "type_text": "類型",
                "node_tag": "節點",
                "y": "Y (cm)",
                "Rx": "RX (kgf)",
                "Ry": "RY (kgf)",
                "Mz": "MZ (kgf-cm)",
            }
        )
        st.dataframe(
            reaction_df.style.format(
                {
                    "Y (cm)": "{:.3f}",
                    "RX (kgf)": "{:.3f}",
                    "RY (kgf)": "{:.3f}",
                    "MZ (kgf-cm)": "{:.3f}",
                }
            ),
            use_container_width=True,
            hide_index=True,
        )

        download_col1, download_col2 = st.columns(2)
        with download_col1:
            render_dataframe_download(
                stress_df,
                "下載應力檢核 CSV",
                "stress_checks.csv",
                "download_stress_csv",
            )
        with download_col2:
            render_dataframe_download(
                reaction_df,
                "下載支承反力 CSV",
                "support_reactions.csv",
                "download_reaction_csv",
            )


def main() -> None:
    st.set_page_config(
        page_title="直立多段梁／帷幕直料分析",
        page_icon="🏗️",
        layout="wide",
    )

    st.title("直立多段梁／帷幕直料分析")
    st.caption(
        "單位：kgf、cm。風載固定向右；"
        "Hinge 為鉸支承、Pin 為內部彎矩釋放、Free end 為自由端。"
    )

    if OPENSEES_IMPORT_ERROR is not None:
        st.error("目前環境缺少必要分析套件，請先安裝後重新啟動。")
        st.code(
            "python -m pip install -r requirements.txt\n"
            "python -m streamlit run Home.py",
            language="bash",
        )
        st.exception(OPENSEES_IMPORT_ERROR)
        st.stop()


    # --------------------------------------------------------
    # 左側：載重與材料斷面
    # --------------------------------------------------------
    st.sidebar.header("載重、材料與斷面")

    report_name_input = st.sidebar.text_input(
        "Word 報告檔名",
        value="",
        placeholder="mullion_analysis_report",
        help="未輸入時，預設使用 mullion_analysis_report.docx。",
    )

    load_input_mode = st.sidebar.radio(
        "風載輸入方式",
        ["風力線載重 (kgf/cm)", "風壓 (kPa)＋受風寬度 (cm)"],
        on_change=clear_analysis_result,
    )

    if load_input_mode == "風力線載重 (kgf/cm)":
        line_load = st.sidebar.number_input(
            "水平風力線載重 w (kgf/cm)",
            min_value=0.000001,
            value=5.7104,
            step=0.1,
            format="%.3f",
            on_change=clear_analysis_result,
        )
        w = float(line_load)
        load_metadata: Dict[str, object] = {
            "風載輸入方式": "風力線載重",
            "輸入風力線載重 (kgf/cm)": w,
        }
    else:
        wind_pressure = st.sidebar.number_input(
            "風壓 p (kPa)",
            min_value=0.000001,
            value=1.0,
            step=0.1,
            format="%.3f",
            on_change=clear_analysis_result,
        )
        tributary_width = st.sidebar.number_input(
            "受風寬度 B (cm)",
            min_value=0.000001,
            value=100.0,
            step=1.0,
            format="%.3f",
            on_change=clear_analysis_result,
        )
        w = float(wind_pressure) * float(tributary_width) * KPA_CM_TO_KGF_PER_CM
        st.sidebar.info(f"換算風力線載重 w = {w:.3f} kgf/cm")
        st.sidebar.caption(
            "單位換算：kgf/m² = kPa × 1000 ÷ 9.81。"
            f"例如：{float(wind_pressure):.3f} kPa ≈ "
            f"{float(wind_pressure) * 1000.0 / 9.81:.3f} kgf/m²。"
            "再乘受風寬度即可換算為線載重 kgf/cm。"
        )
        load_metadata = {
            "風載輸入方式": "風壓＋受風寬度",
            "輸入風壓 p (kPa)": float(wind_pressure),
            "受風寬度 B (cm)": float(tributary_width),
        }

    material_name = st.sidebar.selectbox(
        "材料種類",
        list(MATERIAL_DATABASE.keys()),
        on_change=clear_analysis_result,
    )
    selected_material = MATERIAL_DATABASE[material_name]
    material_grades = selected_material["grades"]
    assert isinstance(material_grades, dict)
    grade_name = st.sidebar.selectbox(
        "材料牌號",
        list(material_grades.keys()),
        on_change=clear_analysis_result,
    )
    selected_grade = material_grades[grade_name]
    assert isinstance(selected_grade, dict)
    st.sidebar.caption(
        f"E = {float(selected_material['E']):,.3f} kgf/cm²　｜　"
        f"Fb = {float(selected_grade['Fb_allow']):,.3f} kgf/cm²"
    )

    section_mode = st.sidebar.radio(
        "斷面輸入方式",
        ["整體直料", "區分公料與母料"],
        on_change=clear_analysis_result,
    )

    profiles: List[Dict[str, object]] = []
    if section_mode == "整體直料":
        profiles.append(
            {
                "name": "整體直料",
                "A": st.sidebar.number_input(
                    "總斷面積 A (cm²)",
                    min_value=0.000001,
                    value=20.0,
                    step=1.0,
                    format="%.3f",
                    key="whole_A",
                    on_change=clear_analysis_result,
                ),
                "Ix": st.sidebar.number_input(
                    "總強軸慣性矩 Ix (cm⁴)",
                    min_value=0.000001,
                    value=1000.0,
                    step=10.0,
                    format="%.3f",
                    key="whole_Ix",
                    on_change=clear_analysis_result,
                ),
                "Sx": st.sidebar.number_input(
                    "總強軸斷面係數 Sx (cm³)",
                    min_value=0.000001,
                    value=100.0,
                    step=1.0,
                    format="%.3f",
                    key="whole_Sx",
                    on_change=clear_analysis_result,
                ),
            }
        )
    else:
        st.sidebar.info("公料與母料的 Ix 必須已換算至共同中性軸。")
        profiles.extend(
            [
                {
                    "name": "公料",
                    "A": st.sidebar.number_input(
                        "公料 A1 (cm²)",
                        min_value=0.000001,
                        value=10.0,
                        step=1.0,
                        format="%.3f",
                        key="male_A",
                        on_change=clear_analysis_result,
                    ),
                    "Ix": st.sidebar.number_input(
                        "公料 Ix1 (cm⁴)",
                        min_value=0.000001,
                        value=500.0,
                        step=10.0,
                        format="%.3f",
                        key="male_Ix",
                        on_change=clear_analysis_result,
                    ),
                    "Sx": st.sidebar.number_input(
                        "公料 Sx1 (cm³)",
                        min_value=0.000001,
                        value=50.0,
                        step=1.0,
                        format="%.3f",
                        key="male_Sx",
                        on_change=clear_analysis_result,
                    ),
                },
                {
                    "name": "母料",
                    "A": st.sidebar.number_input(
                        "母料 A2 (cm²)",
                        min_value=0.000001,
                        value=10.0,
                        step=1.0,
                        format="%.3f",
                        key="female_A",
                        on_change=clear_analysis_result,
                    ),
                    "Ix": st.sidebar.number_input(
                        "母料 Ix2 (cm⁴)",
                        min_value=0.000001,
                        value=500.0,
                        step=10.0,
                        format="%.3f",
                        key="female_Ix",
                        on_change=clear_analysis_result,
                    ),
                    "Sx": st.sidebar.number_input(
                        "母料 Sx2 (cm³)",
                        min_value=0.000001,
                        value=50.0,
                        step=1.0,
                        format="%.3f",
                        key="female_Sx",
                        on_change=clear_analysis_result,
                    ),
                },
            ]
        )

    # --------------------------------------------------------
    # 主畫面最上方：簡支梁建議 Ix 參考表
    # --------------------------------------------------------
    st.header("建議 Ix 參考表")

    reference_span_lengths = [300.0, 350.0, 400.0, 450.0, 500.0, 550.0, 600.0]
    reference_e = float(selected_material["E"])

    reference_rows = []
    for reference_L in reference_span_lengths:
        reference_ix, reference_allowable, reference_formula = (
            recommended_ix_for_simple_span(
                uniform_load_kgf_per_cm=float(w),
                span_length_cm=reference_L,
                elastic_modulus_kgf_cm2=reference_e,
            )
        )
        reference_rows.append(
            {
                "簡支梁跨距 L (cm)": reference_L,
                "容許變形": reference_formula,
                "容許變形值 (cm)": reference_allowable,
                "建議最低 Ix (cm⁴)": reference_ix,
            }
        )

    reference_ix_df = pd.DataFrame(reference_rows)
    st.dataframe(
        round_dataframe_for_display(reference_ix_df),
        use_container_width=True,
        hide_index=True,
    )
    st.caption(
        f"參考條件：目前風力線載重 w = {float(w):.3f} kgf/cm，"
        f"E = {reference_e:,.3f} kgf/cm²。"
        "以單跨簡支梁承受全跨均布風載計算，並依各跨距之容許變形反算最低 Ix；"
        "僅供斷面初選參考，不取代實際多跨／懸臂模型分析與應力檢核。"
    )

    st.divider()

    # --------------------------------------------------------
    # 主畫面：幾何輸入
    # --------------------------------------------------------
    st.header("1. 控制點與各段長度")

    st.subheader("輸入方式說明")
    st.markdown(
        """
        <div style="font-size:1.05rem; line-height:1.85; margin-bottom:0.5rem;">
          <div><strong>1.</strong> 輸入跨數（L 有幾段）</div>
          <div><strong>2.</strong> 輸入最底部的支承形式，可能為鉸接（Hinge）或自由端（Free）</div>
          <div><strong>3.</strong> 輸入各跨長度以及上部控制點的支承形式，可能為鉸接（Hinge）或內部鉸接（Pin）</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    diagram_path = (
        Path(__file__).resolve().parents[1]
        / "assets"
        / "struc_pic.png"
    )
    if diagram_path.exists():
        diagram_col1, diagram_col2, diagram_col3 = st.columns([1, 0.72, 1])
        with diagram_col2:
            st.image(
                str(diagram_path),
                width=300,
            )
    else:
        st.warning(
            "找不到輸入示意圖：assets/mullion_input_diagram.png。"
        )

    st.divider()

    geometry_col1, geometry_col2 = st.columns([1, 1])
    with geometry_col1:
        number_of_segments = int(
            st.number_input(
                "直料總段數",
                min_value=1,
                max_value=30,
                value=3,
                step=1,
                on_change=clear_analysis_result,
            )
        )
    with geometry_col2:
        bottom_type = st.selectbox(
            "最下端控制點類型",
            ["hinge", "free"],
            format_func=station_format,
            on_change=clear_analysis_result,
        )

    segment_lengths: List[float] = []
    upper_types: List[str] = []
    current_lower_type = bottom_type

    for segment_index in range(1, number_of_segments + 1):
        is_last = segment_index == number_of_segments
        with st.container(border=True):
            st.markdown(f"**第 {segment_index} 段**")
            length_col, type_col = st.columns([1, 1.25])
            with length_col:
                default_length = 13.5 if segment_index % 2 == 0 else 386.5
                length = st.number_input(
                    f"L{segment_index} 長度 (cm)",
                    min_value=0.000001,
                    value=float(default_length),
                    step=1.0,
                    format="%.3f",
                    key=f"segment_length_{segment_index}",
                    on_change=clear_analysis_result,
                )
                segment_lengths.append(float(length))

            with type_col:
                widget_key = f"upper_type_{segment_index}"

                if is_last:
                    if current_lower_type == "pin":
                        upper_type = "hinge"
                        st.info("下端為 Pin，最上端自動設為 Hinge。")
                    else:
                        valid_types = ["hinge", "free"]
                        if (
                            widget_key not in st.session_state
                            or st.session_state[widget_key] not in valid_types
                        ):
                            st.session_state[widget_key] = "hinge"
                        upper_type = st.selectbox(
                            "上端類型（整根直料最上端）",
                            valid_types,
                            format_func=station_format,
                            key=widget_key,
                            on_change=clear_analysis_result,
                        )
                else:
                    if current_lower_type == "free":
                        upper_type = "hinge"
                        st.info("由下部 Free end 起始，上端自動設為 Hinge。")
                    else:
                        valid_types = ["hinge", "pin"]
                        if (
                            widget_key not in st.session_state
                            or st.session_state[widget_key] not in valid_types
                        ):
                            st.session_state[widget_key] = "hinge"
                        upper_type = st.selectbox(
                            "上端控制點類型",
                            valid_types,
                            format_func=station_format,
                            key=widget_key,
                            on_change=clear_analysis_result,
                        )

                upper_types.append(upper_type)
                st.caption(
                    f"本段：{station_type_text(current_lower_type)} → "
                    f"{station_type_text(upper_type)}"
                )
                current_lower_type = upper_type

    geometry_error: Optional[str] = None
    try:
        stations, input_segments = build_station_geometry(
            bottom_type,
            segment_lengths,
            upper_types,
        )
    except (ValueError, RuntimeError) as exc:
        geometry_error = str(exc)
        stations = []
        input_segments = []
        st.warning(geometry_error)
    else:
        st.subheader("幾何輸入預覽")
        preview_col1, preview_col2 = st.columns(2)
        with preview_col1:
            preview_station_df = pd.DataFrame(
                [
                    {
                        "名稱": item["name"],
                        "類型": station_type_text(str(item["type"])),
                        "Y (cm)": float(item["y"]),
                    }
                    for item in stations
                ]
            )
            st.dataframe(
                preview_station_df,
                use_container_width=True,
                hide_index=True,
            )
        with preview_col2:
            preview_segment_df = pd.DataFrame(
                [
                    {
                        "段號": item["segment_number"],
                        "連接類型": item["connection_type"],
                        "長度 (cm)": float(item["length"]),
                        "Y 範圍 (cm)": (
                            f"{float(item['start_y']):.3f} → "
                            f"{float(item['end_y']):.3f}"
                        ),
                    }
                    for item in input_segments
                ]
            )
            st.dataframe(
                preview_segment_df,
                use_container_width=True,
                hide_index=True,
            )

    st.header("2. 執行分析")
    analyze_clicked = st.button(
        "執行分析",
        type="primary",
        use_container_width=True,
        disabled=geometry_error is not None,
    )

    if analyze_clicked:
        clear_analysis_result()
        try:
            with st.spinner("正在建立模型、分析並產生圖表……"):
                result = run_analysis(
                    w=float(w),
                    load_metadata=load_metadata,
                    stations=stations,
                    input_segments=input_segments,
                    material_name=material_name,
                    grade_name=grade_name,
                    profiles=profiles,
                )
        except Exception as exc:
            try:
                if ops is not None:
                    ops.wipe()
            except Exception:
                pass
            st.error(f"分析失敗：{exc}")
            with st.expander("顯示完整錯誤資訊"):
                st.exception(exc)
        else:
            st.session_state["mullion_analysis_result"] = result

    stored_result = st.session_state.get("mullion_analysis_result")
    if isinstance(stored_result, dict):
        render_results(stored_result, report_name_input)


if __name__ == "__main__":
    main()
